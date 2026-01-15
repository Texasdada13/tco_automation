"""
Create additional test proposals to reach target of 10+ documents.
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def create_fis_proposal_3():
    """Create third FIS-style proposal - Summit Bank (complex)"""
    doc = Document()

    title = doc.add_heading('FIS Enterprise Banking Solutions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Comprehensive Proposal for Summit National Bank')
    doc.add_paragraph('Vendor: FIS Global')
    doc.add_paragraph('Date: February 2025')
    doc.add_paragraph('Asset Size: $1.2 Billion')
    doc.add_paragraph('')

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'FIS Global presents this enterprise-grade proposal for Summit National Bank. '
        'As a $1.2B institution, Summit requires robust core processing with advanced '
        'digital capabilities and comprehensive treasury management.'
    )

    # 5-Year Term Pricing
    doc.add_heading('5-Year Contract Pricing', level=1)

    table1 = doc.add_table(rows=10, cols=5)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Solution'
    hdr[1].text = 'Category'
    hdr[2].text = 'Monthly'
    hdr[3].text = 'Year 1'
    hdr[4].text = '5-Year Total'

    items_5yr = [
        ('HORIZON Enterprise Core', 'Bundle', '$28,500.00', '$342,000.00', '$1,891,650.00'),
        ('Digital One Enterprise', 'Bundle', '$14,200.00', '$170,400.00', '$942,912.00'),
        ('Payments One Complete', 'Bundle', '$9,800.00', '$117,600.00', '$650,664.00'),
        ('ImageCentre Pro', 'Bundle', '$6,400.00', '$76,800.00', '$424,896.00'),
        ('Enterprise Analytics', 'Bundle', '$4,500.00', '$54,000.00', '$298,620.00'),
        ('Security Suite Premium', 'Bundle', '$3,200.00', '$38,400.00', '$212,352.00'),
        ('API Gateway Enterprise', 'Bundle', '$2,800.00', '$33,600.00', '$185,808.00'),
        ('Compliance Manager Pro', 'Bundle', '$2,100.00', '$25,200.00', '$139,356.00'),
        ('5-YEAR BUNDLE TOTAL', '', '$71,500.00', '$858,000.00', '$4,746,258.00'),
    ]
    for i, item in enumerate(items_5yr, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # 7-Year Term Pricing
    doc.add_heading('7-Year Contract Pricing (Recommended)', level=1)

    table2 = doc.add_table(rows=10, cols=5)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Solution'
    hdr[1].text = 'Category'
    hdr[2].text = 'Monthly'
    hdr[3].text = 'Year 1'
    hdr[4].text = '7-Year Total'

    items_7yr = [
        ('HORIZON Enterprise Core', 'Bundle', '$24,500.00', '$294,000.00', '$2,358,972.00'),
        ('Digital One Enterprise', 'Bundle', '$12,200.00', '$146,400.00', '$1,174,512.00'),
        ('Payments One Complete', 'Bundle', '$8,400.00', '$100,800.00', '$808,416.00'),
        ('ImageCentre Pro', 'Bundle', '$5,500.00', '$66,000.00', '$529,320.00'),
        ('Enterprise Analytics', 'Bundle', '$3,800.00', '$45,600.00', '$365,712.00'),
        ('Security Suite Premium', 'Bundle', '$2,700.00', '$32,400.00', '$259,848.00'),
        ('API Gateway Enterprise', 'Bundle', '$2,400.00', '$28,800.00', '$230,976.00'),
        ('Compliance Manager Pro', 'Bundle', '$1,800.00', '$21,600.00', '$173,232.00'),
        ('7-YEAR BUNDLE TOTAL', '', '$61,300.00', '$735,600.00', '$5,900,988.00'),
    ]
    for i, item in enumerate(items_7yr, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Non-Bundle Required
    doc.add_heading('Non-Bundle Required Services', level=1)

    table3 = doc.add_table(rows=12, cols=4)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Base'
    hdr[2].text = 'Per Transaction'
    hdr[3].text = 'Category'

    nb_items = [
        ('ACH Origination Enterprise', '$1,450.00', '$0.065/item', 'Non-Bundle Required'),
        ('ACH Receipt Processing', '$425.00', '$0.04/item', 'Non-Bundle Required'),
        ('Wire Transfer - Domestic', '$850.00', '$16.00/wire', 'Non-Bundle Required'),
        ('Wire Transfer - International', '$650.00', '$38.00/wire', 'Non-Bundle Required'),
        ('Remote Deposit Capture', '$1,100.00', '$0.06/deposit', 'Non-Bundle Required'),
        ('Positive Pay Enterprise', '$725.00', '$0.04/item', 'Non-Bundle Required'),
        ('Account Reconciliation Pro', '$575.00', '-', 'Non-Bundle Required'),
        ('Mobile Banking Enterprise', '$1,850.00', '-', 'Non-Bundle Required'),
        ('Bill Pay Platform', '$625.00', '$0.35/payment', 'Non-Bundle Required'),
        ('eStatements Enterprise', '$480.00', '$0.12/statement', 'Non-Bundle Required'),
        ('NON-BUNDLE TOTAL', '$8,730.00', '', ''),
    ]
    for i, item in enumerate(nb_items, 1):
        row = table3.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Optional Services
    doc.add_heading('Non-Bundle Optional Services', level=1)

    table4 = doc.add_table(rows=10, cols=3)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Fee'
    hdr[2].text = 'Category'

    opt_items = [
        ('Commercial Cash Management Suite', '$2,400.00', 'Non-Bundle Optional'),
        ('Lockbox Services Enterprise', '$2,800.00', 'Non-Bundle Optional'),
        ('Controlled Disbursement', '$750.00', 'Non-Bundle Optional'),
        ('Sweep Services', '$425.00', 'Non-Bundle Optional'),
        ('Fraud Detection AI', '$950.00', 'Non-Bundle Optional'),
        ('Custom API Development', '$1,200.00', 'Non-Bundle Optional'),
        ('Advanced Reporting Suite', '$650.00', 'Non-Bundle Optional'),
        ('Extended Support 24/7', '$575.00', 'Non-Bundle Optional'),
        ('OPTIONAL TOTAL (if all)', '$9,750.00', ''),
    ]
    for i, item in enumerate(opt_items, 1):
        row = table4.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Third-Party Services
    doc.add_heading('Third-Party Services', level=1)

    table5 = doc.add_table(rows=7, cols=3)
    table5.style = 'Table Grid'
    hdr = table5.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly/Annual'
    hdr[2].text = 'Category'

    tp_items = [
        ('Enterprise Hosting (AWS)', '$5,200.00/mo', 'Third-Party Required'),
        ('Network Security Suite', '$2,800.00/mo', 'Third-Party Required'),
        ('Disaster Recovery', '$1,650.00/mo', 'Third-Party Required'),
        ('Regulatory Compliance', '$1,100.00/mo', 'Third-Party Required'),
        ('Cyber Insurance', '$24,000.00/yr', 'Third-Party Required'),
        ('Penetration Testing', '$18,000.00/yr', 'Third-Party Optional'),
    ]
    for i, item in enumerate(tp_items, 1):
        row = table5.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # One-Time Fees
    doc.add_heading('Implementation Investment', level=1)

    table6 = doc.add_table(rows=9, cols=2)
    table6.style = 'Table Grid'
    hdr = table6.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'One-Time Fee'

    impl_items = [
        ('Enterprise Core Implementation', '$185,000.00'),
        ('Data Migration & Conversion', '$72,000.00'),
        ('Integration Development', '$48,000.00'),
        ('Training Program (Comprehensive)', '$35,000.00'),
        ('Project Management', '$38,000.00'),
        ('Custom Configuration', '$28,000.00'),
        ('Testing & QA', '$22,000.00'),
        ('Enterprise Customer Credit', '($65,000.00)'),
    ]
    for i, item in enumerate(impl_items, 1):
        row = table6.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('Net Implementation: $363,000.00')

    # Terms
    doc.add_heading('Contract Terms', level=1)
    doc.add_paragraph('Recommended Term: 7 Years')
    doc.add_paragraph('Bundle CPI: 6% annually from Year 2')
    doc.add_paragraph('Non-Bundle CPI: 3% annually from Year 2')
    doc.add_paragraph('Early Termination: 80% of remaining contract value')
    doc.add_paragraph('SLA: 99.99% uptime guarantee')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'fis-like', 'Summit_FIS_Enterprise_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_jack_henry_proposal_2():
    """Create second Jack Henry proposal - Lakeview CU"""
    wb = openpyxl.Workbook()

    ws1 = wb.active
    ws1.title = 'Summary'
    ws1['A1'] = 'Jack Henry Symitar Proposal'
    ws1['A1'].font = Font(bold=True, size=14)
    ws1['A2'] = 'Lakeview Credit Union'
    ws1['A3'] = 'Asset Size: $280M'
    ws1['A4'] = 'Members: 32,000'
    ws1['A5'] = 'Date: January 2025'

    # Detailed pricing sheet
    ws2 = wb.create_sheet('Proposal_1')

    headers = ['Product', 'Family', 'Type', 'License', 'Install', 'Monthly', 'Maint', 'Included']
    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    products = [
        ('Episys Core', 'Core', 'Bundle', 62000, 18000, 6200, 11000, 'Y'),
        ('PowerOn', 'Core', 'Bundle', 9500, 4000, 950, 1900, 'Y'),
        ('SymConnect', 'Core', 'Bundle', 7500, 3000, 750, 1500, 'Y'),
        ('Member Connect Online', 'Digital', 'Bundle', 14000, 6500, 1650, 2800, 'Y'),
        ('Banno Digital Platform', 'Digital', 'Bundle', 18000, 8000, 2600, 3400, 'Y'),
        ('Banno Mobile', 'Digital', 'Bundle', 12000, 5000, 1900, 2400, 'Y'),
        ('ACH Services', 'Payments', 'Required', 0, 2000, 520, 950, 'Y'),
        ('Wire Services', 'Payments', 'Required', 0, 1200, 320, 600, 'Y'),
        ('Bill Pay', 'Payments', 'Required', 0, 2500, 425, 850, 'Y'),
        ('Remote Deposit', 'Imaging', 'Required', 0, 3500, 475, 900, 'Y'),
        ('eStatements', 'Imaging', 'Required', 0, 1800, 285, 550, 'Y'),
        ('Card Processing', 'Cards', 'Optional', 6500, 4000, 780, 1400, 'N'),
        ('Loan Origination', 'Lending', 'Optional', 12000, 6000, 1150, 2300, 'N'),
        ('Document Imaging', 'Imaging', 'Optional', 7500, 3800, 700, 1350, 'N'),
        ('Business Intelligence', 'Analytics', 'Optional', 9500, 5000, 900, 1800, 'N'),
        ('Symitar EASE Hosting', 'Hosting', 'Required', 0, 12000, 3600, 6500, 'Y'),
        ('Security Services', 'Hosting', 'Required', 0, 6500, 1450, 2800, 'Y'),
    ]

    for row_num, product in enumerate(products, 2):
        for col_num, value in enumerate(product, 1):
            cell = ws2.cell(row=row_num, column=col_num, value=value)
            if col_num in [4, 5, 6, 7]:
                cell.number_format = '$#,##0.00'

    # Totals
    total_row = 2 + len(products)
    ws2.cell(row=total_row, column=1, value='TOTALS').font = Font(bold=True)

    for ws in [ws1, ws2]:
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 12

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'jack-henry-like', 'Lakeview_CU_JackHenry_2025.xlsx')
    wb.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_csi_proposal_2():
    """Create second CSI proposal - Commerce Bank"""
    doc = Document()

    title = doc.add_heading('CSI Core Banking Services', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Prepared for: Commerce First Bank')
    doc.add_paragraph('Prepared by: Computer Services, Inc.')
    doc.add_paragraph('Date: February 2025')
    doc.add_paragraph('Asset Size: $650 Million')
    doc.add_paragraph('')

    doc.add_heading('NuPoint Platform Overview', level=1)
    doc.add_paragraph(
        'CSI proposes our award-winning NuPoint cloud core banking platform '
        'for Commerce First Bank. This solution provides modern architecture '
        'with open APIs and comprehensive digital banking capabilities.'
    )

    # Core Platform
    doc.add_heading('Core Platform Pricing', level=1)

    table1 = doc.add_table(rows=9, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Category'
    hdr[2].text = 'Monthly'
    hdr[3].text = 'Type'

    core = [
        ('NuPoint Core Engine', 'Bundle', '$18,500.00', 'Monthly F'),
        ('Digital Banking Platform', 'Bundle', '$8,200.00', 'Monthly F'),
        ('Mobile Banking Suite', 'Bundle', '$4,800.00', 'Monthly F'),
        ('Document Management', 'Bundle', '$3,600.00', 'Monthly F'),
        ('Reporting & Analytics', 'Bundle', '$2,900.00', 'Monthly F'),
        ('Security & Compliance', 'Bundle', '$2,200.00', 'Monthly F'),
        ('Integration Layer', 'Bundle', '$1,600.00', 'Monthly F'),
        ('BUNDLE TOTAL', '', '$41,800.00', ''),
    ]
    for i, item in enumerate(core, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Required Services
    doc.add_heading('Required Transaction Services', level=1)

    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Base Monthly'
    hdr[2].text = 'Per-Item'
    hdr[3].text = 'Category'

    req = [
        ('ACH Processing', '$820.00', '$0.055/item', 'Non-Bundle Required'),
        ('Wire Services', '$450.00', '$12.50/wire', 'Non-Bundle Required'),
        ('Remote Deposit', '$625.00', '$0.07/dep', 'Non-Bundle Required'),
        ('Positive Pay', '$385.00', '$0.04/item', 'Non-Bundle Required'),
        ('Bill Pay', '$340.00', '$0.28/pmt', 'Non-Bundle Required'),
        ('eStatements', '$275.00', '$0.09/stmt', 'Non-Bundle Required'),
        ('REQUIRED TOTAL', '$2,895.00', '', ''),
    ]
    for i, item in enumerate(req, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Optional
    doc.add_heading('Optional Enhancements', level=1)

    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly'
    hdr[2].text = 'Category'

    opt = [
        ('Treasury Management Suite', '$1,650.00', 'Non-Bundle Optional'),
        ('Commercial Lockbox', '$1,350.00', 'Non-Bundle Optional'),
        ('International Services', '$520.00', 'Non-Bundle Optional'),
        ('Advanced Fraud AI', '$580.00', 'Non-Bundle Optional'),
        ('API Developer Access', '$875.00', 'Non-Bundle Optional'),
        ('OPTIONAL TOTAL', '$4,975.00', ''),
    ]
    for i, item in enumerate(opt, 1):
        row = table3.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Implementation
    doc.add_heading('Implementation', level=1)

    table4 = doc.add_table(rows=7, cols=2)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Item'
    hdr[1].text = 'Fee'

    impl = [
        ('Core Implementation', '$145,000.00'),
        ('Data Migration', '$48,000.00'),
        ('Training', '$18,500.00'),
        ('Project Management', '$22,000.00'),
        ('Integrations', '$28,000.00'),
        ('Signing Incentive', '($35,000.00)'),
    ]
    for i, item in enumerate(impl, 1):
        row = table4.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('Net Implementation: $226,500.00')

    doc.add_heading('Terms', level=1)
    doc.add_paragraph('Recommended Term: 7 Years')
    doc.add_paragraph('Annual Increase: 4% Bundle, 3% Services')
    doc.add_paragraph('Hosting: CSI Managed Cloud')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'csi-like', 'Commerce_First_CSI_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_finastra_proposal():
    """Create Finastra Fusion proposal"""
    doc = Document()

    title = doc.add_heading('Finastra Fusion Banking Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Proposal for: Gateway Community Bank')
    doc.add_paragraph('Prepared by: Finastra')
    doc.add_paragraph('Date: January 2025')
    doc.add_paragraph('')

    doc.add_heading('Fusion Platform Services', level=1)

    table1 = doc.add_table(rows=8, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Solution'
    hdr[1].text = 'Category'
    hdr[2].text = 'Monthly'
    hdr[3].text = 'Type'

    items = [
        ('Fusion Core Banking', 'Bundle', '$22,400.00', 'Monthly F'),
        ('Fusion Digital Banking', 'Bundle', '$9,800.00', 'Monthly F'),
        ('Fusion Mobile', 'Bundle', '$5,600.00', 'Monthly F'),
        ('Fusion Payments Hub', 'Bundle', '$7,200.00', 'Monthly F'),
        ('Fusion Document Mgmt', 'Bundle', '$4,100.00', 'Monthly F'),
        ('Fusion Analytics', 'Bundle', '$3,400.00', 'Monthly F'),
        ('BUNDLE TOTAL', '', '$52,500.00', ''),
    ]
    for i, item in enumerate(items, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_heading('Transaction Services', level=1)

    table2 = doc.add_table(rows=7, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Base'
    hdr[2].text = 'Per-Item'
    hdr[3].text = 'Category'

    txn = [
        ('ACH Processing', '$980.00', '$0.072/item', 'Non-Bundle Required'),
        ('Wire Services', '$580.00', '$14.50/wire', 'Non-Bundle Required'),
        ('Remote Deposit', '$720.00', '-', 'Non-Bundle Required'),
        ('Positive Pay', '$425.00', '-', 'Non-Bundle Required'),
        ('Bill Pay', '$380.00', '$0.32/pmt', 'Non-Bundle Required'),
        ('SERVICES TOTAL', '$3,085.00', '', ''),
    ]
    for i, item in enumerate(txn, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_heading('Implementation', level=1)
    doc.add_paragraph('Core Implementation: $165,000.00')
    doc.add_paragraph('Data Migration: $55,000.00')
    doc.add_paragraph('Training: $22,000.00')
    doc.add_paragraph('Project Management: $28,000.00')
    doc.add_paragraph('New Customer Credit: ($45,000.00)')
    doc.add_paragraph('Net Implementation: $225,000.00')

    doc.add_heading('Terms', level=1)
    doc.add_paragraph('Contract Term: 7 Years')
    doc.add_paragraph('Annual Increase: 5% Bundle, 3% Services')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'other-vendors', 'Gateway_Finastra_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def main():
    """Create additional proposals"""
    print('Creating additional test proposals...')
    print('=' * 60)

    proposals = []

    print('\nCreating additional FIS proposal...')
    proposals.append(create_fis_proposal_3())

    print('\nCreating additional Jack Henry proposal...')
    proposals.append(create_jack_henry_proposal_2())

    print('\nCreating additional CSI proposal...')
    proposals.append(create_csi_proposal_2())

    print('\nCreating Finastra proposal...')
    proposals.append(create_finastra_proposal())

    print('\n' + '=' * 60)
    print(f'Created {len(proposals)} additional proposals')


if __name__ == '__main__':
    main()
