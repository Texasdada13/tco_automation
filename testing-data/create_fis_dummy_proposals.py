"""
Create dummy FIS proposals based on the ORIGINAL Liberty Capital Bank FIS proposal structure.
These proposals have the same data points and similar values for validation testing.
"""

import os
import sys
import random
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Original FIS proposal data structure (from Liberty Capital Bank)
ORIGINAL_LINE_ITEMS = [
    # Solution Name, Category, Current Pricing, New Pricing
    ("Core: HORIZON", "Core Banking", 11266, 16792),
    ("Digital: D1 Flex, D1 Business, Mobile, Bill Pay", "Digital Banking", 15184, 35253),
    ("Digital: Website Services", "Digital Banking", 568, 568),
    ("Digital: Zelle", "Digital Banking", 641, 641),
    ("EFT: EFT Norcross, 3D Secure, Tokenization, SecurLOCK", "Electronic Funds Transfer", 1736, 6675),
    ("Risk, Fraud & Compliance: Decision Solutions", "Risk Management", 803, 803),
    ("Treasury: eWire (Business)", "Treasury Management", 2382, 2382),
    ("Image Solutions: RTP (Receive)", "Payment Processing", 855, 855),
    ("Item Processing: FCM, IP, Branch Capture, CCX, FXD, DLRR", "Item Processing", 8453, 19353),
    ("FOS: eDelivery, Print/Render, Forms & Envelopes", "Forms and Output", 1168, 2168),
    ("Lending: FLO", "Lending", 1055, 1055),
    ("Lending: CRA Wiz/Fair Lending", "Lending", 1127, 1127),
    ("Information Services: Regulatory Compliance", "Information Services", 5148, 5148),
    ("Information Services: IRS Reporting", "Information Services", 2100, 2100),
]

# Additional line items with one-time fees
ORIGINAL_ONETIME_ITEMS = [
    # Solution Name, One-Time, Monthly
    ("IP: DirectLink Merchant (via RegO)", 21943, 3402.14),
    ("IP: DirectLink Consumer (via FIS)", 25551, 1170.19),
    ("HORIZON: Conditional Processing", 15000, 1076.80),
    ("Treasury: eWire (Consumer)", 1000, 5.00),  # per wire
    ("Treasury: Extended Account Analysis", 65000, 1124.00),
    ("Image Solutions: RTP Send for Business", 7630, 150.50),
    ("Image Solutions: FedNOW (both Send/Receive)", 12500, 3550.00),
    ("Image Solutions: Chargeback Manager", 6064, 250.00),
]

# Relationship credits from original
ORIGINAL_CREDITS = [
    ("Exit Fee Assistance Credit", 750000),
    ("Conversion Cost Credit", 250000),
    ("Contract Renewal Credit", 150000),
    ("Monthly Relationship Credit", 35000),  # per month
    ("One-Time Service Investment Credit", 50000),
]


def vary_price(base_price, variance_pct=0.15):
    """Add random variance to a price."""
    if base_price == 0:
        return 0
    variance = base_price * variance_pct
    return round(base_price + random.uniform(-variance, variance), 2)


def create_fis_dummy_proposal(
    bank_name: str,
    bank_location: str,
    output_filename: str,
    price_multiplier: float = 1.0,
    include_acquisition: bool = True,
    contract_term: int = 7
):
    """
    Create a dummy FIS proposal matching the original structure.

    Args:
        bank_name: Name of the fictional bank
        bank_location: Location
        output_filename: Output file name
        price_multiplier: Scale prices up/down (0.7-1.3 for realism)
        include_acquisition: Include acquisition-related items
        contract_term: Contract term in years
    """
    doc = Document()

    # === HEADER ===
    title = doc.add_paragraph(bank_name)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)

    loc = doc.add_paragraph(bank_location)
    loc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph("FIS Investment Summary")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)

    if include_acquisition:
        doc.add_paragraph("Renewal w/ Acquisition").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Contract Renewal").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("January 2025 (presented to Arriba Advisors)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # === INTRODUCTION ===
    doc.add_paragraph("We are honored to be your bank's technology partner.")
    doc.add_paragraph(
        f"As you review our FIS Investment Summary, please note that we have tried to "
        f"make this easy, while trying to be excellent partners, in terms of pricing."
    )
    doc.add_paragraph()
    doc.add_paragraph("Below are specific details of our proposal:")
    doc.add_paragraph()

    # === FINANCIALS SECTION ===
    doc.add_heading("Financials", level=1)

    if include_acquisition:
        doc.add_paragraph(
            f"If FIS contract renewal and Statement of Work (SOW) are both signed by "
            f"March 31, 2025, FIS is pleased to offer the following:"
        )
        doc.add_paragraph()

        # Relationship credits
        for credit_name, credit_amount in ORIGINAL_CREDITS:
            varied_amount = vary_price(credit_amount * price_multiplier, 0.1)
            if "Monthly" in credit_name:
                doc.add_paragraph(f"• ${varied_amount:,.0f} {credit_name.lower()} for term of contract")
            else:
                doc.add_paragraph(f"• ${varied_amount:,.0f} {credit_name.lower()}")

    doc.add_paragraph()

    # === TERM SECTION ===
    doc.add_heading("Term", level=1)
    doc.add_paragraph(
        f"Term of this agreement will be {contract_term} ({contract_term}) years. "
        f"All solutions will be coterminous, all pricing reflects HORIZON core contract term."
    )
    doc.add_paragraph()

    # === MAIN PRICING TABLE ===
    doc.add_heading("Current Monthly Solutions", level=1)

    table1 = doc.add_table(rows=len(ORIGINAL_LINE_ITEMS) + 2, cols=3)
    table1.style = 'Table Grid'

    # Headers
    hdr = table1.rows[0].cells
    hdr[0].text = "Solution"
    hdr[1].text = "Current Pricing\n(as of April 2024)"
    hdr[2].text = "New Pricing\n(Combined Organization)"

    # Make headers bold
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    total_current = 0
    total_new = 0

    for i, (solution, category, current, new) in enumerate(ORIGINAL_LINE_ITEMS, 1):
        row = table1.rows[i].cells
        row[0].text = solution

        varied_current = vary_price(current * price_multiplier)
        varied_new = vary_price(new * price_multiplier)

        row[1].text = f"${varied_current:,.0f}"
        row[2].text = f"${varied_new:,.0f}"

        total_current += varied_current
        total_new += varied_new

    # Total row
    total_row = table1.rows[len(ORIGINAL_LINE_ITEMS) + 1].cells
    total_row[0].text = "Monthly Subtotal"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[1].text = f"${total_current:,.0f}"
    total_row[2].text = f"${total_new:,.0f}"

    doc.add_paragraph()

    # === NEW SOLUTIONS TABLE ===
    doc.add_heading("New/Optional Solutions", level=1)

    table2 = doc.add_table(rows=len(ORIGINAL_ONETIME_ITEMS) + 1, cols=4)
    table2.style = 'Table Grid'

    # Headers
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Solution"
    hdr2[1].text = "One-Time Investment"
    hdr2[2].text = "Monthly"
    hdr2[3].text = "Notes"

    for cell in hdr2:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for i, (solution, onetime, monthly) in enumerate(ORIGINAL_ONETIME_ITEMS, 1):
        row = table2.rows[i].cells
        row[0].text = solution
        row[1].text = f"${vary_price(onetime * price_multiplier):,.0f}"
        row[2].text = f"${vary_price(monthly * price_multiplier):,.2f}"
        row[3].text = "Optional" if "Optional" in solution or random.random() > 0.7 else "Recommended"

    doc.add_paragraph()

    # === EXIT FEES SECTION ===
    doc.add_heading("Capped Exit and Acquisition Fees", level=1)
    doc.add_paragraph(
        "In the event your bank decides to leave FIS, for any reason, prior to your "
        "contract expiration date; FIS will cap your bank's liquidated damages (LDs)."
    )
    doc.add_paragraph("FIS will agree to cap your bank's deconversion fees, as follows:")
    doc.add_paragraph("• Under $1B ($300,000 fee)")
    doc.add_paragraph("• Over $1B (custom quote required)")
    doc.add_paragraph()
    doc.add_paragraph("Please note, this does not include any third-party fees or image conversion fees.")

    # Save
    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'fis-like', output_filename)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


def main():
    """Create multiple dummy FIS proposals."""
    print("Creating FIS dummy proposals based on original structure...")
    print("=" * 60)

    # Ensure directory exists
    os.makedirs(os.path.join(BASE_DIR, 'dummy-proposals', 'fis-like'), exist_ok=True)

    # Set random seed for reproducibility
    random.seed(42)

    proposals = []

    # Dummy Proposal 1: Similar size bank with acquisition
    proposals.append(create_fis_dummy_proposal(
        bank_name="First National Bank of Texas",
        bank_location="Austin, Texas",
        output_filename="First_National_TX_FIS_Investment_Summary_2025.docx",
        price_multiplier=1.1,  # Slightly larger
        include_acquisition=True,
        contract_term=7
    ))

    # Dummy Proposal 2: Smaller community bank, no acquisition
    proposals.append(create_fis_dummy_proposal(
        bank_name="Community State Bank",
        bank_location="Dallas, Texas",
        output_filename="Community_State_Bank_FIS_Investment_Summary_2025.docx",
        price_multiplier=0.75,  # Smaller bank
        include_acquisition=False,
        contract_term=5
    ))

    # Dummy Proposal 3: Larger regional bank
    proposals.append(create_fis_dummy_proposal(
        bank_name="Southwest Regional Bank",
        bank_location="Houston, Texas",
        output_filename="Southwest_Regional_FIS_Investment_Summary_2025.docx",
        price_multiplier=1.4,  # Larger
        include_acquisition=True,
        contract_term=7
    ))

    # Dummy Proposal 4: Mid-size bank
    proposals.append(create_fis_dummy_proposal(
        bank_name="Heritage Bank & Trust",
        bank_location="San Antonio, Texas",
        output_filename="Heritage_Bank_Trust_FIS_Investment_Summary_2025.docx",
        price_multiplier=0.9,
        include_acquisition=False,
        contract_term=6
    ))

    # Dummy Proposal 5: Similar to original
    proposals.append(create_fis_dummy_proposal(
        bank_name="Lone Star National Bank",
        bank_location="Fort Worth, Texas",
        output_filename="Lone_Star_National_FIS_Investment_Summary_2025.docx",
        price_multiplier=1.0,  # Same size as original
        include_acquisition=True,
        contract_term=7
    ))

    print()
    print("=" * 60)
    print(f"Created {len(proposals)} FIS dummy proposals")
    print()
    print("These proposals have:")
    print("  - Same solution names as original (Core: HORIZON, Digital: D1 Flex, etc.)")
    print("  - Same table structure (Current Pricing vs New Pricing)")
    print("  - Same categories (Core Banking, Digital Banking, EFT, etc.)")
    print("  - Similar pricing values with realistic variance")
    print("  - Same one-time fee structure")
    print("  - Same relationship credits format")


if __name__ == '__main__':
    main()
