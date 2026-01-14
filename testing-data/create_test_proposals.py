"""
Create realistic test proposals for TCO extraction pipeline validation.
Based on research of actual banking vendor pricing structures.
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def create_fis_proposal_1():
    """Create FIS-style Word proposal - Riverside Community Bank"""
    doc = Document()

    # Title
    title = doc.add_heading('FIS Core Banking Services Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    doc.add_paragraph('Prepared for: Riverside Community Bank')
    doc.add_paragraph('Vendor: FIS Global')
    doc.add_paragraph('Date: January 2025')
    doc.add_paragraph('Proposal Valid Through: April 2025')
    doc.add_paragraph('')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'FIS Global is pleased to present this comprehensive proposal for core banking '
        'and digital services to Riverside Community Bank. This proposal includes our '
        'HORIZON Core Processing platform, Digital One digital banking suite, Payments One '
        'payment processing, and ImageCentre document management solutions.'
    )

    # Bundle Pricing
    doc.add_heading('Pricing Summary - Bundle Options', level=1)

    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Contract Term'
    hdr[1].text = 'Monthly Bundle Fee'
    hdr[2].text = 'Annual Total'

    data = [
        ('5-Year Term', '$45,250.00', '$543,000.00'),
        ('7-Year Term', '$41,500.00', '$498,000.00'),
        ('10-Year Term', '$37,750.00', '$453,000.00'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(row_data):
            row[j].text = val

    doc.add_paragraph('')
    doc.add_paragraph('Note: Bundle pricing includes annual CPI adjustment of 6% beginning Year 2.')

    # HORIZON Bundle
    doc.add_heading('Section 1: HORIZON Core Processing Bundle', level=1)

    table2 = doc.add_table(rows=7, cols=5)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Solution Name'
    hdr[1].text = 'Category'
    hdr[2].text = 'Fee Type'
    hdr[3].text = 'Monthly Fee'
    hdr[4].text = 'Per-Item Rate'

    bundle_items = [
        ('HORIZON Core Processing', 'Bundle', 'Monthly F', '$16,500.00', '-'),
        ('Digital One Platform', 'Bundle', 'Monthly F', '$9,250.00', '-'),
        ('Payments One Suite', 'Bundle', 'Monthly F', '$6,750.00', '-'),
        ('ImageCentre', 'Bundle', 'Monthly F', '$4,750.00', '-'),
        ('Data Analytics & Reporting', 'Bundle', 'Monthly F', '$2,500.00', '-'),
        ('Security & Compliance Module', 'Bundle', 'Monthly F', '$1,750.00', '-'),
    ]
    for i, item in enumerate(bundle_items, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('')
    doc.add_paragraph('Bundle Subtotal: $41,500.00/month (7-Year Term)')

    # Non-Bundle Required
    doc.add_heading('Section 2: Non-Bundle Required Services', level=1)

    table3 = doc.add_table(rows=9, cols=5)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Solution Name'
    hdr[1].text = 'Category'
    hdr[2].text = 'Fee Type'
    hdr[3].text = 'Monthly Fee'
    hdr[4].text = 'Per-Item Rate'

    nb_required = [
        ('ACH Origination', 'Non-Bundle Required', 'Monthly V', '$925.00', '$0.09/item'),
        ('Wire Transfer Services', 'Non-Bundle Required', 'Monthly V', '$550.00', '$18.00/wire'),
        ('Remote Deposit Capture', 'Non-Bundle Required', 'Monthly F', '$725.00', '-'),
        ('Positive Pay', 'Non-Bundle Required', 'Monthly F', '$475.00', '-'),
        ('Account Reconciliation', 'Non-Bundle Required', 'Monthly F', '$400.00', '-'),
        ('Bill Pay Services', 'Non-Bundle Required', 'Monthly V', '$350.00', '$0.40/payment'),
        ('Mobile Banking', 'Non-Bundle Required', 'Monthly F', '$1,350.00', '-'),
        ('eStatements', 'Non-Bundle Required', 'Monthly V', '$300.00', '$0.18/statement'),
    ]
    for i, item in enumerate(nb_required, 1):
        row = table3.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('')
    doc.add_paragraph('Non-Bundle Required Subtotal: $5,075.00/month')
    doc.add_paragraph('Note: Non-Bundle services subject to 3% annual CPI increase.')

    # Non-Bundle Optional
    doc.add_heading('Section 3: Non-Bundle Optional Services', level=1)

    table4 = doc.add_table(rows=9, cols=5)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Solution Name'
    hdr[1].text = 'Category'
    hdr[2].text = 'Fee Type'
    hdr[3].text = 'Monthly Fee'
    hdr[4].text = 'Per-Item Rate'

    nb_optional = [
        ('Business Online Banking Premium', 'Non-Bundle Optional', 'Monthly F', '$825.00', '-'),
        ('Commercial Card Services', 'Non-Bundle Optional', 'Monthly V', '$500.00', '2.75%/trans'),
        ('International Wire Module', 'Non-Bundle Optional', 'Monthly F', '$375.00', '$40.00/wire'),
        ('Lockbox Services', 'Non-Bundle Optional', 'Monthly V', '$1,650.00', '$0.50/item'),
        ('ZBA/Sweep Services', 'Non-Bundle Optional', 'Monthly F', '$225.00', '-'),
        ('Fraud Detection Premium', 'Non-Bundle Optional', 'Monthly F', '$600.00', '-'),
        ('API Access Gateway', 'Non-Bundle Optional', 'Monthly F', '$950.00', '-'),
        ('Custom Reporting Package', 'Non-Bundle Optional', 'Monthly F', '$450.00', '-'),
    ]
    for i, item in enumerate(nb_optional, 1):
        row = table4.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Third-Party Required
    doc.add_heading('Section 4: Third-Party Required Services', level=1)

    table5 = doc.add_table(rows=5, cols=5)
    table5.style = 'Table Grid'
    hdr = table5.rows[0].cells
    hdr[0].text = 'Solution Name'
    hdr[1].text = 'Category'
    hdr[2].text = 'Fee Type'
    hdr[3].text = 'Monthly Fee'
    hdr[4].text = 'Annual Fee'

    tp_required = [
        ('Fidelity Bond Coverage', 'Third-Party Required', 'Annual', '-', '$13,500.00'),
        ('Network Security Services', 'Third-Party Required', 'Monthly F', '$2,100.00', '-'),
        ('Core-to-Core Integration', 'Third-Party Required', 'Monthly F', '$1,050.00', '-'),
        ('Regulatory Compliance Suite', 'Third-Party Required', 'Monthly F', '$700.00', '-'),
    ]
    for i, item in enumerate(tp_required, 1):
        row = table5.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # One-Time Fees
    doc.add_heading('Section 5: One-Time Implementation Fees', level=1)

    table6 = doc.add_table(rows=8, cols=3)
    table6.style = 'Table Grid'
    hdr = table6.rows[0].cells
    hdr[0].text = 'Description'
    hdr[1].text = 'Fee Type'
    hdr[2].text = 'Amount'

    one_time = [
        ('Core System Implementation', 'One-Time', '$135,000.00'),
        ('Data Conversion & Migration', 'One-Time', '$48,500.00'),
        ('Training (Initial)', 'One-Time', '$20,000.00'),
        ('Project Management', 'One-Time', '$24,000.00'),
        ('Custom Configuration', 'One-Time', '$16,500.00'),
        ('Interface Development', 'One-Time', '$30,000.00'),
        ('Conversion Credit', 'One-Time', '($40,000.00)'),
    ]
    for i, item in enumerate(one_time, 1):
        row = table6.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('')
    doc.add_paragraph('One-Time Total (Net): $234,000.00')

    # Contract Terms
    doc.add_heading('Section 6: Contract Terms & Conditions', level=1)
    doc.add_paragraph('Annual Price Increases:')
    doc.add_paragraph('  - Bundle Services: 6% CPI adjustment annually, beginning Year 2')
    doc.add_paragraph('  - Non-Bundle Services: 3% CPI adjustment annually, beginning Year 2')
    doc.add_paragraph('')
    doc.add_paragraph('Volume Assumptions:')
    doc.add_paragraph('  - Average monthly DDA accounts: 18,000')
    doc.add_paragraph('  - Average monthly transactions: 525,000')
    doc.add_paragraph('  - Average ACH items: 85,000/month')
    doc.add_paragraph('  - Average wire transfers: 300/month')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'fis-like', 'Riverside_FIS_Proposal_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_fis_proposal_2():
    """Create second FIS-style proposal - Valley State Bank"""
    doc = Document()

    title = doc.add_heading('FIS Banking Services Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Prepared for: Valley State Bank')
    doc.add_paragraph('Vendor: FIS Global')
    doc.add_paragraph('Date: February 2025')
    doc.add_paragraph('')

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'FIS Global presents this proposal for comprehensive core banking services '
        'to Valley State Bank, a $450M asset community bank. The proposed solution '
        'leverages our proven HORIZON platform with tailored digital banking capabilities.'
    )

    # Bundle Options
    doc.add_heading('Bundle Pricing - 7 Year Term (Recommended)', level=1)

    table1 = doc.add_table(rows=8, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Solution'
    hdr[1].text = 'Category'
    hdr[2].text = 'Monthly Fee'
    hdr[3].text = 'Fee Type'

    items = [
        ('HORIZON Core Platform', 'Bundle', '$12,800.00', 'Monthly F'),
        ('Digital One Online Banking', 'Bundle', '$7,500.00', 'Monthly F'),
        ('Payments One', 'Bundle', '$5,200.00', 'Monthly F'),
        ('ImageCentre Document Mgmt', 'Bundle', '$3,800.00', 'Monthly F'),
        ('Analytics Dashboard', 'Bundle', '$2,100.00', 'Monthly F'),
        ('Compliance Manager', 'Bundle', '$1,400.00', 'Monthly F'),
        ('BUNDLE TOTAL', '', '$32,800.00', ''),
    ]
    for i, item in enumerate(items, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Non-Bundle Services
    doc.add_heading('Non-Bundle Required Services', level=1)

    table2 = doc.add_table(rows=7, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Base'
    hdr[2].text = 'Per-Item Fee'
    hdr[3].text = 'Fee Type'

    nb_items = [
        ('ACH Processing', '$750.00', '$0.07/item', 'Monthly V'),
        ('Wire Services', '$425.00', '$15.00/domestic', 'Monthly V'),
        ('Remote Deposit', '$575.00', '-', 'Monthly F'),
        ('Positive Pay', '$350.00', '-', 'Monthly F'),
        ('Mobile Banking App', '$1,100.00', '-', 'Monthly F'),
        ('Bill Payment', '$275.00', '$0.32/payment', 'Monthly V'),
    ]
    for i, item in enumerate(nb_items, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('Non-Bundle Required Total: $3,475.00/month')

    # Optional Services
    doc.add_heading('Optional Services', level=1)

    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Fee'
    hdr[2].text = 'Notes'

    opt_items = [
        ('Commercial Lockbox', '$1,400.00', 'Plus $0.45/item'),
        ('International Wires', '$300.00', 'Plus $35/wire'),
        ('API Gateway Access', '$800.00', 'Up to 50,000 calls'),
        ('Fraud Analytics Pro', '$525.00', 'Real-time monitoring'),
        ('Extended Support', '$375.00', '24/7 coverage'),
    ]
    for i, item in enumerate(opt_items, 1):
        row = table3.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Implementation
    doc.add_heading('Implementation Costs', level=1)

    table4 = doc.add_table(rows=6, cols=2)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Item'
    hdr[1].text = 'One-Time Fee'

    impl_items = [
        ('Core Implementation', '$98,000.00'),
        ('Data Migration', '$35,000.00'),
        ('Training Package', '$15,000.00'),
        ('Project Management', '$18,000.00'),
        ('Conversion Credit', '($25,000.00)'),
    ]
    for i, item in enumerate(impl_items, 1):
        row = table4.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('Implementation Total: $141,000.00')

    doc.add_heading('Terms', level=1)
    doc.add_paragraph('Contract Term: 7 Years')
    doc.add_paragraph('Annual CPI: 5% Bundle, 3% Non-Bundle')
    doc.add_paragraph('Early Termination: 75% of remaining contract value')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'fis-like', 'Valley_State_FIS_Proposal_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_jack_henry_proposal():
    """Create Jack Henry-style Excel proposal"""
    wb = openpyxl.Workbook()

    # Proposal Summary Sheet
    ws1 = wb.active
    ws1.title = 'Proposal Summary'

    ws1['A1'] = 'Jack Henry & Associates'
    ws1['A1'].font = Font(bold=True, size=16)
    ws1['A2'] = 'Symitar Core Banking Proposal'
    ws1['A3'] = 'Prepared for: Mountain Credit Union'
    ws1['A4'] = 'Date: January 2025'
    ws1['A6'] = 'Proposal Options:'
    ws1['A7'] = 'Proposal_1: Standard Package'
    ws1['A8'] = 'Proposal_2: Enhanced Package'
    ws1['A9'] = 'Proposal_3: Premium Package'

    # Proposal_1 Sheet
    ws2 = wb.create_sheet('Proposal_1')

    headers = ['Product Name', 'Product Family', 'Category', 'License Fee',
               'Installation', 'Monthly Fee', 'Annual Maintenance', 'Optional', 'Notes']
    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    products = [
        ('Symitar Episys Core', 'SilverLake', 'Bundle', 85000, 25000, 8500, 15000, 'N', 'Core processing platform'),
        ('PowerOn Programming', 'SilverLake', 'Bundle', 12000, 5000, 1200, 2500, 'N', 'Custom programming'),
        ('Member Connect', 'Xperience', 'Bundle', 18000, 8000, 2100, 3500, 'N', 'Member portal'),
        ('Banno Digital', 'Banno', 'Bundle', 22000, 10000, 3200, 4200, 'N', 'Digital banking'),
        ('Banno Mobile', 'Banno', 'Bundle', 15000, 6000, 2400, 2800, 'N', 'Mobile app'),
        ('ACH Services', 'Payments', 'Non-Bundle Required', 0, 2500, 650, 1200, 'N', '$0.06/item'),
        ('Wire Transfer', 'Payments', 'Non-Bundle Required', 0, 1500, 400, 800, 'N', '$12/wire'),
        ('Bill Pay', 'Payments', 'Non-Bundle Required', 0, 3000, 525, 1100, 'N', '$0.28/payment'),
        ('Remote Deposit', 'Imaging', 'Non-Bundle Required', 0, 4000, 575, 1000, 'N', 'RDC solution'),
        ('eStatements', 'Imaging', 'Non-Bundle Required', 0, 2000, 350, 700, 'N', '$0.12/statement'),
        ('Card Management', 'Cards', 'Non-Bundle Optional', 8000, 5000, 950, 1800, 'Y', 'Debit/credit cards'),
        ('Loan Origination', 'Lending', 'Non-Bundle Optional', 15000, 7500, 1400, 2800, 'Y', 'LOS module'),
        ('Collections', 'Lending', 'Non-Bundle Optional', 6000, 3000, 700, 1200, 'Y', 'Collections mgmt'),
        ('Document Management', 'Imaging', 'Non-Bundle Optional', 9000, 4500, 850, 1600, 'Y', 'Doc imaging'),
        ('Business Intelligence', 'Analytics', 'Non-Bundle Optional', 12000, 6000, 1100, 2200, 'Y', 'BI platform'),
        ('Symitar Hosted', 'Infrastructure', 'Third-Party Required', 0, 15000, 4500, 8000, 'N', 'Hosting services'),
        ('Network Security', 'Infrastructure', 'Third-Party Required', 0, 8000, 1800, 3500, 'N', 'Security suite'),
        ('DR/Business Continuity', 'Infrastructure', 'Third-Party Optional', 0, 12000, 2200, 4000, 'Y', 'DR services'),
    ]

    for row_num, product in enumerate(products, 2):
        for col_num, value in enumerate(product, 1):
            cell = ws2.cell(row=row_num, column=col_num, value=value)
            if col_num in [4, 5, 6, 7]:  # Currency columns
                cell.number_format = '$#,##0.00'

    # Add totals row
    total_row = len(products) + 3
    ws2.cell(row=total_row, column=1, value='TOTALS').font = Font(bold=True)
    ws2.cell(row=total_row, column=4, value=f'=SUM(D2:D{len(products)+1})').number_format = '$#,##0.00'
    ws2.cell(row=total_row, column=5, value=f'=SUM(E2:E{len(products)+1})').number_format = '$#,##0.00'
    ws2.cell(row=total_row, column=6, value=f'=SUM(F2:F{len(products)+1})').number_format = '$#,##0.00'
    ws2.cell(row=total_row, column=7, value=f'=SUM(G2:G{len(products)+1})').number_format = '$#,##0.00'

    # Add comments
    ws2.cell(row=2, column=6).comment = openpyxl.comments.Comment('Includes base processing fees', 'JH Sales')
    ws2.cell(row=4, column=6).comment = openpyxl.comments.Comment('Per-member pricing available', 'JH Sales')

    # Proposal_2 Sheet (Enhanced - slightly higher)
    ws3 = wb.create_sheet('Proposal_2')
    for col, header in enumerate(headers, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    for row_num, product in enumerate(products, 2):
        adjusted_product = list(product)
        # Increase prices by ~15%
        if adjusted_product[3]:
            adjusted_product[3] = int(adjusted_product[3] * 1.15)
        if adjusted_product[5]:
            adjusted_product[5] = int(adjusted_product[5] * 1.15)
        # Add more features
        if 'Optional' in adjusted_product[7]:
            adjusted_product[7] = 'N'  # Include optionals
        for col_num, value in enumerate(adjusted_product, 1):
            cell = ws3.cell(row=row_num, column=col_num, value=value)
            if col_num in [4, 5, 6, 7]:
                cell.number_format = '$#,##0.00'

    # Proposal_3 Sheet (Premium)
    ws4 = wb.create_sheet('Proposal_3')
    for col, header in enumerate(headers, 1):
        cell = ws4.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    for row_num, product in enumerate(products, 2):
        adjusted_product = list(product)
        # Premium pricing ~25% higher
        if adjusted_product[3]:
            adjusted_product[3] = int(adjusted_product[3] * 1.25)
        if adjusted_product[5]:
            adjusted_product[5] = int(adjusted_product[5] * 1.25)
        adjusted_product[7] = 'N'  # All included
        for col_num, value in enumerate(adjusted_product, 1):
            cell = ws4.cell(row=row_num, column=col_num, value=value)
            if col_num in [4, 5, 6, 7]:
                cell.number_format = '$#,##0.00'

    # Column widths
    for ws in [ws1, ws2, ws3, ws4]:
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['I'].width = 25

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'jack-henry-like', 'Mountain_CU_JackHenry_2025.xlsx')
    wb.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_csi_proposal():
    """Create CSI NuPoint-style proposal"""
    doc = Document()

    title = doc.add_heading('CSI NuPoint Core Banking Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Prepared for: Heritage National Bank')
    doc.add_paragraph('Vendor: Computer Services, Inc. (CSI)')
    doc.add_paragraph('Date: January 2025')
    doc.add_paragraph('')

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'CSI is pleased to present this proposal for NuPoint core banking services '
        'to Heritage National Bank. NuPoint provides a modern, cloud-native core '
        'banking platform with integrated digital solutions.'
    )

    # Core Services
    doc.add_heading('NuPoint Core Platform Services', level=1)

    table1 = doc.add_table(rows=8, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Category'
    hdr[2].text = 'Monthly Fee'
    hdr[3].text = 'Fee Type'

    core_items = [
        ('NuPoint Core Processing', 'Bundle', '$14,200.00', 'Monthly F'),
        ('Digital Banking Suite', 'Bundle', '$6,800.00', 'Monthly F'),
        ('NuPoint Mobile', 'Bundle', '$4,100.00', 'Monthly F'),
        ('Document Imaging', 'Bundle', '$3,200.00', 'Monthly F'),
        ('Analytics & Reporting', 'Bundle', '$2,400.00', 'Monthly F'),
        ('Security Center', 'Bundle', '$1,800.00', 'Monthly F'),
        ('CORE BUNDLE TOTAL', '', '$32,500.00', ''),
    ]
    for i, item in enumerate(core_items, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Transaction Services
    doc.add_heading('Transaction Services', level=1)

    table2 = doc.add_table(rows=7, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Base Monthly'
    hdr[2].text = 'Per Transaction'
    hdr[3].text = 'Category'

    txn_items = [
        ('ACH Origination', '$680.00', '$0.065/item', 'Non-Bundle Required'),
        ('Wire Services', '$380.00', '$14.00/domestic', 'Non-Bundle Required'),
        ('Remote Deposit Capture', '$520.00', '$0.08/deposit', 'Non-Bundle Required'),
        ('Bill Pay Platform', '$290.00', '$0.30/payment', 'Non-Bundle Required'),
        ('Positive Pay', '$340.00', '$0.05/check', 'Non-Bundle Required'),
        ('eStatements', '$245.00', '$0.10/statement', 'Non-Bundle Required'),
    ]
    for i, item in enumerate(txn_items, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('Transaction Services Total: $2,455.00/month base')

    # Optional Services
    doc.add_heading('Optional Services', level=1)

    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Fee'
    hdr[2].text = 'Description'

    opt_items = [
        ('Commercial Cash Management', '$1,250.00', 'Full TM suite'),
        ('Lockbox Processing', '$1,100.00', 'Plus $0.38/item'),
        ('International Services', '$425.00', 'Plus $32/wire'),
        ('API Developer Portal', '$725.00', 'Open banking APIs'),
        ('Advanced Fraud Detection', '$480.00', 'AI-powered monitoring'),
    ]
    for i, item in enumerate(opt_items, 1):
        row = table3.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Implementation
    doc.add_heading('Implementation Investment', level=1)

    table4 = doc.add_table(rows=7, cols=2)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'One-Time Fee'

    impl_items = [
        ('NuPoint Implementation', '$110,000.00'),
        ('Data Migration', '$38,000.00'),
        ('Training & Certification', '$16,500.00'),
        ('Project Management', '$19,500.00'),
        ('Custom Integrations', '$22,000.00'),
        ('New Customer Credit', '($30,000.00)'),
    ]
    for i, item in enumerate(impl_items, 1):
        row = table4.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    doc.add_paragraph('Net Implementation: $176,000.00')

    doc.add_heading('Contract Terms', level=1)
    doc.add_paragraph('Recommended Term: 7 Years')
    doc.add_paragraph('Annual Increase: 4% on core bundle, 3% on transaction services')
    doc.add_paragraph('Hosting: Cloud-hosted in CSIs secure data centers')
    doc.add_paragraph('SLA: 99.95% uptime guarantee')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'csi-like', 'Heritage_CSI_NuPoint_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_treasury_management_proposal():
    """Create treasury management services proposal"""
    doc = Document()

    title = doc.add_heading('Treasury Management Services Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Prepared for: Metro City School District')
    doc.add_paragraph('Submitted by: First Regional Bank')
    doc.add_paragraph('Date: January 2025')
    doc.add_paragraph('')

    doc.add_heading('Proposed Banking Services', level=1)

    # Account Services
    doc.add_heading('Account Maintenance Services', level=2)

    table1 = doc.add_table(rows=8, cols=3)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Fee'
    hdr[2].text = 'Per-Item Fee'

    acct_items = [
        ('Primary Operating Account', '$75.00', '-'),
        ('Payroll Account', '$50.00', '-'),
        ('Accounts Payable Account', '$50.00', '-'),
        ('Investment Sweep Account', '$100.00', '-'),
        ('Deposit Processing', '-', '$0.15/deposit'),
        ('Check Processing', '-', '$0.12/check'),
        ('Statement Rendering', '$15.00', '-'),
    ]
    for i, item in enumerate(acct_items, 1):
        row = table1.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Cash Management
    doc.add_heading('Cash Management Services', level=2)

    table2 = doc.add_table(rows=9, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Base'
    hdr[2].text = 'Transaction Fee'

    cm_items = [
        ('Online Banking Platform', '$125.00', '-'),
        ('ACH Origination', '$85.00', '$0.08/item'),
        ('ACH Receipt', '-', '$0.06/item'),
        ('Wire Transfer - Domestic', '$35.00', '$15.00/wire'),
        ('Wire Transfer - International', '$35.00', '$35.00/wire'),
        ('Positive Pay', '$75.00', '$0.05/item'),
        ('Remote Deposit Capture', '$95.00', '$0.10/deposit'),
        ('Zero Balance Accounting', '$50.00', '-'),
    ]
    for i, item in enumerate(cm_items, 1):
        row = table2.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Reporting
    doc.add_heading('Reporting & Analytics', level=2)

    table3 = doc.add_table(rows=5, cols=2)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Service'
    hdr[1].text = 'Monthly Fee'

    rpt_items = [
        ('Daily Balance Reporting', '$45.00'),
        ('Account Analysis Statement', '$25.00'),
        ('Custom Report Package', '$75.00'),
        ('API Data Access', '$150.00'),
    ]
    for i, item in enumerate(rpt_items, 1):
        row = table3.rows[i].cells
        for j, val in enumerate(item):
            row[j].text = val

    # Pricing Summary
    doc.add_heading('Pricing Summary', level=1)
    doc.add_paragraph('Estimated Monthly Services: $1,035.00')
    doc.add_paragraph('Estimated Transaction Costs: $450.00 (based on volume assumptions)')
    doc.add_paragraph('Total Estimated Monthly: $1,485.00')
    doc.add_paragraph('')
    doc.add_paragraph('Earnings Credit Rate: Fed Funds - 0.25%')
    doc.add_paragraph('Contract Term: 3 Years with 2 optional renewal years')

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'other-vendors', 'Metro_City_Treasury_Proposal_2025.docx')
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def create_fiserv_proposal():
    """Create Fiserv DNA-style proposal"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Fiserv DNA Pricing'

    # Header
    ws['A1'] = 'Fiserv DNA Core Banking Platform'
    ws['A1'].font = Font(bold=True, size=14)
    ws['A2'] = 'Pricing Proposal for Sunrise Community Bank'
    ws['A3'] = 'Prepared: January 2025'

    # Headers
    headers = ['Component', 'Category', 'License', 'Monthly', 'Annual Support', 'Implementation']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=header)
        cell.font = Font(bold=True)

    # DNA Products
    products = [
        ('DNA Core Platform', 'Bundle', 75000, 7500, 12000, 45000),
        ('DNA Teller', 'Bundle', 18000, 2200, 3500, 12000),
        ('DNA New Accounts', 'Bundle', 15000, 1800, 2800, 10000),
        ('DNA Lending', 'Bundle', 22000, 2800, 4200, 15000),
        ('DNA Digital Banking', 'Bundle', 28000, 3500, 5500, 18000),
        ('DNA Mobile', 'Bundle', 16000, 2400, 3800, 11000),
        ('DNA IVR', 'Non-Bundle Optional', 8000, 1200, 1800, 6000),
        ('DNA Business Intelligence', 'Non-Bundle Optional', 12000, 1600, 2500, 8000),
        ('ACH Processing', 'Non-Bundle Required', 0, 580, 0, 3500),
        ('Wire Services', 'Non-Bundle Required', 0, 320, 0, 2000),
        ('Remote Deposit', 'Non-Bundle Required', 0, 480, 0, 4000),
        ('Positive Pay', 'Non-Bundle Required', 0, 280, 0, 2500),
        ('Fiserv Hosting', 'Third-Party Required', 0, 4200, 7500, 20000),
        ('Network Security', 'Third-Party Required', 0, 1650, 3000, 8000),
    ]

    for row_num, product in enumerate(products, 6):
        for col_num, value in enumerate(product, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            if col_num >= 3:
                cell.number_format = '$#,##0.00'

    # Totals
    total_row = 6 + len(products) + 1
    ws.cell(row=total_row, column=1, value='TOTALS').font = Font(bold=True)
    for col in [3, 4, 5, 6]:
        ws.cell(row=total_row, column=col,
                value=f'=SUM({chr(64+col)}6:{chr(64+col)}{total_row-1})').number_format = '$#,##0.00'

    # Notes
    ws.cell(row=total_row+2, column=1, value='Notes:').font = Font(bold=True)
    ws.cell(row=total_row+3, column=1, value='- 7-year term recommended for best pricing')
    ws.cell(row=total_row+4, column=1, value='- Annual increase: 4% on platform, 3% on services')
    ws.cell(row=total_row+5, column=1, value='- Implementation credit of $25,000 available for Q1 signing')

    # Column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 20

    output_path = os.path.join(BASE_DIR, 'dummy-proposals', 'other-vendors', 'Sunrise_Fiserv_DNA_2025.xlsx')
    wb.save(output_path)
    print(f'Created: {output_path}')
    return output_path


def main():
    """Create all test proposals"""
    print('Creating test proposals for TCO extraction validation...')
    print('=' * 60)

    # Ensure directories exist
    for subdir in ['fis-like', 'jack-henry-like', 'csi-like', 'other-vendors']:
        os.makedirs(os.path.join(BASE_DIR, 'dummy-proposals', subdir), exist_ok=True)

    # Create proposals
    proposals = []

    print('\nCreating FIS-style proposals...')
    proposals.append(create_fis_proposal_1())
    proposals.append(create_fis_proposal_2())

    print('\nCreating Jack Henry-style proposal...')
    proposals.append(create_jack_henry_proposal())

    print('\nCreating CSI-style proposal...')
    proposals.append(create_csi_proposal())

    print('\nCreating other vendor proposals...')
    proposals.append(create_treasury_management_proposal())
    proposals.append(create_fiserv_proposal())

    print('\n' + '=' * 60)
    print(f'Successfully created {len(proposals)} test proposals:')
    for p in proposals:
        print(f'  - {os.path.basename(p)}')

    print('\nProposals are ready for testing in:')
    print(f'  {os.path.join(BASE_DIR, "dummy-proposals")}')


if __name__ == '__main__':
    main()
