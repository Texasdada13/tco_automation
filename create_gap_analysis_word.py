"""
Create professional Word document for Gap Analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading_with_color(doc, text, level, color_rgb=(54, 96, 146)):
    """Add a heading with custom color"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_colored_text(paragraph, text, color_rgb, bold=False):
    """Add colored text to a paragraph"""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(*color_rgb)
    if bold:
        run.bold = True
    return run

def create_gap_analysis_document():
    """Create comprehensive gap analysis Word document"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading('TCO Automation Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Gap Analysis: Client Requirements vs. Current Implementation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(54, 96, 146)

    doc.add_paragraph()

    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_para = doc.add_paragraph('Status: Analysis Complete')
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_para.runs[0].bold = True

    doc.add_page_break()

    # Executive Summary
    add_heading_with_color(doc, 'Executive Summary', 1)

    p = doc.add_paragraph()
    p.add_run('Overall Project Completion: ').bold = True
    add_colored_text(p, '70%', (0, 128, 0), bold=True)

    doc.add_paragraph(
        'The TCO Automation project has successfully built a robust extraction and standardization '
        'pipeline. However, critical gaps exist between the delivered solution and client requirements. '
        'The system currently creates new standardized Excel files, but the client needs integration '
        'with their existing TCO template with Excel formulas for tweakable assumptions.'
    )

    doc.add_paragraph()

    # Key Findings
    add_heading_with_color(doc, 'Key Findings', 2, (192, 0, 0))

    findings = [
        ('CRITICAL GAP', 'System creates new files instead of populating client\'s existing TCO template'),
        ('CRITICAL GAP', 'Values are hardcoded, not Excel formulas - client can\'t tweak assumptions'),
        ('CRITICAL GAP', 'No Excel macros or "click to run" automation as requested'),
        ('HIGH GAP', 'Accuracy not validated against client\'s actual manual TCOs'),
        ('MEDIUM GAP', 'CPI (Consumer Price Index) not separately implemented'),
        ('MEDIUM GAP', 'Jack Henry (JHA) vendor support not demonstrated')
    ]

    for priority, finding in findings:
        p = doc.add_paragraph(style='List Bullet')
        color = (192, 0, 0) if 'CRITICAL' in priority else (255, 128, 0) if 'HIGH' in priority else (0, 0, 255)
        add_colored_text(p, f'{priority}: ', color, bold=True)
        p.add_run(finding)

    doc.add_page_break()

    # Client Requirements
    add_heading_with_color(doc, '1. Client Requirements (From Problem Statement)', 1)

    doc.add_paragraph(
        'The client submitted two detailed problem statements outlining their needs:'
    )

    add_heading_with_color(doc, 'Initial Problem Statement', 2)

    quote1 = doc.add_paragraph(
        '"Today, our process requires manually ingesting each proposal into our 5-year TCO Excel model '
        '(Echelon_Primary TCO). This spreadsheet is what we ultimately present to the bank, and as '
        'negotiations progress, we continually refine the TCOs. Right now, this ingestion process is '
        'extremely time-consuming and prone to human error."',
        style='Quote'
    )
    quote1.runs[0].italic = True

    add_heading_with_color(doc, 'Refined Requirements', 2)

    quote2 = doc.add_paragraph(
        '"I\'d like the output structured so it can be dropped into our financial modeling with clear '
        'traceability for the client. What I\'m envisioning is an Excel or PDF we can upload in a '
        'consistent format, plus some automation (macros/formulas with a click) to capture monthly '
        'and annual totals, with growth and CPI built in (with tweakable assumptions on our end). '
        'It doesn\'t need to be 100% accurate - goal is to limit manual customization to ~10-15% '
        'and avoid repetitive copy/paste."',
        style='Quote'
    )
    quote2.runs[0].italic = True

    add_heading_with_color(doc, 'Core Requirements Summary', 2)

    requirements = [
        'Automate extraction from vendor proposals (PDF, Word, Excel) for FIS and JHA',
        'Standardize different vendor formats into "apples-to-apples" structure',
        'Populate 5-year TCO template automatically',
        'Excel/PDF output in consistent format for upload to financial modeling',
        'Macros/formulas to capture monthly and annual totals with a click',
        'Growth and CPI built in with tweakable assumptions',
        '90% accuracy - limit manual customization to 10-15%',
        'Clear traceability for the client'
    ]

    for req in requirements:
        p = doc.add_paragraph(req, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # What's Working Well
    add_heading_with_color(doc, '2. What\'s Working Well (Strengths)', 1, (0, 128, 0))

    doc.add_paragraph(
        'The project has achieved significant success in several key areas:'
    )

    strengths = [
        {
            'title': 'Extraction Pipeline',
            'rating': '95%',
            'description': 'Successfully extracts from PDF proposals using AI-powered Anthropic Claude API. Works with FIS and CSI vendors. Extracted 22-38 line items per vendor for Echelon Bank, Liberty Capital Bank, and FSB.'
        },
        {
            'title': 'Standardized Schema',
            'rating': '90%',
            'description': 'Universal 17-column format works across all vendors. Consistent fee types (Monthly F, Monthly V, Annual, One-Time) and standardized categories. Complete data dictionary for mappings.'
        },
        {
            'title': 'Credit Detection',
            'rating': '95%',
            'description': 'Automatically detects missing credits in extractions. Fixed critical $1.5M missing credit bug for Echelon Bank FIS. Properly handles negative values for implementation credits.'
        },
        {
            'title': 'Multi-Vendor Comparison',
            'rating': '85%',
            'description': 'Side-by-side vendor comparison with summary sheet and TCO calculations. Successfully tested with FIS vs CSI for Liberty Capital Bank.'
        },
        {
            'title': 'Year-by-Year Projections',
            'rating': '80%',
            'description': 'Supports projections for Years 1, 2, 3, 5, 7 (extendable to 10). Applies growth rate (20% default) with compounding calculations.'
        },
        {
            'title': 'Documentation',
            'rating': '95%',
            'description': 'Comprehensive documentation with 13 user guides, 5 reference documents, and 15 project documents. Well-organized structure with clear README files.'
        },
        {
            'title': 'Code Quality',
            'rating': '90%',
            'description': 'Clean, organized project structure with 130+ files organized into logical folders. Modular, maintainable code with version control and clean git history.'
        }
    ]

    for strength in strengths:
        add_heading_with_color(doc, f'{strength["title"]} - {strength["rating"]}', 3, (0, 128, 0))
        doc.add_paragraph(strength['description'])
        doc.add_paragraph()

    doc.add_page_break()

    # Critical Gaps
    add_heading_with_color(doc, '3. Critical Gaps (Must Fix)', 1, (192, 0, 0))

    doc.add_paragraph(
        'The following gaps represent critical misalignments between client requirements and current implementation:'
    )

    # Gap 1
    add_heading_with_color(doc, 'GAP #1: Not Populating Client\'s Actual TCO Template', 2, (192, 0, 0))

    p = doc.add_paragraph()
    add_colored_text(p, 'Priority: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'CRITICAL', (192, 0, 0), bold=True)

    p = doc.add_paragraph()
    add_colored_text(p, 'Impact: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'HIGH', (192, 0, 0), bold=True)

    add_heading_with_color(doc, 'Client Expects:', 3)
    doc.add_paragraph(
        '"Automatically populate the 5-year TCO template"',
        style='Quote'
    ).runs[0].italic = True

    add_heading_with_color(doc, 'What\'s Built:', 3)
    issues = [
        'Creates NEW standardized Excel files',
        'Does NOT populate their existing "Echelon_Primary TCO.xlsx" template',
        'Generates separate outputs instead of integrating with their model'
    ]
    for issue in issues:
        p = doc.add_paragraph(issue, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    add_heading_with_color(doc, 'Consequence:', 3)
    doc.add_paragraph(
        'Client still needs to manually copy data from our output files to their TCO template. '
        'This defeats the core purpose of automation and still requires significant manual work.'
    )

    add_heading_with_color(doc, 'What\'s Needed:', 3)
    needs = [
        'Obtain actual "Echelon_Primary TCO.xlsx" from client',
        'Build mapper to populate THAT specific template',
        'Preserve their existing formulas, formatting, and structure',
        'Only populate the data cells they need filled'
    ]
    for need in needs:
        p = doc.add_paragraph(need, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    # Gap 2
    add_heading_with_color(doc, 'GAP #2: No Excel Formulas - Values Only', 2, (192, 0, 0))

    p = doc.add_paragraph()
    add_colored_text(p, 'Priority: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'CRITICAL', (192, 0, 0), bold=True)

    p = doc.add_paragraph()
    add_colored_text(p, 'Impact: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'HIGH', (192, 0, 0), bold=True)

    add_heading_with_color(doc, 'Client Expects:', 3)
    doc.add_paragraph(
        '"macros/formulas with a click to capture monthly and annual totals, with growth and CPI '
        'built in (with tweakable assumptions on our end)"',
        style='Quote'
    ).runs[0].italic = True

    add_heading_with_color(doc, 'What\'s Built:', 3)
    doc.add_paragraph(
        'Python calculates all values and writes hardcoded numbers to Excel cells. There are no '
        'Excel formulas for recalculation. Client cannot tweak assumptions without re-running the '
        'Python script.'
    )

    add_heading_with_color(doc, 'Code Example (Current):', 3)
    code = doc.add_paragraph(
        'year_costs[1] = 180000  # Hardcoded value\n'
        'ws.cell(row, col).value = year_costs[1]  # Written as static number'
    )
    code.runs[0].font.name = 'Consolas'
    code.runs[0].font.size = Pt(9)

    add_heading_with_color(doc, 'What Should Be (Excel Formulas):', 3)
    code = doc.add_paragraph(
        '# Parameter cells:\n'
        'B1: Growth Rate (20%)\n'
        'B2: CPI Rate (2.5%)\n\n'
        '# Formula-based calculations:\n'
        'Year 1: =E7*12\n'
        'Year 2: =E7*12*(1+$B$1)\n'
        'Year 6: =E7*12*(1+$B$1)^5*(1+$B$2)'
    )
    code.runs[0].font.name = 'Consolas'
    code.runs[0].font.size = Pt(9)

    add_heading_with_color(doc, 'Consequence:', 3)
    consequences = [
        'Client cannot adjust growth rates in Excel',
        'Cannot tweak CPI assumptions',
        'No "what-if" analysis capability',
        'Must re-run Python script for any parameter changes'
    ]
    for consequence in consequences:
        p = doc.add_paragraph(consequence, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    add_heading_with_color(doc, 'What\'s Needed:', 3)
    needs = [
        'Replace hardcoded values with Excel formulas',
        'Add parameter cells for growth rate, CPI, and other assumptions',
        'Make all year projections formula-based',
        'Enable client to tweak assumptions directly in Excel'
    ]
    for need in needs:
        p = doc.add_paragraph(need, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # Gap 3
    add_heading_with_color(doc, 'GAP #3: No Macros or "Click" Automation', 2, (255, 128, 0))

    p = doc.add_paragraph()
    add_colored_text(p, 'Priority: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'HIGH', (255, 128, 0), bold=True)

    p = doc.add_paragraph()
    add_colored_text(p, 'Impact: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'HIGH', (255, 128, 0), bold=True)

    add_heading_with_color(doc, 'Client Expects:', 3)
    doc.add_paragraph(
        '"macros/formulas with a click"',
        style='Quote'
    ).runs[0].italic = True

    add_heading_with_color(doc, 'What\'s Built:', 3)
    doc.add_paragraph(
        'Command-line Python scripts only. No Excel macros. No "click to run" functionality. '
        'Requires technical knowledge to operate.'
    )

    add_heading_with_color(doc, 'Consequence:', 3)
    doc.add_paragraph(
        'Not user-friendly for client. Cannot be used by non-technical staff. Requires Python '
        'environment setup.'
    )

    add_heading_with_color(doc, 'What\'s Needed:', 3)
    needs = [
        'Excel VBA macros for automation',
        'Buttons in Excel to trigger actions (e.g., "Load Proposal", "Calculate TCO")',
        'Alternative: Simple web interface for uploading proposals',
        'Alternative: Excel add-in for integration'
    ]
    for need in needs:
        p = doc.add_paragraph(need, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    # Gap 4
    add_heading_with_color(doc, 'GAP #4: Accuracy Not Validated Against Client TCOs', 2, (255, 128, 0))

    p = doc.add_paragraph()
    add_colored_text(p, 'Priority: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'HIGH', (255, 128, 0), bold=True)

    p = doc.add_paragraph()
    add_colored_text(p, 'Impact: ', (0, 0, 0), bold=True)
    add_colored_text(p, 'HIGH', (255, 128, 0), bold=True)

    add_heading_with_color(doc, 'Client Expects:', 3)
    doc.add_paragraph(
        '"goal is to limit manual customization to ~10-15%"',
        style='Quote'
    ).runs[0].italic = True

    add_heading_with_color(doc, 'What\'s Built:', 3)
    doc.add_paragraph(
        'Internal gap analysis shows 31% → 73% coverage improvement for Echelon FIS. However, '
        'this has NOT been validated against the client\'s actual manually-created TCO. No '
        'measurement of the 10-15% manual customization target.'
    )

    add_heading_with_color(doc, 'Consequence:', 3)
    doc.add_paragraph(
        'Unknown if we\'re meeting the 90% accuracy goal. Cannot measure success. Don\'t know '
        'which specific items are systematically missing or incorrect.'
    )

    add_heading_with_color(doc, 'What\'s Needed:', 3)
    needs = [
        'Obtain client\'s manually-created TCO for Echelon Bank and Liberty Capital Bank',
        'Compare our output line-by-line against their manual TCO',
        'Measure accuracy percentage',
        'Identify systematic gaps and patterns',
        'Create QA workflow for future validations'
    ]
    for need in needs:
        p = doc.add_paragraph(need, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # Medium Priority Gaps
    add_heading_with_color(doc, '4. Medium Priority Gaps', 1, (255, 128, 0))

    # Gap 5
    add_heading_with_color(doc, 'GAP #5: CPI Not Explicitly Implemented', 2, (255, 128, 0))

    doc.add_paragraph(
        'Client expects both growth and CPI. Current implementation only applies a generic '
        'growth rate (20% default). CPI (Consumer Price Index) inflation should be applied '
        'separately, typically starting in Year 6-7.'
    )

    add_heading_with_color(doc, 'What\'s Needed:', 3)
    doc.add_paragraph('Add separate CPI parameter (e.g., 2-3%)')
    doc.add_paragraph('Apply CPI starting Year 6 (after initial growth period)')
    doc.add_paragraph('Make both growth rate and CPI tweakable in Excel')

    doc.add_paragraph()

    # Gap 6
    add_heading_with_color(doc, 'GAP #6: Jack Henry (JHA) Support Not Demonstrated', 2, (255, 128, 0))

    doc.add_paragraph(
        'Client specifically requested FIS and JHA vendor support. Current implementation has '
        'FIS and CSI working, but no Jack Henry (JHA) proposal samples have been tested.'
    )

    add_heading_with_color(doc, 'What\'s Needed:', 3)
    doc.add_paragraph('Obtain JHA proposal samples from client')
    doc.add_paragraph('Test JHA extraction with existing pipeline')
    doc.add_paragraph('Validate JHA standardized output format')
    doc.add_paragraph('Create JHA vs FIS comparison')

    doc.add_page_break()

    # Gap Summary Table
    add_heading_with_color(doc, '5. Gap Summary Matrix', 1)

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Priority'
    header_cells[1].text = 'Gap'
    header_cells[2].text = 'Status'
    header_cells[3].text = 'Impact'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    gaps_data = [
        ('🔴 CRITICAL', 'Not populating client\'s TCO template', '❌ Missing', 'HIGH'),
        ('🔴 CRITICAL', 'No Excel formulas (values only)', '❌ Missing', 'HIGH'),
        ('🔴 CRITICAL', 'No financial modeling integration', '❌ Unknown', 'HIGH'),
        ('🔴 HIGH', 'No macros/"click" automation', '❌ Missing', 'HIGH'),
        ('🔴 HIGH', 'Accuracy not validated', '❌ Missing', 'HIGH'),
        ('🟡 MEDIUM', 'CPI not implemented', '⚠️ Partial', 'MEDIUM'),
        ('🟡 MEDIUM', 'No JHA support demonstrated', '⚠️ Partial', 'MEDIUM')
    ]

    for i, (priority, gap, status, impact) in enumerate(gaps_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = priority
        row_cells[1].text = gap
        row_cells[2].text = status
        row_cells[3].text = impact

    doc.add_page_break()

    # Core Issue
    add_heading_with_color(doc, '6. Core Issue: System vs. Client Expectations', 1)

    add_heading_with_color(doc, 'What We Built:', 2)
    p = doc.add_paragraph()
    p.add_run('Extract Proposal ').font.name = 'Courier New'
    p.add_run('→ ').bold = True
    p.add_run('Create NEW Standardized Excel File').font.name = 'Courier New'

    doc.add_paragraph()

    add_heading_with_color(doc, 'What Client Wants:', 2)
    p = doc.add_paragraph()
    p.add_run('Extract Proposal ').font.name = 'Courier New'
    p.add_run('→ ').bold = True
    p.add_run('Populate THEIR TCO Template with Formulas ').font.name = 'Courier New'
    p.add_run('→ ').bold = True
    p.add_run('They Tweak & Use').font.name = 'Courier New'

    doc.add_paragraph()

    add_heading_with_color(doc, 'Analogy:', 2)
    doc.add_paragraph(
        'We built a great data extraction pipeline but didn\'t connect it to their final destination. '
        'It\'s like building a state-of-the-art water filtration system but not connecting it to '
        'the house\'s existing plumbing - the water is clean, but it doesn\'t flow where it needs to go.'
    ).runs[0].italic = True

    doc.add_page_break()

    # Recommendations
    add_heading_with_color(doc, '7. Recommended Action Plan', 1)

    add_heading_with_color(doc, 'Phase 1: Critical Fixes (Must Do)', 2, (192, 0, 0))

    actions1 = [
        ('Get Client\'s Actual TCO Template',
         'Request "Echelon_Primary TCO.xlsx" and understand their exact structure, formulas, and workflow'),

        ('Replace Values with Excel Formulas',
         'Rewrite output generation to create formulas instead of values. Add parameter cells for growth rate, CPI, and other assumptions. Enable "what-if" analysis.'),

        ('Validate Against Client\'s Manual TCO',
         'Obtain their manually-created TCOs, compare line-by-line, measure accuracy percentage, identify systematic gaps'),

        ('Build Template Population Script',
         'Create script to populate THEIR template (not create new files). Preserve their formulas and formatting.')
    ]

    for i, (action, description) in enumerate(actions1, 1):
        add_heading_with_color(doc, f'{i}. {action}', 3)
        doc.add_paragraph(description)
        doc.add_paragraph()

    add_heading_with_color(doc, 'Phase 2: High Priority', 2, (255, 128, 0))

    actions2 = [
        ('Add Excel Macros',
         'Create VBA macros with buttons for "Load Proposal", "Calculate TCO", "Compare Vendors"'),

        ('Implement CPI',
         'Add explicit CPI handling separate from growth rate. Apply starting Year 6-7 (typical industry practice). Make tweakable.'),

        ('Test JHA Support',
         'Get JHA proposal samples, test extraction, validate output, create comparison with FIS')
    ]

    for i, (action, description) in enumerate(actions2, 5):
        add_heading_with_color(doc, f'{i}. {action}', 3)
        doc.add_paragraph(description)
        doc.add_paragraph()

    doc.add_page_break()

    # Questions for Client
    add_heading_with_color(doc, '8. Critical Questions for Client', 1)

    doc.add_paragraph(
        'To properly address the gaps, we need clarification on the following:'
    )

    questions = [
        ('Can you provide your actual TCO template?',
         'We need "Echelon_Primary TCO.xlsx" or similar to understand your exact structure and integrate properly.'),

        ('What is your "financial modeling" system?',
         'Is it Excel-based? Do you use specific software? What format does it need?'),

        ('Do you have a manually-created TCO we can validate against?',
         'We need to measure our accuracy against your manual work to verify the 90% accuracy goal.'),

        ('What assumptions do you typically tweak?',
         'Growth rate? CPI? Volume/quantities? Pricing tiers? This helps us make the right things formula-based.'),

        ('Do you have Jack Henry (JHA) proposal samples?',
         'We need actual JHA proposals to test and validate the extraction.'),

        ('How do you envision "clicking" to run automation?',
         'Excel macro button? Web interface? Desktop application? What would be most useful for your workflow?')
    ]

    for i, (question, context) in enumerate(questions, 1):
        add_heading_with_color(doc, f'Question {i}: {question}', 3, (0, 0, 255))
        doc.add_paragraph(context)
        doc.add_paragraph()

    doc.add_page_break()

    # Progress Summary
    add_heading_with_color(doc, '9. Current Progress Summary', 1)

    progress_items = [
        ('Extraction Pipeline', '95%', 'AI-powered extraction working excellently'),
        ('Standardization', '90%', 'Universal schema working across vendors'),
        ('Documentation', '95%', 'Comprehensive guides and references'),
        ('Code Quality', '90%', 'Clean, organized, maintainable'),
        ('Integration', '20%', 'Creates new files, not integrated with client workflow'),
        ('Interactivity', '10%', 'Command-line only, no macros or UI'),
        ('Validation', '0%', 'Not validated against client\'s actual TCOs')
    ]

    table = doc.add_table(rows=len(progress_items)+1, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Completion'
    header_cells[2].text = 'Status'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Data
    for i, (component, completion, status) in enumerate(progress_items, 1):
        cells = table.rows[i].cells
        cells[0].text = component
        cells[1].text = completion
        cells[2].text = status

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Overall Project Completion: ').bold = True
    add_colored_text(p, '70%', (0, 128, 0), bold=True)

    doc.add_paragraph(
        'Strong foundation built, but critical integration and interactivity gaps must be addressed '
        'to meet client expectations.'
    )

    doc.add_page_break()

    # Conclusion
    add_heading_with_color(doc, '10. Conclusion & Next Steps', 1)

    doc.add_paragraph(
        'The TCO Automation project has successfully built robust extraction, standardization, and '
        'analysis capabilities. The technical foundation is solid with excellent code quality and '
        'documentation.'
    )

    doc.add_paragraph(
        'However, critical gaps exist in the integration with client workflows. The system currently '
        'creates standalone output files, but the client needs:'
    )

    needs = [
        'Integration with their existing TCO template',
        'Excel formulas (not hardcoded values) for tweakable assumptions',
        'User-friendly "click to run" automation via Excel macros',
        'Validation against their actual manual TCOs'
    ]

    for need in needs:
        doc.add_paragraph(need, style='List Number')

    doc.add_paragraph()

    add_heading_with_color(doc, 'Immediate Next Steps:', 2)

    next_steps = [
        'Schedule meeting with client to discuss gaps and get their TCO template',
        'Obtain "Echelon_Primary TCO.xlsx" and client\'s manually-created TCOs',
        'Prioritize Phase 1 critical fixes (template integration, Excel formulas)',
        'Get JHA proposal samples for testing',
        'Define success criteria and validation process'
    ]

    for i, step in enumerate(next_steps, 1):
        p = doc.add_paragraph(f'{i}. ', style='List Number')
        p.add_run(step)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    add_colored_text(p, 'Ready for client discussion and Phase 1 implementation', (0, 128, 0))

    # Save document
    output_file = 'TCO_Automation_Gap_Analysis.docx'
    doc.save(output_file)
    print(f'\n[OK] Created: {output_file}')
    return output_file

if __name__ == '__main__':
    create_gap_analysis_document()
