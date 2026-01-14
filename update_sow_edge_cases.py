"""
Update SOW Section 2.1.1 with Comprehensive Edge Cases
Adds detailed scenarios for what the tool can and cannot handle
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_edge_cases_to_sow():
    """Add comprehensive edge cases to section 2.1.1"""

    # Load SOW document
    doc = Document('Karishma Proposal Documents/SOW_Phase2_TCO_Automation sm updates 1_9.docx')

    print('[INFO] Adding edge cases to Section 2.1.1...')

    # Find section 2.1.1
    section_index = None
    for i, para in enumerate(doc.paragraphs):
        if '2.1.1' in para.text and 'Volume Processing Infrastructure' in para.text:
            section_index = i
            print(f'Found section 2.1.1 at paragraph {i}')
            break

    if section_index is None:
        print('[ERROR] Section 2.1.1 not found')
        return

    # Insert content after the heading
    insert_pos = section_index + 1

    # Edge cases content
    edge_cases_content = [
        {
            'heading': 'Supported Document Types & Formats',
            'content': [
                ('Word Documents (.docx, .doc)', [
                    'Native Microsoft Word format proposals (FIS, Jack Henry, etc.)',
                    'Structured pricing tables with clear column headers',
                    'Standard formatting (normal fonts, no extreme styling)',
                    'File size: Up to 50MB'
                ]),
                ('Excel Spreadsheets (.xlsx, .xls)', [
                    'Native Excel deal sheets and pricing models',
                    'Multiple sheets supported (scenarios, bundles, modules)',
                    'Formula preservation and extraction',
                    'Hidden rows/columns detection',
                    'Comments and notes extraction',
                    'File size: Up to 100MB'
                ]),
                ('PDF Documents (.pdf)', [
                    'Text-based PDFs (created digitally, not scanned)',
                    'Scanned PDFs with clear text (300+ DPI recommended)',
                    'Complex layouts with side-by-side tables',
                    'Merged cells and nested category structures',
                    'Multi-page pricing sections',
                    'File size: Up to 200MB',
                    'Note: Direct PDF extraction via Claude AI (multimodal) achieves 95-99% accuracy on complex layouts'
                ])
            ]
        },
        {
            'heading': 'Known Limitations & Edge Cases',
            'content': [
                ('Document Format Edge Cases', [
                    'Password-protected or encrypted documents (NOT SUPPORTED - require unencrypted files)',
                    'Heavily corrupted files with missing data or unreadable sections',
                    'Documents with extreme formatting (rainbow colors, artistic layouts, non-standard fonts)',
                    'Image-only PDFs without OCR layer (requires OCR preprocessing)',
                    'Handwritten proposals or hand-annotated documents (OCR accuracy <60%)',
                    'Non-English proposals or multi-language documents (English-only supported in Phase 2)',
                    'Proprietary vendor-specific formats (.proprietary, .custom) not convertible to standard formats'
                ]),
                ('Table Structure Complexities', [
                    'Highly complex nested tables (3+ levels deep) may require manual validation',
                    'Tables spanning multiple pages with broken continuity',
                    'Rotated or sideways tables (landscape orientation within portrait pages)',
                    'Tables with irregular cell merging (random merged cells without pattern)',
                    'Pricing tables embedded in narrative text (non-tabular proposals)',
                    'Side-by-side comparison tables with unclear column separation',
                    'Tables using images or graphics instead of text (charts, diagrams)',
                    'Dynamically generated tables from database dumps (inconsistent structure)'
                ]),
                ('Data Quality & Content Issues', [
                    'Missing critical pricing data (blank cells, "TBD", "Contact for pricing")',
                    'Inconsistent units (mix of monthly/annual, per-user/per-account without clear labels)',
                    'Ambiguous vendor terminology (non-standard product names, internal codes)',
                    'Contradictory pricing (table vs narrative text showing different amounts)',
                    'Incomplete proposals (missing implementation fees, no contract terms)',
                    'Proposals with conditional pricing ("if X, then Y; otherwise Z") requiring business logic',
                    'Volume-based graduated pricing with 10+ tiers (may truncate to top 5 tiers)',
                    'Proposals referencing external appendices or separate pricing documents not provided'
                ]),
                ('Vendor-Specific Anomalies', [
                    'New vendors not previously encountered (require 2-3 sample proposals for pattern learning)',
                    'Vendor format changes mid-contract (old vs new template structure)',
                    'Custom proposals with non-standard layouts (RFP responses, custom quotes)',
                    'Proposals mixing multiple vendor products (bundled deals from 2+ vendors)',
                    'White-labeled vendor proposals (rebadged products with altered terminology)',
                    'Proposals with proprietary pricing models (non-standard fee structures)',
                    'Redacted proposals with pricing information obscured or removed'
                ]),
                ('TCO Template Matching Edge Cases', [
                    'Client-specific custom TCO templates not in template registry (requires manual template upload)',
                    'TCO templates with 50+ custom fields (may require mapping configuration)',
                    'Templates requiring complex calculations not supported by standard formulas',
                    'Templates with conditional formatting based on external data (regulatory caps, market rates)',
                    'Multi-currency TCO templates (USD-only supported in Phase 2)',
                    'Templates requiring manual approval workflows not yet configured',
                    'Historical templates from 5+ years ago (obsolete field structures)'
                ]),
                ('Volume & Performance Edge Cases', [
                    'Batch uploads exceeding 50 proposals simultaneously (may trigger rate limiting)',
                    'Individual proposal files exceeding 200MB (require manual splitting)',
                    'Proposals with 500+ line items (processing time may exceed 5 minutes)',
                    'Concurrent processing requests exceeding worker pool capacity (queue delays)',
                    'API rate limits during peak usage (Claude API throttling)',
                    'Disk storage approaching capacity (proposals >1GB total size)',
                    'Network interruptions during long-running extractions (require resume capability)'
                ])
            ]
        },
        {
            'heading': 'Mitigation Strategies & Fallback Procedures',
            'content': [
                ('When Edge Cases Occur', [
                    'Confidence Scoring: All extractions include per-field confidence scores (0-100%)',
                    'Two-Bucket Routing: Low-confidence items (<90%) automatically flagged for manual review',
                    'Manual Review Queue: Web UI allows reviewers to validate, correct, and approve flagged items',
                    'Feedback Loop: Manual corrections captured and used to improve future extractions',
                    'Partial Success Handling: System extracts what it can, flags incomplete sections',
                    'Expert Escalation: Complex edge cases routed to designated subject matter experts'
                ]),
                ('Preprocessing Requirements', [
                    'Password-protected files: Client must provide unencrypted versions',
                    'Scanned PDFs: Client provides OCR-processed files or originals for OCR preprocessing',
                    'Non-English proposals: Client provides English translations or waits for multi-language support',
                    'Missing appendices: Client uploads all referenced documents for complete extraction',
                    'Corrupted files: Client re-exports from source system or provides alternative format',
                    'Proprietary formats: Client converts to Word, Excel, or PDF before upload'
                ]),
                ('Discovery Phase Validation', [
                    'Historical inventory identifies edge case frequency (% of proposals affected)',
                    'Sample testing on 10-20 historical files validates accuracy across formats',
                    'Template analysis documents custom fields requiring special handling',
                    'Vendor pattern analysis identifies vendors requiring custom extraction rules',
                    'Volume testing confirms system capacity for client-specific proposal sizes',
                    'Edge case playbook developed documenting known issues and resolutions'
                ])
            ]
        },
        {
            'heading': 'Out-of-Scope Scenarios (Require Future Enhancement)',
            'content': [
                ('Phase 2 Exclusions', [
                    'Real-time proposal analysis (proposals analyzed on-demand, not during vendor meetings)',
                    'Proposals in languages other than English (multi-language support planned for Phase 3)',
                    'Integration with vendor portals for automatic proposal download (manual upload required)',
                    'Automated contract term extraction (focus on pricing only; terms require manual entry)',
                    'Predictive pricing modeling (benchmarking available; forecasting requires ML training)',
                    'Multi-currency support and FX conversion (USD-only in Phase 2)',
                    'Custom business rule engines for client-specific pricing logic',
                    'Proposals requiring external data enrichment (market rates, competitor pricing)'
                ])
            ]
        },
        {
            'heading': 'Success Criteria & Quality Gates',
            'content': [
                ('Acceptance Thresholds', [
                    'Extraction Accuracy: ≥95% for well-structured proposals (clean Word/Excel/PDF)',
                    'Extraction Accuracy: ≥85% for complex/scanned proposals (with OCR)',
                    'Auto-Accept Rate: ≥70% of proposals processed without manual intervention',
                    'Processing Time: <5 minutes per proposal (end-to-end)',
                    'Template Matching: ≥95% correct template selection',
                    'Data Completeness: ≥90% of required fields populated',
                    'False Positive Rate: <5% (incorrect data accepted as confident)'
                ]),
                ('Quality Validation Process', [
                    'Random sampling: 10% of auto-accepted proposals manually audited weekly',
                    'Confidence calibration: Periodic validation that 90% confidence = 90% accuracy',
                    'Edge case tracking: All flagged items logged with resolution outcomes',
                    'Accuracy trending: Monthly reports on extraction quality by vendor and format',
                    'Client feedback loops: Quarterly reviews of accuracy and edge case handling',
                    'Continuous improvement: Bi-weekly model retraining with corrected data'
                ])
            ]
        }
    ]

    # Insert content
    for section_data in edge_cases_content:
        # Add heading
        p = doc.paragraphs[insert_pos].insert_paragraph_before(section_data['heading'])
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        insert_pos += 1

        # Add blank line
        doc.paragraphs[insert_pos].insert_paragraph_before('')
        insert_pos += 1

        # Add content
        for sub_heading, items in section_data['content']:
            # Sub-heading
            p = doc.paragraphs[insert_pos].insert_paragraph_before(sub_heading)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.left_indent = Pt(18)
            insert_pos += 1

            # Bullet points
            for item in items:
                p = doc.paragraphs[insert_pos].insert_paragraph_before(f'• {item}')
                p.runs[0].font.size = Pt(10)
                p.paragraph_format.left_indent = Pt(36)
                insert_pos += 1

            # Blank line after each sub-section
            doc.paragraphs[insert_pos].insert_paragraph_before('')
            insert_pos += 1

    # Save document
    doc.save('Karishma Proposal Documents/SOW_Phase2_TCO_Automation sm updates 1_9.docx')

    print('\n[SUCCESS] Edge cases added to Section 2.1.1!')
    print(f'\nContent added:')
    print(f'  - Supported Document Types (Word, Excel, PDF with specifications)')
    print(f'  - 6 Categories of Edge Cases:')
    print(f'    1. Document Format Edge Cases (7 scenarios)')
    print(f'    2. Table Structure Complexities (8 scenarios)')
    print(f'    3. Data Quality & Content Issues (8 scenarios)')
    print(f'    4. Vendor-Specific Anomalies (7 scenarios)')
    print(f'    5. TCO Template Matching Edge Cases (7 scenarios)')
    print(f'    6. Volume & Performance Edge Cases (7 scenarios)')
    print(f'  - Mitigation Strategies (3 categories)')
    print(f'  - Out-of-Scope Scenarios (8 items)')
    print(f'  - Success Criteria & Quality Gates')
    print(f'\nTotal Edge Cases Documented: 44 specific scenarios')

if __name__ == '__main__':
    add_edge_cases_to_sow()
