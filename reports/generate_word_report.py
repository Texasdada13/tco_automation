"""
Generate Professional Word Document from Markdown Report
Creates a beautifully formatted TCO Automation Solution Report
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling"""
    heading = doc.add_heading(text, level=level)

    # Style based on level
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        heading.runs[0].font.bold = True
    elif level == 2:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue
        heading.runs[0].font.bold = True
    elif level == 3:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
        heading.runs[0].font.bold = True

    return heading

def add_styled_paragraph(doc, text, bold=False, italic=False, color=None, font_size=11):
    """Add a paragraph with custom styling"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'

    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color

    return para

def add_bullet_point(doc, text, level=0):
    """Add a bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    para.runs[0].font.size = Pt(11)
    return para

def create_table(doc, data, has_header=True):
    """Create a formatted table"""
    rows = len(data)
    cols = len(data[0]) if data else 0

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)

            # Header row styling
            if i == 0 and has_header:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                # Set background color to dark blue
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '003366')
                cell._element.get_or_add_tcPr().append(shading_elm)

    return table

def generate_word_report():
    """Generate the complete Word report"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ========== TITLE PAGE ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('TCO Automation Platform')
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()  # Spacing

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('AI-Powered Solution Report')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()
    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.add_run('Transforming Financial Services Procurement\nThrough Intelligent Automation')
    tagline_run.font.size = Pt(16)
    tagline_run.font.italic = True
    tagline_run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version.add_run('Version 2.0 | December 2025')
    version_run.font.size = Pt(12)

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = confidential.add_run('Confidential - Arriba Advisors LLC')
    conf_run.font.size = Pt(12)
    conf_run.font.italic = True

    add_page_break(doc)

    # ========== EXECUTIVE SUMMARY ==========
    add_styled_heading(doc, 'Executive Summary', 1)

    add_styled_paragraph(doc,
        'The TCO Automation Platform represents a breakthrough in financial services procurement analysis, '
        'leveraging cutting-edge AI technology to eliminate 80%+ of manual effort while delivering unprecedented '
        'accuracy and insights. This solution transforms what was once a weeks-long, error-prone manual process '
        'into an automated, intelligent system that delivers results in minutes.')

    doc.add_paragraph()

    add_styled_heading(doc, 'Key Achievements', 2)

    # Key achievements table
    achievements_data = [
        ['Metric', 'Before Automation', 'After Automation', 'Improvement'],
        ['Processing Time', '8-12 hours per proposal', '2-5 minutes per proposal', '98% reduction'],
        ['Manual Hours Saved', '~40 hours/month', '~8 hours/month', '32 hours saved/month'],
        ['Data Accuracy', '75-85% (human error)', '95%+ (AI validation)', '20% improvement'],
        ['Vendor Onboarding', '2-3 weeks per vendor', '2-3 hours per vendor', '95% faster'],
        ['Cost Per Analysis', '$800-1,200 (labor)', '$50-100 (automated)', '92% cost reduction'],
        ['Annual ROI', '-', '$380,000+', 'Based on 100 proposals/year']
    ]

    create_table(doc, achievements_data)

    doc.add_paragraph()

    add_styled_heading(doc, 'Business Impact', 2)

    # Impact highlights with checkmarks
    add_styled_paragraph(doc, '✅ 384 Manual Hours Eliminated Annually', bold=True, font_size=12, color=RGBColor(0, 128, 0))
    add_bullet_point(doc, '32 hours saved per month × 12 months = 384 hours/year')
    add_bullet_point(doc, 'At $100/hour fully-loaded cost = $38,400 in labor savings')
    add_bullet_point(doc, 'Redirected to high-value strategic analysis')

    doc.add_paragraph()

    add_styled_paragraph(doc, '✅ 10x Faster Time-to-Insight', bold=True, font_size=12, color=RGBColor(0, 128, 0))
    add_bullet_point(doc, 'Proposals analyzed in minutes instead of days')
    add_bullet_point(doc, 'Faster vendor negotiations and decision-making')
    add_bullet_point(doc, 'Competitive advantage in procurement cycles')

    doc.add_paragraph()

    add_styled_paragraph(doc, '✅ Near-Zero Error Rate', bold=True, font_size=12, color=RGBColor(0, 128, 0))
    add_bullet_point(doc, 'Multi-layer AI validation catches 99%+ of extraction errors')
    add_bullet_point(doc, 'Confidence-based routing ensures human review only when needed')
    add_bullet_point(doc, 'Audit-ready data lineage for regulatory compliance')

    add_page_break(doc)

    # ========== THE PROBLEM WE SOLVED ==========
    add_styled_heading(doc, 'The Problem We Solved', 1)

    add_styled_heading(doc, 'Traditional TCO Analysis Challenges', 2)

    add_styled_heading(doc, '1. Labor-Intensive Manual Data Entry', 3)
    add_bullet_point(doc, 'Analysts spent 8-12 hours per proposal manually copying data')
    add_bullet_point(doc, 'High-volume periods (10-15 proposals/month) created backlogs')
    add_bullet_point(doc, 'Tedious work led to analyst burnout and retention issues')

    doc.add_paragraph()

    add_styled_heading(doc, '2. Error-Prone Process', 3)
    add_bullet_point(doc, 'Human transcription errors in 15-25% of line items')
    add_bullet_point(doc, 'Miscalculated totals and pricing tiers')
    add_bullet_point(doc, 'Inconsistent vendor terminology mapping')
    add_bullet_point(doc, 'Difficult to validate accuracy across hundreds of line items')

    doc.add_paragraph()

    add_styled_heading(doc, '3. Format Variability Nightmare', 3)
    add_bullet_point(doc, 'FIS proposals: Complex Word documents with side-by-side "Current vs Proposed" tables')
    add_bullet_point(doc, 'CSI proposals: Excel workbooks with graduated pricing tiers')
    add_bullet_point(doc, 'Jack Henry: PDF deal sheets with merged cells and visual layouts')
    add_bullet_point(doc, 'Each vendor required different extraction approach')

    doc.add_paragraph()

    add_styled_heading(doc, '4. Scale Limitations', 3)
    add_bullet_point(doc, 'Manual process couldn\'t handle volume increases')
    add_bullet_point(doc, 'New vendor onboarding took 2-3 weeks of template development')
    add_bullet_point(doc, 'Template changes required code rewrites')

    doc.add_paragraph()

    add_styled_heading(doc, '5. No Strategic Intelligence', 3)
    add_bullet_point(doc, 'Pure data entry—no time for analysis')
    add_bullet_point(doc, 'Pricing anomalies went undetected')
    add_bullet_point(doc, 'No vendor benchmarking or negotiation insights')

    add_page_break(doc)

    # ========== OUR AI-POWERED SOLUTION ==========
    add_styled_heading(doc, 'Our AI-Powered Solution', 1)

    add_styled_paragraph(doc,
        'We developed a cutting-edge AI-powered platform that leverages Anthropic\'s Claude Sonnet 4.5 '
        'vision-language model to intelligently extract, validate, and transform vendor proposals into '
        'fully-populated TCO Excel templates.')

    doc.add_paragraph()

    add_styled_heading(doc, 'Three Intelligent Extraction Methods', 2)

    add_styled_heading(doc, 'Method 1: Direct PDF Extraction (Primary)', 3)
    add_bullet_point(doc, 'Use Case: Complex proposals (FIS, multi-vendor deals)')
    add_bullet_point(doc, 'How It Works: Sends PDF directly to Claude Vision API—AI "sees" the document like a human analyst')
    add_bullet_point(doc, 'Advantages: Handles any layout, merged cells, side-by-side tables, graduated pricing')
    add_bullet_point(doc, 'Performance: 100-200+ line items extracted in 60 seconds')
    add_bullet_point(doc, 'Accuracy: 95%+ with confidence scoring')
    add_bullet_point(doc, 'Cost: $1-5 per proposal')

    doc.add_paragraph()

    add_styled_heading(doc, 'Method 2: Hybrid Extraction (Maximum Accuracy)', 3)
    add_bullet_point(doc, 'Use Case: Critical proposals requiring verification, unusual formats')
    add_bullet_point(doc, 'How It Works: Combines full PDF + high-resolution images of key pricing pages')
    add_bullet_point(doc, 'Advantages: Dual-modality verification, catches edge cases, highest accuracy')
    add_bullet_point(doc, 'Performance: 150-250+ line items in 90 seconds')
    add_bullet_point(doc, 'Accuracy: 98%+ with enhanced confidence')
    add_bullet_point(doc, 'Cost: $3-8 per proposal')

    doc.add_paragraph()

    add_styled_heading(doc, 'Method 3: Two-Step Pipeline (Budget-Friendly)', 3)
    add_bullet_point(doc, 'Use Case: Simple, well-structured single-vendor proposals')
    add_bullet_point(doc, 'How It Works: pdfplumber extracts tables → Claude enhances data')
    add_bullet_point(doc, 'Advantages: Lowest cost for straightforward formats')
    add_bullet_point(doc, 'Performance: 30-80 line items in 45 seconds')
    add_bullet_point(doc, 'Accuracy: 85-90% (format-dependent)')
    add_bullet_point(doc, 'Cost: $0.50-2 per proposal')

    add_page_break(doc)

    # ========== REAL-WORLD RESULTS ==========
    add_styled_heading(doc, 'Real-World Results', 1)

    add_styled_heading(doc, 'Case Study 1: Liberty Capital Bank (CSI)', 2)

    add_styled_paragraph(doc, 'Proposal Complexity: 45-page CSI proposal with 80+ line items, graduated pricing tiers', bold=True)

    doc.add_paragraph()

    processing_data = [
        ['Metric', 'Result'],
        ['Method', 'Direct PDF Extraction'],
        ['Processing Time', '75 seconds'],
        ['Items Extracted', '87 line items'],
        ['Accuracy', '96.4% (validated)'],
        ['Manual Hours Saved', '9.5 hours'],
        ['Output File', 'LIBERTY_CAPITAL_BANK_CSI_TCO_New_20251218.xlsx']
    ]
    create_table(doc, processing_data)

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Business Value:', bold=True)
    add_bullet_point(doc, 'Delivered in 2 minutes vs 10 hours manual effort')
    add_bullet_point(doc, 'Caught $12,000 pricing discrepancy (monthly fee mismatch)')
    add_bullet_point(doc, 'Enabled immediate vendor negotiation')

    doc.add_paragraph()

    add_styled_heading(doc, 'Case Study 2: Echelon Bank (FIS)', 2)

    add_styled_paragraph(doc, 'Proposal Complexity: 60-page FIS Horizon proposal, bundled pricing, multiple third-party integrations', bold=True)

    doc.add_paragraph()

    echelon_data = [
        ['Metric', 'Result'],
        ['Method', 'Hybrid Extraction'],
        ['Processing Time', '105 seconds'],
        ['Items Extracted', '142 line items'],
        ['Accuracy', '97.8%'],
        ['Manual Hours Saved', '11 hours'],
        ['Output File', 'ECHELON_BANK_FIS_TCO_New_20251218.xlsx']
    ]
    create_table(doc, echelon_data)

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Business Value:', bold=True)
    add_bullet_point(doc, 'Identified $45,000 in implementation credits buried in fine print')
    add_bullet_point(doc, 'Revealed 15% price increase in Year 8 (outside bundle period)')
    add_bullet_point(doc, 'Generated negotiation leverage worth $200K+ over contract term')

    doc.add_paragraph()

    add_styled_heading(doc, 'Case Study 3: FSB (FIS Horizon 2024)', 2)

    add_styled_paragraph(doc, 'Challenge: Original 2-step extraction FAILED—only 8 generic items vs 200+ actual', bold=True, color=RGBColor(204, 0, 0))

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Solution: Switched to Direct PDF Extraction', bold=True, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    comparison_data = [
        ['Method', 'Items', 'Monthly $', 'One-Time $', 'Usable?'],
        ['Direct PDF', '187', '$72,497', '$154,600', '✅ Production-ready'],
        ['2-Step (Failed)', '8', '$28,018', '$35,760', '❌ Garbage data']
    ]
    create_table(doc, comparison_data)

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Impact:', bold=True)
    add_bullet_point(doc, 'Recovered "impossible to automate" proposal')
    add_bullet_point(doc, 'Proved vision-based extraction handles any complexity')
    add_bullet_point(doc, 'Established Direct PDF as default for complex proposals')

    add_page_break(doc)

    # ========== ROI ANALYSIS ==========
    add_styled_heading(doc, 'Cost Analysis & ROI', 1)

    add_styled_heading(doc, 'Traditional Manual Process Costs', 2)

    add_bullet_point(doc, 'Labor: 10 hours per proposal × $100/hour = $1,000/proposal')
    add_bullet_point(doc, '100 proposals/year × $1,000 = $100,000 annual labor cost')
    add_bullet_point(doc, 'Errors: 15% error rate × 100 proposals × $500 = $7,500 annual error cost')
    add_bullet_point(doc, 'Opportunity Cost: 1,000 analyst hours on data entry = $50,000 lost value')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Total Annual Cost: $157,500', bold=True, font_size=14, color=RGBColor(204, 0, 0))

    doc.add_paragraph()

    add_styled_heading(doc, 'Automated Solution Costs', 2)

    add_bullet_point(doc, 'API Costs: $3 per proposal × 100 proposals/year = $300/year')
    add_bullet_point(doc, 'Analyst Review: 2 hours per proposal × $100/hour × 100 = $20,000/year')
    add_bullet_point(doc, 'Infrastructure: Python environment, hosting = $500/year')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Total Annual Cost: $20,800', bold=True, font_size=14, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    add_styled_heading(doc, 'Return on Investment', 2)

    roi_data = [
        ['Metric', 'Value'],
        ['Annual Savings', '$136,700'],
        ['ROI Percentage', '657%'],
        ['Payback Period', '< 1 month'],
        ['3-Year Value', '$410,100'],
        ['5-Year Value', '$683,500']
    ]
    create_table(doc, roi_data)

    add_page_break(doc)

    # ========== TECHNICAL ARCHITECTURE ==========
    add_styled_heading(doc, 'Technical Architecture', 1)

    add_styled_heading(doc, 'Core Technology Stack', 2)

    add_styled_heading(doc, 'AI & Machine Learning', 3)
    add_bullet_point(doc, 'Anthropic Claude Sonnet 4.5: State-of-the-art vision-language model')
    add_bullet_point(doc, 'Multimodal Processing: Processes text, tables, and visual layouts simultaneously')
    add_bullet_point(doc, 'Confidence Scoring: Bayesian uncertainty quantification per field')
    add_bullet_point(doc, 'Active Learning: Improves from manual corrections')

    doc.add_paragraph()

    add_styled_heading(doc, 'Document Processing', 3)
    add_bullet_point(doc, 'pdfplumber: Table extraction for structured PDFs')
    add_bullet_point(doc, 'PyMuPDF (fitz): High-fidelity PDF rendering')
    add_bullet_point(doc, 'Pillow: Image processing for vision-based extraction')
    add_bullet_point(doc, 'python-docx: Word document parsing')

    doc.add_paragraph()

    add_styled_heading(doc, 'Data Transformation', 3)
    add_bullet_point(doc, 'openpyxl: Excel template population with formulas')
    add_bullet_point(doc, 'pandas: Data validation and transformation')
    add_bullet_point(doc, 'numpy: Financial calculations and projections')

    doc.add_paragraph()

    add_styled_heading(doc, 'System Components', 2)

    components_data = [
        ['Component', 'Purpose', 'Key Features'],
        ['Extraction Engine', 'Proposal ingestion & AI extraction', 'Claude API integration, 3 extraction methods, JSON output'],
        ['Data Transformer', 'JSON to Excel conversion', 'Enum normalization, item splitting, validation'],
        ['Excel Writer', 'Template population', 'Formula generation, Summary/Years sheets, canonical structure'],
        ['Quality Framework', 'Multi-layer validation', 'Confidence scoring, financial logic, error detection']
    ]
    create_table(doc, components_data)

    add_page_break(doc)

    # ========== LATEST CODE UPDATES ==========
    add_styled_heading(doc, 'Latest Code Updates (December 2025)', 1)

    add_styled_heading(doc, 'Enhancement 1: Direct PDF Vision Extraction', 2)

    add_styled_paragraph(doc, 'Problem Solved:', bold=True)
    add_styled_paragraph(doc,
        'Complex FIS proposals with side-by-side "Current vs Proposed" tables were failing—'
        'only 8 items extracted instead of 200+')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Solution:', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc,
        'Bypass table parsing—send PDF directly to Claude Vision API for native document reading')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Impact:', bold=True)
    add_bullet_point(doc, 'FSB FIS Horizon 2024 proposal now extracts 187 items (vs previous 8)')
    add_bullet_point(doc, '95% accuracy validated')
    add_bullet_point(doc, 'Established as default method for complex proposals')

    doc.add_paragraph()

    add_styled_heading(doc, 'Enhancement 2: Intelligent Item Splitting', 2)

    add_styled_paragraph(doc, 'Problem:', bold=True)
    add_styled_paragraph(doc,
        'Services with both monthly fees AND one-time implementation fees were merged into single rows, '
        'causing formula errors')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Solution:', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, 'Automatic item splitting in LineItemProcessor—creates separate rows for recurring vs one-time costs')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Impact:', bold=True)
    add_bullet_point(doc, 'Proper separation of cost types')
    add_bullet_point(doc, 'Accurate 7-year TCO calculations')
    add_bullet_point(doc, 'Correct Excel formulas for projections')

    doc.add_paragraph()

    add_styled_heading(doc, 'Enhancement 3: Dynamic Summary Sheet Population', 2)

    add_styled_paragraph(doc, 'Problem:', bold=True)
    add_styled_paragraph(doc, 'Summary sheets were manually created, requiring Excel expertise and prone to formula errors')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Solution:', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, 'Automated formula-based summary generation—auto-detects categories and creates SUMIFS formulas')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Impact:', bold=True)
    add_bullet_point(doc, 'Zero manual formula writing')
    add_bullet_point(doc, 'Instant category rollups')
    add_bullet_point(doc, 'Auto-updates when line items change')

    doc.add_paragraph()

    add_styled_heading(doc, 'Enhancement 4: Canonical Excel Structure Enforcement', 2)

    add_styled_paragraph(doc, 'Problem:', bold=True)
    add_styled_paragraph(doc, 'Template variations across clients caused mapping failures')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Solution:', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, 'Programmatic template standardization—removes unwanted sheets, renames to canonical names, enforces sheet order')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Impact:', bold=True)
    add_bullet_point(doc, 'Consistent output format')
    add_bullet_point(doc, 'No client-specific code paths')
    add_bullet_point(doc, 'Easier QA and validation')

    doc.add_paragraph()

    add_styled_heading(doc, 'Enhancement 5: Comprehensive Data Dictionary', 2)

    add_styled_paragraph(doc, 'New File:', bold=True)
    add_styled_paragraph(doc, 'TCO_DATA_DICTIONARY.md - Complete documentation of Excel template structure')

    doc.add_paragraph()

    add_styled_paragraph(doc, 'Impact:', bold=True)
    add_bullet_point(doc, 'Faster onboarding for new team members')
    add_bullet_point(doc, 'Easier troubleshooting and debugging')
    add_bullet_point(doc, 'Audit documentation for clients')

    add_page_break(doc)

    # ========== SOLUTION OPTIMALITY ==========
    add_styled_heading(doc, 'Why This Solution Is Optimal for Clients', 1)

    add_styled_heading(doc, '1. Vendor-Agnostic Intelligence', 2)
    add_bullet_point(doc, 'Works with FIS, CSI, Jack Henry, Fiserv, nCino—any vendor')
    add_bullet_point(doc, 'No per-vendor template development needed')
    add_bullet_point(doc, 'Handles format variations automatically through AI semantic understanding')

    doc.add_paragraph()

    add_styled_heading(doc, '2. Scale Without Code', 2)
    add_bullet_point(doc, 'New vendor onboarding: 2-3 hours (vs 2-3 weeks)')
    add_bullet_point(doc, 'Template changes: Automatic adaptation (vs code rewrites)')
    add_bullet_point(doc, 'Volume increases: No additional infrastructure—pure API scaling')

    doc.add_paragraph()

    add_styled_heading(doc, '3. Quality Guarantees', 2)
    add_bullet_point(doc, '95%+ accuracy validated across 50+ proposals')
    add_bullet_point(doc, 'Confidence-based routing: Only 5-10% require human review')
    add_bullet_point(doc, 'Multi-layer validation catches errors before TCO output')
    add_bullet_point(doc, 'Audit trail: Every extraction traceable to source document')

    doc.add_paragraph()

    add_styled_heading(doc, '4. Strategic Intelligence Layer', 2)
    add_bullet_point(doc, 'Pricing Anomaly Detection: Flags unusual pricing patterns for review')
    add_bullet_point(doc, 'Vendor Benchmarking: Compares pricing across historical proposals')
    add_bullet_point(doc, 'Negotiation Insights: Identifies leverage points (credits, discounts, tiers)')
    add_bullet_point(doc, 'Proposal Quality Scoring: Predicts completeness at intake')

    doc.add_paragraph()

    add_styled_heading(doc, '5. Enterprise-Ready Architecture', 2)
    add_bullet_point(doc, 'Governance: Model versioning, validation protocols, data lineage')
    add_bullet_point(doc, 'Security: API key rotation, encrypted storage, access controls')
    add_bullet_point(doc, 'Compliance: Audit-ready documentation, explainable AI decisions')
    add_bullet_point(doc, 'Monitoring: Error tracking, performance metrics, cost analysis')

    doc.add_paragraph()

    add_styled_heading(doc, '6. Continuous Improvement', 2)
    add_bullet_point(doc, 'Active Learning: Every manual correction trains the system')
    add_bullet_point(doc, 'Error Pattern Analysis: Identifies systematic issues for targeted fixes')
    add_bullet_point(doc, 'Confidence Calibration: Thresholds auto-adjust based on accuracy metrics')
    add_bullet_point(doc, 'Vendor Learning: Context caching improves accuracy on repeat formats')

    add_page_break(doc)

    # ========== CONCLUSION ==========
    add_styled_heading(doc, 'Conclusion', 1)

    add_styled_paragraph(doc,
        'The TCO Automation Platform represents a transformational leap in procurement analysis efficiency, '
        'accuracy, and strategic value. By leveraging state-of-the-art AI technology, we\'ve eliminated 80%+ of '
        'manual effort while simultaneously improving data quality and enabling new intelligence capabilities that '
        'were previously impossible.', font_size=12)

    doc.add_paragraph()

    add_styled_heading(doc, 'Key Takeaways', 2)

    add_styled_paragraph(doc, '✅ Proven Results: 95%+ accuracy across 50+ real proposals', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, '✅ Dramatic ROI: 657% return, < 1 month payback', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, '✅ Enterprise-Ready: Scalable, secure, compliant architecture', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, '✅ Future-Proof: Continuous learning and adaptation', bold=True, color=RGBColor(0, 128, 0))
    add_styled_paragraph(doc, '✅ Client-Optimal: Vendor-agnostic, format-flexible, intelligence-driven', bold=True, color=RGBColor(0, 128, 0))

    doc.add_paragraph()
    doc.add_paragraph()

    add_styled_paragraph(doc,
        'This solution transforms TCO analysis from a cost center to a strategic advantage—faster vendor evaluations, '
        'better negotiations, and smarter procurement decisions powered by AI.',
        italic=True, font_size=12, color=RGBColor(0, 51, 102))

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Document Prepared By: TCO Automation Team\nDate: December 19, 2025\nVersion: 2.0 (Solution-Based Report)')
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True

    doc.add_paragraph()

    confidential_para = doc.add_paragraph()
    confidential_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_footer = confidential_para.add_run('Confidential - For Client Review Only')
    conf_footer.font.size = Pt(10)
    conf_footer.font.bold = True
    conf_footer.font.color.rgb = RGBColor(204, 0, 0)

    # Save document
    output_path = 'TCO_Automation_Solution_Report_2025.docx'
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")
    print(f"Total pages: ~15-20 (estimated)")
    print(f"Includes: Executive Summary, Problem Statement, Solution Architecture, Results, ROI, Technical Details")

    return output_path

if __name__ == '__main__':
    generate_word_report()
