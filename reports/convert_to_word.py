"""
Convert Markdown files to Word documents with proper formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_word(md_file, output_file):
    """Convert Markdown file to Word document"""

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    setup_styles(doc)

    # Split content into lines
    lines = content.split('\n')

    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and line.strip():
            if not in_table:
                in_table = True
                table_rows = []

            # Skip separator rows
            if re.match(r'^[\|\-\s:]+$', line):
                i += 1
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells
            table_rows.append(cells)
            i += 1
            continue
        else:
            # End of table
            if in_table:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False

        # Handle page breaks
        if line.strip() == '---' and i > 0 and lines[i-1].strip() == '':
            doc.add_page_break()
            i += 1
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()

            if level == 1:
                add_heading(doc, text, 1)
            elif level == 2:
                add_heading(doc, text, 2)
            elif level == 3:
                add_heading(doc, text, 3)
            else:
                add_heading(doc, text, 4)
            i += 1
            continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            add_bullet(doc, text, 0)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            add_numbered(doc, text)
            i += 1
            continue

        # Handle checkboxes
        if '- [ ]' in line or '- [x]' in line or '- [X]' in line:
            checked = '[x]' in line.lower()
            text = re.sub(r'-\s*\[.\]\s*', '', line.strip())
            symbol = '☑' if checked else '☐'
            add_bullet(doc, f"{symbol} {text}", 0)
            i += 1
            continue

        # Handle blockquotes
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            add_quote(doc, text)
            i += 1
            continue

        # Handle empty lines
        if not line.strip():
            # Don't add multiple empty paragraphs in a row
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()
            i += 1
            continue

        # Regular paragraph with inline formatting
        if line.strip():
            add_paragraph_with_formatting(doc, line.strip())

        i += 1

    # Add any remaining table
    if table_rows:
        add_table(doc, table_rows)

    # Save document
    doc.save(output_file)
    print(f"Created: {output_file}")

def setup_styles(doc):
    """Set up custom styles for the document"""
    styles = doc.styles

    # Heading styles
    for i in range(1, 5):
        style_name = f'Heading {i}'
        if style_name in styles:
            heading = styles[style_name]
            heading.font.color.rgb = RGBColor(0, 51, 102)
            if i == 1:
                heading.font.size = Pt(18)
            elif i == 2:
                heading.font.size = Pt(16)
            elif i == 3:
                heading.font.size = Pt(14)
            else:
                heading.font.size = Pt(12)

def add_heading(doc, text, level):
    """Add a heading with specified level"""
    h = doc.add_heading(text, level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

def add_paragraph_with_formatting(doc, text):
    """Add paragraph with bold, italic, and code formatting"""
    p = doc.add_paragraph()

    # Split by code blocks first
    parts = re.split(r'`([^`]+)`', text)

    for i, part in enumerate(parts):
        if i % 2 == 1:  # Code
            run = p.add_run(part)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 0, 0)
        else:
            # Handle bold and italic
            segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)', part)

            for segment in segments:
                if segment.startswith('**') and segment.endswith('**'):
                    run = p.add_run(segment[2:-2])
                    run.bold = True
                elif segment.startswith('__') and segment.endswith('__'):
                    run = p.add_run(segment[2:-2])
                    run.bold = True
                elif segment.startswith('*') and segment.endswith('*'):
                    run = p.add_run(segment[1:-1])
                    run.italic = True
                elif segment.startswith('_') and segment.endswith('_'):
                    run = p.add_run(segment[1:-1])
                    run.italic = True
                else:
                    p.add_run(segment)

def add_bullet(doc, text, level):
    """Add a bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    add_inline_formatting(p, text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)

def add_numbered(doc, text):
    """Add a numbered item"""
    p = doc.add_paragraph(style='List Number')
    add_inline_formatting(p, text)

def add_inline_formatting(p, text):
    """Add text with inline formatting to paragraph"""
    # Handle code
    parts = re.split(r'`([^`]+)`', text)
    for i, part in enumerate(parts):
        if i % 2 == 1:
            run = p.add_run(part)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            # Handle bold/italic
            segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', part)
            for segment in segments:
                if segment.startswith('**') and segment.endswith('**'):
                    run = p.add_run(segment[2:-2])
                    run.bold = True
                elif segment.startswith('*') and segment.endswith('*'):
                    run = p.add_run(segment[1:-1])
                    run.italic = True
                else:
                    if segment:
                        p.add_run(segment)

def add_quote(doc, text):
    """Add a blockquote"""
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)

    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(96, 96, 96)

def add_code_block(doc, lines):
    """Add a code block"""
    code_text = '\n'.join(lines)
    p = doc.add_paragraph(code_text)
    p.style = 'No Spacing'

    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Add background color
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)

def add_table(doc, rows):
    """Add a table from rows"""
    if not rows or len(rows) < 2:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Populate table
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_data

                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    # Add spacing after table
    doc.add_paragraph()

if __name__ == '__main__':
    # Convert both documents
    parse_markdown_to_word(
        'TCO_Quick_Implementation_Plan.md',
        'TCO_Quick_Implementation_Plan.docx'
    )

    parse_markdown_to_word(
        'TCO_Detailed_Project_Plan.md',
        'TCO_Detailed_Project_Plan.docx'
    )

    print("\nConversion complete!")
    print("Created:")
    print("  - TCO_Quick_Implementation_Plan.docx")
    print("  - TCO_Detailed_Project_Plan.docx")
