"""
Document Loader Module

Unified document ingestion for PDFs, DOCX, images with OCR support.
Normalizes all document types to structured text for downstream processing.
"""

import os
import logging
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from dataclasses import dataclass, field
from enum import Enum

# PDF processing
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# OCR
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Word documents
from docx import Document

# Excel
import openpyxl


logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class ExtractedPage:
    """Represents a single page/sheet of extracted content."""
    page_number: int
    text: str
    tables: List[List[List[str]]] = field(default_factory=list)
    images: List[bytes] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractedDocument:
    """Complete extracted document with all pages and metadata."""
    file_path: str
    document_type: DocumentType
    pages: List[ExtractedPage]
    metadata: Dict[str, Any]
    raw_text: str  # Combined text from all pages
    tables: List[List[List[str]]]  # All tables combined
    ocr_applied: bool = False

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def has_tables(self) -> bool:
        return len(self.tables) > 0


class DocumentLoader:
    """
    Unified document loader supporting multiple formats.

    Handles:
    - PDF files (text extraction + OCR for scanned pages)
    - Word documents (.docx)
    - Excel spreadsheets (.xlsx)
    - Images (OCR)
    """

    # Supported image extensions
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif'}

    def __init__(self, ocr_language: str = 'eng', ocr_config: str = ''):
        """
        Initialize the document loader.

        Args:
            ocr_language: Tesseract language code (default: 'eng')
            ocr_config: Additional Tesseract configuration
        """
        self.ocr_language = ocr_language
        self.ocr_config = ocr_config

        # Check available libraries
        self._check_dependencies()

    def _check_dependencies(self):
        """Log available processing capabilities."""
        if not HAS_PDFPLUMBER:
            logger.warning("pdfplumber not installed - PDF text extraction limited")
        if not HAS_PYMUPDF:
            logger.warning("PyMuPDF not installed - PDF image extraction unavailable")
        if not HAS_OCR:
            logger.warning("pytesseract/Pillow not installed - OCR unavailable")

    def detect_type(self, file_path: str) -> DocumentType:
        """
        Detect document type from file extension.

        Args:
            file_path: Path to the document

        Returns:
            DocumentType enum value
        """
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentType.PDF
        elif ext == '.docx':
            return DocumentType.DOCX
        elif ext in ('.xlsx', '.xls'):
            return DocumentType.XLSX
        elif ext in self.IMAGE_EXTENSIONS:
            return DocumentType.IMAGE
        else:
            return DocumentType.UNKNOWN

    def load(self, file_path: str, apply_ocr: bool = True) -> ExtractedDocument:
        """
        Load and extract content from a document.

        Args:
            file_path: Path to the document
            apply_ocr: Whether to apply OCR for images/scanned PDFs

        Returns:
            ExtractedDocument with all content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        doc_type = self.detect_type(file_path)

        if doc_type == DocumentType.PDF:
            return self._load_pdf(file_path, apply_ocr)
        elif doc_type == DocumentType.DOCX:
            return self._load_docx(file_path)
        elif doc_type == DocumentType.XLSX:
            return self._load_xlsx(file_path)
        elif doc_type == DocumentType.IMAGE:
            return self._load_image(file_path)
        else:
            raise ValueError(f"Unsupported document type: {file_path}")

    def _load_pdf(self, file_path: str, apply_ocr: bool = True) -> ExtractedDocument:
        """
        Extract content from PDF file.

        Uses pdfplumber for text/tables, PyMuPDF for images,
        and OCR as fallback for scanned pages.
        """
        pages = []
        all_tables = []
        all_text = []
        ocr_applied = False
        metadata = {}

        if not HAS_PDFPLUMBER:
            raise ImportError("pdfplumber required for PDF processing")

        with pdfplumber.open(file_path) as pdf:
            metadata = {
                'page_count': len(pdf.pages),
                'metadata': pdf.metadata or {}
            }

            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ''
                page_tables = []

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    page_tables = tables
                    all_tables.extend(tables)

                # Check if page needs OCR (little to no text but has images)
                if apply_ocr and len(page_text.strip()) < 50 and HAS_OCR and HAS_PYMUPDF:
                    ocr_text = self._ocr_pdf_page(file_path, i)
                    if ocr_text:
                        page_text = ocr_text
                        ocr_applied = True

                all_text.append(page_text)

                extracted_page = ExtractedPage(
                    page_number=i + 1,
                    text=page_text,
                    tables=page_tables,
                    metadata={'width': page.width, 'height': page.height}
                )
                pages.append(extracted_page)

        return ExtractedDocument(
            file_path=file_path,
            document_type=DocumentType.PDF,
            pages=pages,
            metadata=metadata,
            raw_text='\n\n'.join(all_text),
            tables=all_tables,
            ocr_applied=ocr_applied
        )

    def _ocr_pdf_page(self, file_path: str, page_num: int) -> str:
        """
        Apply OCR to a specific PDF page.

        Args:
            file_path: Path to PDF
            page_num: Zero-indexed page number

        Returns:
            OCR extracted text
        """
        if not HAS_PYMUPDF or not HAS_OCR:
            return ''

        try:
            doc = fitz.open(file_path)
            page = doc[page_num]

            # Render page to image
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Apply OCR
            text = pytesseract.image_to_string(
                img,
                lang=self.ocr_language,
                config=self.ocr_config
            )

            doc.close()
            return text.strip()

        except Exception as e:
            logger.error(f"OCR failed for page {page_num}: {e}")
            return ''

    def _load_docx(self, file_path: str) -> ExtractedDocument:
        """
        Extract content from Word document.
        """
        doc = Document(file_path)

        all_text = []
        all_tables = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            all_tables.append(table_data)

        # Create single page for DOCX
        page = ExtractedPage(
            page_number=1,
            text='\n'.join(all_text),
            tables=all_tables,
            metadata={
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
        )

        metadata = {
            'core_properties': {
                'author': doc.core_properties.author,
                'title': doc.core_properties.title,
                'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None
            }
        }

        return ExtractedDocument(
            file_path=file_path,
            document_type=DocumentType.DOCX,
            pages=[page],
            metadata=metadata,
            raw_text='\n'.join(all_text),
            tables=all_tables
        )

    def _load_xlsx(self, file_path: str) -> ExtractedDocument:
        """
        Extract content from Excel spreadsheet.
        """
        wb = openpyxl.load_workbook(file_path, data_only=True)

        pages = []
        all_tables = []
        all_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Extract all data as table
            table_data = []
            text_parts = []

            for row in ws.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    table_data.append(row_data)
                    text_parts.append(' | '.join(row_data))

            if table_data:
                all_tables.append(table_data)

            sheet_text = '\n'.join(text_parts)
            all_text.append(f"=== {sheet_name} ===\n{sheet_text}")

            page = ExtractedPage(
                page_number=len(pages) + 1,
                text=sheet_text,
                tables=[table_data] if table_data else [],
                metadata={
                    'sheet_name': sheet_name,
                    'row_count': ws.max_row,
                    'col_count': ws.max_column
                }
            )
            pages.append(page)

        wb.close()

        return ExtractedDocument(
            file_path=file_path,
            document_type=DocumentType.XLSX,
            pages=pages,
            metadata={'sheet_count': len(wb.sheetnames), 'sheet_names': wb.sheetnames},
            raw_text='\n\n'.join(all_text),
            tables=all_tables
        )

    def _load_image(self, file_path: str) -> ExtractedDocument:
        """
        Extract text from image using OCR.
        """
        if not HAS_OCR:
            raise ImportError("pytesseract and Pillow required for image OCR")

        img = Image.open(file_path)

        # Apply OCR
        text = pytesseract.image_to_string(
            img,
            lang=self.ocr_language,
            config=self.ocr_config
        )

        page = ExtractedPage(
            page_number=1,
            text=text.strip(),
            metadata={
                'width': img.width,
                'height': img.height,
                'format': img.format,
                'mode': img.mode
            }
        )

        return ExtractedDocument(
            file_path=file_path,
            document_type=DocumentType.IMAGE,
            pages=[page],
            metadata={'image_size': img.size, 'image_format': img.format},
            raw_text=text.strip(),
            tables=[],
            ocr_applied=True
        )

    def load_directory(self, directory: str, recursive: bool = False) -> List[ExtractedDocument]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Path to directory
            recursive: Whether to search subdirectories

        Returns:
            List of ExtractedDocument objects
        """
        documents = []
        path = Path(directory)

        if recursive:
            files = path.rglob('*')
        else:
            files = path.glob('*')

        for file_path in files:
            if file_path.is_file():
                doc_type = self.detect_type(str(file_path))
                if doc_type != DocumentType.UNKNOWN:
                    try:
                        doc = self.load(str(file_path))
                        documents.append(doc)
                        logger.info(f"Loaded: {file_path}")
                    except Exception as e:
                        logger.error(f"Failed to load {file_path}: {e}")

        return documents


def load_document(file_path: str, apply_ocr: bool = True) -> ExtractedDocument:
    """
    Convenience function to load a single document.

    Args:
        file_path: Path to the document
        apply_ocr: Whether to apply OCR if needed

    Returns:
        ExtractedDocument with all content
    """
    loader = DocumentLoader()
    return loader.load(file_path, apply_ocr)


def load_documents(directory: str, recursive: bool = False) -> List[ExtractedDocument]:
    """
    Convenience function to load all documents from a directory.

    Args:
        directory: Path to directory
        recursive: Whether to search subdirectories

    Returns:
        List of ExtractedDocument objects
    """
    loader = DocumentLoader()
    return loader.load_directory(directory, recursive)


if __name__ == '__main__':
    # Example usage
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        doc = load_document(file_path)

        print(f"Document: {doc.file_path}")
        print(f"Type: {doc.document_type.value}")
        print(f"Pages: {doc.page_count}")
        print(f"Tables: {len(doc.tables)}")
        print(f"OCR Applied: {doc.ocr_applied}")
        print(f"\nText Preview (first 500 chars):")
        print(doc.raw_text[:500])
    else:
        print("Usage: python document_loader.py <file_path>")
