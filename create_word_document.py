"""
Script to convert QA Validation Processes markdown to Word document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_styled_heading(doc, text, level):
    """Add a heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(12)
    heading_format.space_after = Pt(6)
    return heading

def add_styled_paragraph(doc, text, style='Normal'):
    """Add a paragraph with custom styling"""
    para = doc.add_paragraph(text, style=style)
    para_format = para.paragraph_format
    para_format.space_after = Pt(6)
    para_format.line_spacing = 1.15
    return para

def add_bullet_point(doc, text):
    """Add a bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25)
    return para

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Add light gray background effect
    para.paragraph_format.line_spacing = 1.0
    return para

def add_table_from_markdown(doc, lines):
    """Create a table from markdown table lines"""
    # Parse markdown table
    rows = []
    for line in lines:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and not all(cell.startswith('-') for cell in cells):
                rows.append(cells)

    if not rows:
        return None

    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # Populate table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            # Bold header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    return table

def process_markdown_to_word(md_file, output_file):
    """Convert markdown file to Word document"""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Process lines
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                add_code_block(doc, code_text)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and '-|-' in line:
            in_table = True
            table_lines = [lines[i-1].rstrip(), line]  # Include header
            i += 1
            continue

        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                add_table_from_markdown(doc, table_lines)
                doc.add_paragraph()  # Add spacing after table
                table_lines = []
                in_table = False

        # Skip horizontal rules
        if line.startswith('---') or line.startswith('***'):
            doc.add_paragraph()  # Add spacing
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            text = line.replace('# ', '').strip()
            add_styled_heading(doc, text, 1)
        elif line.startswith('## '):
            text = line.replace('## ', '').strip()
            add_styled_heading(doc, text, 2)
        elif line.startswith('### '):
            text = line.replace('### ', '').strip()
            add_styled_heading(doc, text, 3)
        elif line.startswith('#### '):
            text = line.replace('#### ', '').strip()
            add_styled_heading(doc, text, 4)

        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = re.sub(r'^[-*]\s+', '', line)
            # Remove markdown bold/italic markers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            add_bullet_point(doc, text)

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            para = doc.add_paragraph(text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.25)

        # Handle blockquotes
        elif line.startswith('>'):
            text = line.replace('>', '').strip()
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Inches(0.5)
            para_format = para.paragraph_format
            para_format.space_before = Pt(6)
            para_format.space_after = Pt(6)
            # Add italic
            for run in para.runs:
                run.italic = True

        # Handle regular paragraphs
        elif line.strip():
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)

            # Check if it's a bold line (key metrics, etc.)
            if line.startswith('**') and line.endswith('**'):
                para = add_styled_paragraph(doc, text.strip('*'))
                for run in para.runs:
                    run.font.bold = True
            else:
                add_styled_paragraph(doc, text)

        # Empty line - just skip
        else:
            pass

        i += 1

    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    doc.save(output_file)
    print(f"Word document created successfully: {output_file}")

if __name__ == "__main__":
    md_file = r"C:\Users\dada_\OneDrive\Documents\TCO_Automation\QA_Validation_Processes_Client_Document.md"
    output_file = r"C:\Users\dada_\OneDrive\Documents\TCO_Automation\QA_Validation_Processes_Client_Document.docx"

    process_markdown_to_word(md_file, output_file)
