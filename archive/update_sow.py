"""
Update SOW Document with Project Plan Details
Updates timelines, automation metrics, and project details from Excel plan
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl

def read_project_plan_data():
    """Read key data from the Excel project plan"""
    wb = openpyxl.load_workbook('Karishma Proposal Documents/Project_Plan_Phase2_DETAILED_20260109_130815.xlsx')
    ws_milestones = wb['Milestones']

    # Get project summary from A2
    project_summary = ws_milestones['A2'].value

    # Extract milestones
    milestones = []
    for i in range(5, 30):
        milestone_id = ws_milestones.cell(i, 1).value
        if milestone_id and str(milestone_id).startswith('M') and ws_milestones.cell(i, 2).value:
            name = ws_milestones.cell(i, 2).value
            timeline = ws_milestones.cell(i, 4).value
            payment = ws_milestones.cell(i, 13).value
            deliverables = ws_milestones.cell(i, 7).value

            milestones.append({
                'id': milestone_id,
                'name': name,
                'timeline': timeline,
                'payment': payment,
                'deliverables': deliverables
            })

    return {
        'summary': project_summary,
        'milestones': milestones,
        'total_weeks': 28,
        'total_days': 134,
        'total_milestones': len(milestones)
    }

def update_sow_document():
    """Update SOW document with correct information"""

    # Load SOW document
    doc = Document('Karishma Proposal Documents/SOW_Phase2_TCO_Automation sm updates 1_9.docx')

    # Read project plan data
    plan_data = read_project_plan_data()

    print('[INFO] Updating SOW document...')
    print(f'Project timeline: {plan_data["total_weeks"]} weeks ({plan_data["total_days"]} days)')
    print(f'Total milestones: {plan_data["total_milestones"]}')

    # Track changes
    changes_made = []

    # Update paragraphs
    for i, para in enumerate(doc.paragraphs):
        original_text = para.text

        # 1. Update automation time from "under 60 seconds" to more realistic time
        if 'under 60 seconds per vendor' in para.text:
            para.text = para.text.replace(
                'under 60 seconds per vendor',
                '3-5 minutes per vendor (end-to-end including extraction, validation, and TCO generation)'
            )
            changes_made.append(f'Para {i}: Updated automation time from 60 seconds to 3-5 minutes')

        # 2. Update "Reduce manual TCO creation time from 2-4 hours"
        if 'Reduce manual TCO creation time from 2-4 hours' in para.text:
            para.text = para.text.replace(
                'Reduce manual TCO creation time from 2-4 hours to under 60 seconds per vendor',
                'Reduce manual TCO creation time from 2-4 hours to 3-5 minutes per vendor (90%+ time savings)'
            )
            changes_made.append(f'Para {i}: Updated time savings description')

    # Find and update sections that need project timeline details
    section_5_found = False
    section_10_found = False

    for i, para in enumerate(doc.paragraphs):
        # Update Section 5 - Project Plan
        if para.text.strip() == '5. Project Plan' or para.text.strip() == '5.2 Milestone Schedule':
            section_5_found = True

        if section_5_found and para.text.strip() == 'Critical Gate: Development work packages (WP1-WP8) do not begin until Discovery (M0) is complete and findings are accepted.':
            # Insert project timeline details before the critical gate
            insert_pos = i

            # Add timeline overview
            p1 = doc.paragraphs[insert_pos].insert_paragraph_before('Project Duration: 28 weeks (134 days / ~5 months)')
            p1.runs[0].bold = True
            p1.runs[0].font.size = Pt(11)

            p2 = doc.paragraphs[insert_pos + 1].insert_paragraph_before('')

            p3 = doc.paragraphs[insert_pos + 2].insert_paragraph_before('Milestone Overview:')
            p3.runs[0].bold = True
            p3.runs[0].font.size = Pt(11)

            # Add milestone summary
            milestone_summary_lines = []
            for m in plan_data['milestones']:
                payment_text = f" ({m['payment']} payment)" if m['payment'] and m['payment'] != '-' else ''
                milestone_summary_lines.append(f"• {m['id']}: {m['name']}{payment_text}")

            p4 = doc.paragraphs[insert_pos + 3].insert_paragraph_before('\n'.join(milestone_summary_lines))
            p4.runs[0].font.size = Pt(10)

            p5 = doc.paragraphs[insert_pos + 4].insert_paragraph_before('')

            changes_made.append(f'Section 5: Added project timeline overview and milestone summary')
            section_5_found = False
            break

    # Update Section 10 - Payment Schedule
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == '10. Commercial Terms' or para.text.strip() == '10.2 Payment Schedule':
            section_10_found = True

        if section_10_found and para.text.strip().startswith('Note:'):
            # Insert payment schedule before the note
            insert_pos = i

            p1 = doc.paragraphs[insert_pos].insert_paragraph_before('Payment Milestones:')
            p1.runs[0].bold = True
            p1.runs[0].font.size = Pt(11)

            # Add payment milestones
            payment_milestones = []
            for m in plan_data['milestones']:
                if m['payment'] and m['payment'] != '-':
                    payment_milestones.append(f"• {m['id']} ({m['name']}): {m['payment']}")

            payment_milestones.append('• Retention: 15% (warranty period)')

            p2 = doc.paragraphs[insert_pos + 1].insert_paragraph_before('\n'.join(payment_milestones))
            p2.runs[0].font.size = Pt(10)

            p3 = doc.paragraphs[insert_pos + 2].insert_paragraph_before('')

            changes_made.append(f'Section 10: Added payment schedule details')
            section_10_found = False
            break

    # Add deliverables details if section exists
    deliverables_section_found = False
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == '3. Deliverables' or para.text.strip() == '3.1 Software Deliverables':
            deliverables_section_found = True

        if deliverables_section_found and (para.text.strip() == '3.2 Documentation Deliverables' or para.text.strip() == '3.3 Process Deliverables'):
            # We're past the software deliverables section, add summary before
            insert_pos = i

            # Count deliverables
            total_deliverables = 0
            for m in plan_data['milestones']:
                if m['deliverables']:
                    # Count bullet points
                    total_deliverables += m['deliverables'].count('•')

            p1 = doc.paragraphs[insert_pos].insert_paragraph_before(f'Total Project Deliverables: {total_deliverables} items across {plan_data["total_milestones"]} milestones')
            p1.runs[0].bold = True
            p1.runs[0].font.size = Pt(10)

            p2 = doc.paragraphs[insert_pos + 1].insert_paragraph_before('')

            changes_made.append(f'Section 3: Added deliverables count')
            deliverables_section_found = False
            break

    # Save document
    doc.save('Karishma Proposal Documents/SOW_Phase2_TCO_Automation sm updates 1_9.docx')

    print('\n[SUCCESS] SOW document updated!')
    print(f'\nChanges made ({len(changes_made)}):')
    for change in changes_made:
        print(f'  • {change}')

    print(f'\nKey updates:')
    print(f'  - Automation time: 60 seconds -> 3-5 minutes (more realistic)')
    print(f'  - Time savings: Now shows 90%+ savings (2-4 hours -> 3-5 min)')
    print(f'  - Project duration: {plan_data["total_weeks"]} weeks ({plan_data["total_days"]} days / ~5 months)')
    print(f'  - Milestone count: {plan_data["total_milestones"]}')
    print(f'  - Total deliverables: Added count across all milestones')
    print(f'  - Payment schedule: Added detailed milestone breakdown')

if __name__ == '__main__':
    update_sow_document()
