"""
Generate SOW and Project Plan documents in Word and Excel formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_to_doc(doc, headers, rows, header_color="1F4E79"):
    """Add a formatted table to the document"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = str(cell_data) if cell_data else ""

    return table

def create_sow_word_document():
    """Create the Statement of Work as a Word document"""
    doc = Document()

    # Title
    title = doc.add_heading('Statement of Work (SOW)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('TCO Automation System - Phase 2: Enterprise Scale-Up', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document info
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_data = [
        ("Document Version:", "1.0"),
        ("Date:", datetime.now().strftime("%B %d, %Y")),
        ("Prepared For:", "[Client Name]"),
        ("Prepared By:", "[Consultant Name]")
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = value

    doc.add_paragraph()
    doc.add_paragraph("_" * 80)

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph(
        "Phase 1 of the TCO Automation System has been successfully completed, demonstrating the ability to:"
    )
    bullets = [
        "Extract pricing data from FIS Word proposals and Jack Henry Excel deal sheets",
        "Normalize vendor-specific data into a standardized TCO schema with confidence scoring",
        "Populate Excel TCO templates with side-by-side vendor comparisons",
        "Achieve 95-99% extraction accuracy with complete audit trails",
        "Reduce manual TCO creation time from 2-4 hours to under 60 seconds per vendor"
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('Phase 1 Deliverables Completed:', level=3)
    deliverables = [
        ("TCO_Test_Output.xlsx", "FIS-only populated TCO template (42 normalized line items)"),
        ("TCO_Complete_Comparison.xlsx", "Side-by-side vendor comparison (176 total line items)"),
        ("Extracted JSON Files", "6 validated extraction outputs with confidence scores")
    ]
    for name, desc in deliverables:
        p = doc.add_paragraph(style='List Number')
        p.add_run(name).bold = True
        p.add_run(f" - {desc}")

    doc.add_heading('1.2 Phase 2 Objectives', level=2)
    objectives = [
        ("Volume Processing:", "Handle 20+ new proposals per month"),
        ("Historical Analysis:", "Process and analyze historical proposal archives"),
        ("Multi-Model Support:", "Accommodate multiple TCO output models (various template formats)"),
        ("Production Operations:", "Establish sustainable operational workflows")
    ]
    for title, desc in objectives:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f" {desc}")

    # Section 2: Scope of Work
    doc.add_page_break()
    doc.add_heading('2. Scope of Work', level=1)

    # Discovery Phase
    doc.add_heading('2.0 Discovery Phase (Pre-Development)', level=2)
    doc.add_paragraph(
        "Before development begins, a structured Discovery Phase will validate assumptions, "
        "gather requirements, and reduce project risk. Discovery findings will directly inform "
        "the technical approach and may result in scope adjustments."
    )

    doc.add_heading('2.0.1 Discovery Objectives', level=3)
    discovery_objectives = [
        ("D0.1", "Validate volume assumptions", "Confirmed monthly proposal volume and peak periods"),
        ("D0.2", "Inventory historical proposals", "File count, formats, date ranges, storage locations"),
        ("D0.3", "Catalog TCO templates", "Complete list of template variants with sample files"),
        ("D0.4", "Assess legacy format complexity", "Format compatibility matrix with effort estimates"),
        ("D0.5", "Define exception handling requirements", "Business rules for manual review triggers"),
        ("D0.6", "Identify integration touchpoints", "Upstream/downstream system dependencies"),
        ("D0.7", "Document operational workflows", "Current state vs. future state process maps"),
    ]
    add_table_to_doc(doc, ["ID", "Objective", "Output"], discovery_objectives)

    doc.add_heading('2.0.2 Discovery Activities', level=3)
    discovery_activities = [
        ("Stakeholder Interviews", "Meet with analysts, managers, IT to gather requirements", "2-3 sessions"),
        ("Historical Data Analysis", "Sample and analyze historical proposal files", "Collaborative"),
        ("Template Inventory", "Collect and document all TCO template variants", "Collaborative"),
        ("Current State Mapping", "Document existing manual TCO creation process", "1 session"),
        ("Technical Assessment", "Evaluate infrastructure, access, deployment options", "1 session"),
        ("Discovery Findings Review", "Present findings and confirm scope", "1 session"),
    ]
    add_table_to_doc(doc, ["Activity", "Description", "Duration"], discovery_activities)

    doc.add_heading('2.0.3 Discovery Deliverables', level=3)
    discovery_deliverables = [
        ("DD1", "Discovery Report", "Comprehensive findings document with recommendations"),
        ("DD2", "Historical Proposal Inventory", "Catalog of all historical files with metadata"),
        ("DD3", "Template Catalog", "All TCO templates with mapping requirements"),
        ("DD4", "Format Compatibility Matrix", "Support level for each document format found"),
        ("DD5", "Updated Risk Register", "Risks identified during discovery with mitigations"),
        ("DD6", "Refined Scope Document", "Any scope adjustments based on discovery findings"),
    ]
    add_table_to_doc(doc, ["#", "Deliverable", "Description"], discovery_deliverables)

    doc.add_heading('2.0.4 Discovery Exit Criteria', level=3)
    doc.add_paragraph("Discovery Phase is complete when:")
    exit_criteria = [
        "All discovery questions (Appendix D) have been answered",
        "Historical proposal inventory is complete (file count, formats, locations)",
        "All TCO template variants have been collected and cataloged",
        "Stakeholder sign-off on discovery findings",
        "Scope adjustments (if any) have been agreed upon",
        "Development phase can proceed with validated requirements"
    ]
    for criteria in exit_criteria:
        doc.add_paragraph(f"[ ] {criteria}", style='List Bullet')

    # In-Scope Activities
    doc.add_page_break()
    doc.add_heading('2.1 In-Scope Activities', level=2)

    # Volume Processing
    doc.add_heading('2.1.1 Volume Processing Infrastructure', level=3)
    volume_items = [
        ("2.1.1.1", "Batch Processing Engine", "Automated system to process multiple proposals in queue"),
        ("2.1.1.2", "Job Scheduler", "Configurable scheduling for processing windows (daily/weekly)"),
        ("2.1.1.3", "Progress Dashboard", "Real-time visibility into processing status and queue"),
        ("2.1.1.4", "Notification System", "Email/Slack alerts for completion, errors, and exceptions"),
        ("2.1.1.5", "Parallel Processing", "Concurrent processing of multiple proposals"),
    ]
    add_table_to_doc(doc, ["ID", "Deliverable", "Description"], volume_items)

    # Historical Processing
    doc.add_heading('2.1.2 Historical Proposal Processing', level=3)
    historical_items = [
        ("2.1.2.1", "Bulk Import Utility", "Tool to ingest historical proposal archives"),
        ("2.1.2.2", "Document Classification", "Auto-detect vendor type (FIS/JH) and document format"),
        ("2.1.2.3", "Version Handling", "Support for legacy document formats and variations"),
        ("2.1.2.4", "De-duplication Logic", "Identify and handle duplicate proposals"),
        ("2.1.2.5", "Historical Data Store", "Structured storage for processed historical data"),
    ]
    add_table_to_doc(doc, ["ID", "Deliverable", "Description"], historical_items)

    # Multi-Model TCO
    doc.add_heading('2.1.3 Multi-Model TCO Support', level=3)
    template_items = [
        ("2.1.3.1", "Template Registry", "Centralized management of multiple TCO templates"),
        ("2.1.3.2", "Template Detection", "Auto-match proposals to appropriate TCO model"),
        ("2.1.3.3", "Dynamic Column Mapping", "Configurable field-to-column mappings per template"),
        ("2.1.3.4", "Template Versioning", "Track and manage template versions"),
        ("2.1.3.5", "Custom Template Support", "Ability to add new TCO models without code changes"),
    ]
    add_table_to_doc(doc, ["ID", "Deliverable", "Description"], template_items)

    # QA
    doc.add_heading('2.1.4 Quality Assurance & Validation', level=3)
    qa_items = [
        ("2.1.4.1", "Confidence Threshold Rules", "Configurable thresholds per field type"),
        ("2.1.4.2", "Exception Queue", "Manual review workflow for low-confidence extractions"),
        ("2.1.4.3", "Validation Reports", "Detailed accuracy reports per batch"),
        ("2.1.4.4", "Audit Trail Enhancement", "Complete lineage from source to output"),
        ("2.1.4.5", "Reconciliation Tools", "Compare automated vs. manual TCO outputs"),
    ]
    add_table_to_doc(doc, ["ID", "Deliverable", "Description"], qa_items)

    # Operations
    doc.add_heading('2.1.5 Operational Infrastructure', level=3)
    ops_items = [
        ("2.1.5.1", "Configuration Management", "Environment-based settings (dev/test/prod)"),
        ("2.1.5.2", "Logging & Monitoring", "Centralized logging with error tracking"),
        ("2.1.5.3", "Backup & Recovery", "Data backup and disaster recovery procedures"),
        ("2.1.5.4", "Security Hardening", "Access controls, encryption, secure storage"),
        ("2.1.5.5", "Performance Optimization", "Handle 20+ proposals/month efficiently"),
    ]
    add_table_to_doc(doc, ["ID", "Deliverable", "Description"], ops_items)

    # Out of Scope
    doc.add_heading('2.2 Out-of-Scope Activities', level=2)
    doc.add_paragraph("The following items are explicitly excluded from this engagement:")
    out_of_scope = [
        "Integration with external CRM/ERP systems (future phase)",
        "Mobile application development",
        "Custom vendor extractors beyond FIS and Jack Henry",
        "Cloud hosting infrastructure setup (on-premise deployment only)",
        "End-user training beyond documented procedures",
        "Support for non-Excel TCO output formats"
    ]
    for item in out_of_scope:
        doc.add_paragraph(item, style='List Bullet')

    # Section 3: Deliverables
    doc.add_page_break()
    doc.add_heading('3. Deliverables', level=1)

    doc.add_heading('3.1 Software Deliverables', level=2)
    sw_deliverables = [
        ("D1", "Batch Processing Module", "Python Package", "Queue management and parallel processing"),
        ("D2", "Job Scheduler", "Python + Config", "Automated scheduling with cron-like syntax"),
        ("D3", "Template Registry System", "Python + JSON", "Multi-template management and mapping"),
        ("D4", "Bulk Import Utility", "CLI Tool", "Historical proposal ingestion"),
        ("D5", "Exception Management UI", "Web Interface", "Review queue for manual validation"),
        ("D6", "Monitoring Dashboard", "Web Interface", "Real-time processing status"),
        ("D7", "Reporting Module", "Python + Excel", "Batch processing and accuracy reports"),
    ]
    add_table_to_doc(doc, ["#", "Deliverable", "Format", "Description"], sw_deliverables)

    doc.add_heading('3.2 Documentation Deliverables', level=2)
    doc_deliverables = [
        ("D8", "Operations Manual", "Step-by-step procedures for daily operations"),
        ("D9", "Template Configuration Guide", "How to add/modify TCO templates"),
        ("D10", "Troubleshooting Guide", "Common issues and resolution steps"),
        ("D11", "API Documentation", "Technical reference for all modules"),
        ("D12", "System Architecture Document", "Updated architecture with Phase 2 components"),
    ]
    add_table_to_doc(doc, ["#", "Deliverable", "Description"], doc_deliverables)

    doc.add_heading('3.3 Process Deliverables', level=2)
    proc_deliverables = [
        ("D13", "Standard Operating Procedures", "Runbooks for common operational tasks"),
        ("D14", "Exception Handling Workflow", "Process for manual review cases"),
        ("D15", "Quality Control Checklist", "Validation procedures for output review"),
    ]
    add_table_to_doc(doc, ["#", "Deliverable", "Description"], proc_deliverables)

    # Section 5: Project Plan
    doc.add_page_break()
    doc.add_heading('5. Project Plan', level=1)

    doc.add_heading('5.2 Milestone Schedule', level=2)
    milestones = [
        ("M0", "Discovery Complete", "Discovery Phase (Section 2.0)"),
        ("M1", "Foundation Complete", "WP1"),
        ("M2", "Batch Processing Operational", "WP2"),
        ("M3", "Multi-Template Support", "WP3"),
        ("M4", "Historical Processing Ready", "WP4"),
        ("M5", "Exception Management Live", "WP5"),
        ("M6", "Monitoring & Reporting", "WP6"),
        ("M7", "Testing Complete", "WP7"),
        ("M8", "Documentation & Go-Live", "WP8"),
    ]
    add_table_to_doc(doc, ["Milestone", "Description", "Work Packages"], milestones)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Critical Gate: ").bold = True
    p.add_run("Development work packages (WP1-WP8) do not begin until Discovery (M0) is complete and findings are accepted.")

    # Section 10: Commercial Terms
    doc.add_page_break()
    doc.add_heading('10. Commercial Terms', level=1)

    doc.add_heading('10.2 Payment Schedule', level=2)
    payments = [
        ("Project Kickoff", "15%", "SOW Execution"),
        ("Discovery Complete", "10%", "Discovery Report Accepted"),
        ("M4 Complete (Historical Processing)", "25%", "Milestone Acceptance"),
        ("M7 Complete (Testing)", "25%", "UAT Sign-off"),
        ("M8 Complete (Go-Live)", "25%", "Final Acceptance"),
    ]
    add_table_to_doc(doc, ["Milestone", "Payment", "Trigger"], payments)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("Discovery Phase findings may result in scope adjustments. Any material scope changes will be documented and may affect subsequent milestone pricing.")

    # Section 12: Signatures
    doc.add_page_break()
    doc.add_heading('12. Signatures', level=1)
    doc.add_paragraph("This Statement of Work is agreed upon by:")

    doc.add_paragraph()
    doc.add_heading('Client:', level=3)
    for field in ["Name:", "Title:", "Signature:", "Date:"]:
        p = doc.add_paragraph()
        p.add_run(field).bold = True
        p.add_run(" " + "_" * 40)

    doc.add_paragraph()
    doc.add_heading('Consultant:', level=3)
    for field in ["Name:", "Title:", "Signature:", "Date:"]:
        p = doc.add_paragraph()
        p.add_run(field).bold = True
        p.add_run(" " + "_" * 40)

    # Save document
    doc.save('SOW_Phase2_TCO_Automation.docx')
    print("Created: SOW_Phase2_TCO_Automation.docx")

def create_discovery_questions_excel():
    """Create Discovery Questions Checklist as Excel workbook"""
    wb = openpyxl.Workbook()

    # Styles
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    section_fill = PatternFill(start_color="D6DCE4", end_color="D6DCE4", fill_type="solid")
    section_font = Font(bold=True, size=12)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    wrap_alignment = Alignment(wrap_text=True, vertical='top')

    # Remove default sheet
    wb.remove(wb.active)

    # Sheet 1: All Discovery Questions
    ws = wb.create_sheet("Discovery Questions")

    # Headers
    headers = ["#", "Category", "Question", "Priority", "Answer", "Notes", "Status"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')

    # Question data organized by category
    questions = [
        # D.1 Volume & Capacity
        ("D.1.1", "Volume & Capacity", "What is the current monthly volume of new vendor proposals?", "High"),
        ("D.1.2", "Volume & Capacity", "What is the expected volume growth over the next 12-24 months?", "Medium"),
        ("D.1.3", "Volume & Capacity", "Are there seasonal peaks (e.g., quarter-end, budget cycles)?", "Medium"),
        ("D.1.4", "Volume & Capacity", "What is the maximum volume during peak periods?", "High"),
        ("D.1.5", "Volume & Capacity", "What is the required turnaround time for new proposals?", "High"),
        ("D.1.6", "Volume & Capacity", "How many analysts will be using the system concurrently?", "Medium"),

        # D.2 Historical Proposals
        ("D.2.1", "Historical Proposals", "How many historical proposals exist in total?", "High"),
        ("D.2.2", "Historical Proposals", "What date range do the historical proposals cover?", "High"),
        ("D.2.3", "Historical Proposals", "Where are historical proposals currently stored? (File shares, SharePoint, local drives)", "High"),
        ("D.2.4", "Historical Proposals", "Are historical proposals organized by any structure? (By year, client, vendor)", "Medium"),
        ("D.2.5", "Historical Proposals", "What file formats are present? (Word versions, Excel versions, PDFs)", "High"),
        ("D.2.6", "Historical Proposals", "Are there any scanned/image-based documents requiring OCR?", "High"),
        ("D.2.7", "Historical Proposals", "Do historical files include the corresponding completed TCO outputs?", "Medium"),
        ("D.2.8", "Historical Proposals", "Are there any access restrictions or permissions issues?", "Medium"),
        ("D.2.9", "Historical Proposals", "What is the priority for processing historical data? (All at once, phased by date)", "Medium"),
        ("D.2.10", "Historical Proposals", "Are there any historical proposals that should be excluded?", "Low"),

        # D.3 TCO Templates
        ("D.3.1", "TCO Templates", "How many distinct TCO template variants are currently in use?", "High"),
        ("D.3.2", "TCO Templates", "What are the key differences between template variants?", "High"),
        ("D.3.3", "TCO Templates", "Are templates organized by term length? (5-year, 7-year, 10-year)", "High"),
        ("D.3.4", "TCO Templates", "Are there client-specific or deal-specific template customizations?", "Medium"),
        ("D.3.5", "TCO Templates", "Who owns/maintains the TCO templates?", "Medium"),
        ("D.3.6", "TCO Templates", "How often do templates change?", "Medium"),
        ("D.3.7", "TCO Templates", "Is there a master/golden template that others derive from?", "Medium"),
        ("D.3.8", "TCO Templates", "Do templates contain macros or VBA code?", "High"),
        ("D.3.9", "TCO Templates", "Are there formula dependencies that must be preserved?", "High"),
        ("D.3.10", "TCO Templates", "Can we obtain a copy of each template variant for analysis?", "High"),

        # D.4 Vendor & Document Formats
        ("D.4.1", "Vendor & Documents", "Are FIS and Jack Henry the only vendors, or are there others?", "High"),
        ("D.4.2", "Vendor & Documents", "If other vendors exist, what is the volume for each?", "High"),
        ("D.4.3", "Vendor & Documents", "Have FIS proposal formats changed over time? (Layout, sections, tables)", "High"),
        ("D.4.4", "Vendor & Documents", "Have Jack Henry proposal formats changed over time?", "High"),
        ("D.4.5", "Vendor & Documents", "Are there multiple proposal types per vendor? (Initial, renewal, amendment)", "Medium"),
        ("D.4.6", "Vendor & Documents", "Do proposals include attachments or supplementary files?", "Medium"),
        ("D.4.7", "Vendor & Documents", "Are there password-protected or encrypted documents?", "Medium"),
        ("D.4.8", "Vendor & Documents", "Do proposals contain handwritten annotations or comments?", "Low"),

        # D.5 Exception Handling
        ("D.5.1", "Exception Handling", "What confidence level triggers manual review? (Current: 90%)", "High"),
        ("D.5.2", "Exception Handling", "Who is responsible for reviewing exceptions?", "High"),
        ("D.5.3", "Exception Handling", "What is the expected exception rate?", "Medium"),
        ("D.5.4", "Exception Handling", "What is the acceptable turnaround for exception resolution?", "Medium"),
        ("D.5.5", "Exception Handling", "What fields are most critical for accuracy?", "High"),
        ("D.5.6", "Exception Handling", "Are there business rules that should trigger review? (e.g., large dollar variances)", "Medium"),
        ("D.5.7", "Exception Handling", "How should conflicts between automated and manual values be resolved?", "Medium"),
        ("D.5.8", "Exception Handling", "Should corrections feed back to improve future extractions?", "Medium"),

        # D.6 Operational & Process
        ("D.6.1", "Operations & Process", "What is the current manual process for creating TCOs?", "High"),
        ("D.6.2", "Operations & Process", "How long does manual TCO creation currently take?", "High"),
        ("D.6.3", "Operations & Process", "Who are the primary users of the system? (Titles, number of users)", "High"),
        ("D.6.4", "Operations & Process", "What business hours should the system be available?", "Medium"),
        ("D.6.5", "Operations & Process", "Are there SLAs for TCO delivery to clients?", "Medium"),
        ("D.6.6", "Operations & Process", "How are completed TCOs currently delivered/stored?", "Medium"),
        ("D.6.7", "Operations & Process", "Is there an existing approval workflow for TCOs?", "Medium"),
        ("D.6.8", "Operations & Process", "What happens when a proposal is revised/updated?", "Medium"),

        # D.7 Technical Infrastructure
        ("D.7.1", "Technical Infrastructure", "Where will the system be deployed? (On-premise server, cloud, user workstations)", "High"),
        ("D.7.2", "Technical Infrastructure", "What operating system is available? (Windows Server, Linux)", "High"),
        ("D.7.3", "Technical Infrastructure", "What database platform is preferred? (SQLite, PostgreSQL, SQL Server)", "Medium"),
        ("D.7.4", "Technical Infrastructure", "Are there firewall or network restrictions?", "Medium"),
        ("D.7.5", "Technical Infrastructure", "Is there an existing authentication system to integrate with? (AD, SSO)", "Medium"),
        ("D.7.6", "Technical Infrastructure", "What email system is available for notifications? (SMTP, Exchange, O365)", "Medium"),
        ("D.7.7", "Technical Infrastructure", "Is Slack or Teams available for notifications?", "Low"),
        ("D.7.8", "Technical Infrastructure", "What is the backup/disaster recovery strategy?", "Medium"),
        ("D.7.9", "Technical Infrastructure", "Who will provide IT support for the deployed system?", "Medium"),
        ("D.7.10", "Technical Infrastructure", "Are there any security or compliance requirements? (SOC2, data residency)", "High"),

        # D.8 Integration
        ("D.8.1", "Integration", "Where do new proposals currently come from? (Email, file drop, portal)", "High"),
        ("D.8.2", "Integration", "Should the system integrate with a CRM or deal management system?", "Medium"),
        ("D.8.3", "Integration", "Should completed TCOs be uploaded to a document management system?", "Medium"),
        ("D.8.4", "Integration", "Are there reporting or BI systems that should consume the data?", "Low"),
        ("D.8.5", "Integration", "Are there any APIs or systems we need to call or receive calls from?", "Medium"),

        # D.9 Reporting & Analytics
        ("D.9.1", "Reporting & Analytics", "What reports are needed for daily operations?", "High"),
        ("D.9.2", "Reporting & Analytics", "What management/executive reports are required?", "Medium"),
        ("D.9.3", "Reporting & Analytics", "What metrics are most important to track?", "High"),
        ("D.9.4", "Reporting & Analytics", "How should reports be delivered? (Dashboard, email, scheduled export)", "Medium"),
        ("D.9.5", "Reporting & Analytics", "Are there historical trend analysis requirements?", "Medium"),
        ("D.9.6", "Reporting & Analytics", "Should the system support ad-hoc queries or custom reports?", "Low"),

        # D.10 Training & Change Management
        ("D.10.1", "Training & Change", "Who needs training on the new system?", "High"),
        ("D.10.2", "Training & Change", "What is the preferred training format? (In-person, remote, self-service)", "Medium"),
        ("D.10.3", "Training & Change", "Is there a change management process for new tools?", "Medium"),
        ("D.10.4", "Training & Change", "Who will be the system administrator(s)?", "High"),
        ("D.10.5", "Training & Change", "What level of ongoing support is expected post-launch?", "High"),
    ]

    # Write questions
    for row, (num, category, question, priority) in enumerate(questions, 2):
        ws.cell(row=row, column=1, value=num).border = thin_border
        ws.cell(row=row, column=2, value=category).border = thin_border
        cell = ws.cell(row=row, column=3, value=question)
        cell.border = thin_border
        cell.alignment = wrap_alignment
        ws.cell(row=row, column=4, value=priority).border = thin_border
        ws.cell(row=row, column=5, value="").border = thin_border
        ws.cell(row=row, column=6, value="").border = thin_border
        ws.cell(row=row, column=7, value="Open").border = thin_border

        # Color priority cells
        priority_cell = ws.cell(row=row, column=4)
        if priority == "High":
            priority_cell.fill = PatternFill(start_color="F4B084", end_color="F4B084", fill_type="solid")
        elif priority == "Medium":
            priority_cell.fill = PatternFill(start_color="FFE699", end_color="FFE699", fill_type="solid")
        else:
            priority_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")

    # Set column widths
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 70
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 40
    ws.column_dimensions['F'].width = 30
    ws.column_dimensions['G'].width = 12

    # Freeze header row
    ws.freeze_panes = 'A2'

    # Sheet 2: Summary Dashboard
    ws2 = wb.create_sheet("Summary")

    ws2['A1'] = "Discovery Questions Summary"
    ws2['A1'].font = Font(bold=True, size=16)

    ws2['A3'] = "Category"
    ws2['B3'] = "Total"
    ws2['C3'] = "High Priority"
    ws2['D3'] = "Answered"
    ws2['E3'] = "% Complete"

    for col in range(1, 6):
        ws2.cell(row=3, column=col).fill = header_fill
        ws2.cell(row=3, column=col).font = header_font
        ws2.cell(row=3, column=col).border = thin_border

    categories = [
        ("Volume & Capacity", 6, 4),
        ("Historical Proposals", 10, 6),
        ("TCO Templates", 10, 7),
        ("Vendor & Documents", 8, 5),
        ("Exception Handling", 8, 3),
        ("Operations & Process", 8, 3),
        ("Technical Infrastructure", 10, 4),
        ("Integration", 5, 1),
        ("Reporting & Analytics", 6, 2),
        ("Training & Change", 5, 3),
    ]

    for row, (cat, total, high) in enumerate(categories, 4):
        ws2.cell(row=row, column=1, value=cat).border = thin_border
        ws2.cell(row=row, column=2, value=total).border = thin_border
        ws2.cell(row=row, column=3, value=high).border = thin_border
        ws2.cell(row=row, column=4, value=0).border = thin_border
        ws2.cell(row=row, column=5, value="0%").border = thin_border

    # Total row
    total_row = 4 + len(categories)
    ws2.cell(row=total_row, column=1, value="TOTAL").font = Font(bold=True)
    ws2.cell(row=total_row, column=2, value=f"=SUM(B4:B{total_row-1})")
    ws2.cell(row=total_row, column=3, value=f"=SUM(C4:C{total_row-1})")

    ws2.column_dimensions['A'].width = 25
    ws2.column_dimensions['B'].width = 10
    ws2.column_dimensions['C'].width = 15
    ws2.column_dimensions['D'].width = 12
    ws2.column_dimensions['E'].width = 12

    # Sheet 3: Template Inventory
    ws3 = wb.create_sheet("Template Inventory")

    template_headers = ["Template Name", "Term Length", "Usage/Purpose", "Owner", "Last Updated", "Complexity", "Macros?", "Notes"]
    for col, header in enumerate(template_headers, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    # Add empty rows for data entry
    for row in range(2, 12):
        for col in range(1, 9):
            ws3.cell(row=row, column=col).border = thin_border

    ws3.column_dimensions['A'].width = 25
    ws3.column_dimensions['B'].width = 15
    ws3.column_dimensions['C'].width = 30
    ws3.column_dimensions['D'].width = 15
    ws3.column_dimensions['E'].width = 15
    ws3.column_dimensions['F'].width = 12
    ws3.column_dimensions['G'].width = 10
    ws3.column_dimensions['H'].width = 40

    # Sheet 4: Historical File Inventory
    ws4 = wb.create_sheet("Historical Inventory")

    inv_headers = ["File Path", "File Name", "Vendor", "Client", "Date", "Format", "Size", "Compatible?", "Notes"]
    for col, header in enumerate(inv_headers, 1):
        cell = ws4.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    # Add empty rows
    for row in range(2, 22):
        for col in range(1, 10):
            ws4.cell(row=row, column=col).border = thin_border

    ws4.column_dimensions['A'].width = 40
    ws4.column_dimensions['B'].width = 30
    ws4.column_dimensions['C'].width = 12
    ws4.column_dimensions['D'].width = 20
    ws4.column_dimensions['E'].width = 12
    ws4.column_dimensions['F'].width = 12
    ws4.column_dimensions['G'].width = 10
    ws4.column_dimensions['H'].width = 12
    ws4.column_dimensions['I'].width = 40

    wb.save('Discovery_Questions_Checklist.xlsx')
    print("Created: Discovery_Questions_Checklist.xlsx")

def create_project_plan_excel():
    """Create Project Plan as Excel workbook"""
    wb = openpyxl.Workbook()

    # Styles
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    section_fill = PatternFill(start_color="D6DCE4", end_color="D6DCE4", fill_type="solid")
    milestone_fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    wrap_alignment = Alignment(wrap_text=True, vertical='top')

    # Remove default sheet
    wb.remove(wb.active)

    # Sheet 1: Milestones
    ws = wb.create_sheet("Milestones")

    headers = ["Milestone", "Description", "Work Package", "Deliverables", "Exit Criteria", "Payment %", "Status"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    milestones = [
        ("M0", "Discovery Complete", "Discovery Phase", "DD1-DD6", "Discovery exit criteria met", "10%", "Not Started"),
        ("M1", "Foundation Complete", "WP1", "Database, Queue, Config, Logging", "Infrastructure operational", "-", "Not Started"),
        ("M2", "Batch Processing Operational", "WP2", "D1, D2", "20 proposals processed successfully", "-", "Not Started"),
        ("M3", "Multi-Template Support", "WP3", "D3", "All templates configured", "-", "Not Started"),
        ("M4", "Historical Processing Ready", "WP4", "D4", "Historical import complete", "25%", "Not Started"),
        ("M5", "Exception Management Live", "WP5", "D5", "Review workflow operational", "-", "Not Started"),
        ("M6", "Monitoring & Reporting", "WP6", "D6, D7", "Dashboard and reports live", "-", "Not Started"),
        ("M7", "Testing Complete", "WP7", "Test Reports", "UAT sign-off", "25%", "Not Started"),
        ("M8", "Documentation & Go-Live", "WP8", "D8-D15", "Final acceptance", "25%", "Not Started"),
    ]

    for row, milestone in enumerate(milestones, 2):
        for col, value in enumerate(milestone, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.border = thin_border
            cell.alignment = wrap_alignment

    # Column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 35
    ws.column_dimensions['E'].width = 35
    ws.column_dimensions['F'].width = 12
    ws.column_dimensions['G'].width = 15

    # Sheet 2: Work Packages
    ws2 = wb.create_sheet("Work Packages")

    headers = ["WP", "Task ID", "Task Name", "Description", "Deliverable", "Dependencies", "Status"]
    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    tasks = [
        # WP1
        ("WP1", "1.1", "Database Schema Design", "Design data model for proposals, extractions, outputs", "Schema docs", "-", "Not Started"),
        ("WP1", "1.2", "Job Queue Implementation", "Implement job submission and tracking", "Queue module", "1.1", "Not Started"),
        ("WP1", "1.3", "Configuration Management", "Environment-based settings", "Config module", "-", "Not Started"),
        ("WP1", "1.4", "Logging Infrastructure", "Structured logging with rotation", "Logging module", "-", "Not Started"),
        # WP2
        ("WP2", "2.1", "Queue Management System", "Proposal ingestion and queue", "Queue manager", "WP1", "Not Started"),
        ("WP2", "2.2", "Parallel Worker Implementation", "Worker pool management", "Worker module", "2.1", "Not Started"),
        ("WP2", "2.3", "Error Handling & Retry", "Retry logic and dead letter queue", "Error handling", "2.2", "Not Started"),
        ("WP2", "2.4", "Progress Tracking", "Real-time progress updates", "Progress API", "2.2", "Not Started"),
        # WP3
        ("WP3", "3.1", "Template Storage & Versioning", "Template registry structure", "Registry module", "WP1", "Not Started"),
        ("WP3", "3.2", "Dynamic Mapping Engine", "Field-to-cell mapping", "Mapping engine", "3.1", "Not Started"),
        ("WP3", "3.3", "Template Detection Logic", "Auto-match templates", "Detection module", "3.2", "Not Started"),
        ("WP3", "3.4", "Configuration UI", "Template management interface", "Config UI", "3.3", "Not Started"),
        # WP4
        ("WP4", "4.1", "Bulk Import Utility", "Directory scanner and importer", "Import CLI", "WP1", "Not Started"),
        ("WP4", "4.2", "Document Classification", "Vendor and format detection", "Classifier", "4.1", "Not Started"),
        ("WP4", "4.3", "Legacy Format Handlers", "Support for older formats", "Legacy extractors", "4.2", "Not Started"),
        ("WP4", "4.4", "De-duplication Logic", "Identify and handle duplicates", "De-dup module", "4.2", "Not Started"),
        # WP5
        ("WP5", "5.1", "Review Queue Backend", "Exception queue data model", "Queue backend", "WP2, WP3", "Not Started"),
        ("WP5", "5.2", "Web-based Review UI", "Review interface", "Review UI", "5.1", "Not Started"),
        ("WP5", "5.3", "Approval Workflow", "Approval state machine", "Workflow engine", "5.2", "Not Started"),
        ("WP5", "5.4", "Feedback Loop Integration", "Learning from corrections", "Feedback module", "5.3", "Not Started"),
        # WP6
        ("WP6", "6.1", "Status Dashboard", "Real-time status display", "Dashboard", "WP2", "Not Started"),
        ("WP6", "6.2", "Notification System", "Email/Slack notifications", "Notifications", "WP2", "Not Started"),
        ("WP6", "6.3", "Analytics & Reporting", "Standard reports", "Reports module", "WP5", "Not Started"),
        ("WP6", "6.4", "Audit Trail Enhancement", "Extended audit logging", "Audit module", "WP5", "Not Started"),
        # WP7
        ("WP7", "7.1", "Unit Test Expansion", "80% code coverage", "Test suite", "WP1-WP6", "Not Started"),
        ("WP7", "7.2", "Integration Testing", "End-to-end tests", "Integration tests", "7.1", "Not Started"),
        ("WP7", "7.3", "Volume Testing", "20+ proposal stress test", "Volume report", "7.2", "Not Started"),
        ("WP7", "7.4", "User Acceptance Testing", "UAT execution", "UAT sign-off", "7.3", "Not Started"),
        # WP8
        ("WP8", "8.1", "Technical Documentation", "Architecture and API docs", "Tech docs", "WP7", "Not Started"),
        ("WP8", "8.2", "Operations Manual", "Daily operations procedures", "Ops manual", "WP7", "Not Started"),
        ("WP8", "8.3", "User Guides", "End-user documentation", "User guides", "WP7", "Not Started"),
        ("WP8", "8.4", "Knowledge Transfer", "Training sessions", "Training materials", "8.1-8.3", "Not Started"),
    ]

    for row, task in enumerate(tasks, 2):
        for col, value in enumerate(task, 1):
            cell = ws2.cell(row=row, column=col, value=value)
            cell.border = thin_border
            cell.alignment = wrap_alignment

    ws2.column_dimensions['A'].width = 8
    ws2.column_dimensions['B'].width = 10
    ws2.column_dimensions['C'].width = 30
    ws2.column_dimensions['D'].width = 45
    ws2.column_dimensions['E'].width = 20
    ws2.column_dimensions['F'].width = 15
    ws2.column_dimensions['G'].width = 12

    # Sheet 3: Risk Register
    ws3 = wb.create_sheet("Risk Register")

    headers = ["ID", "Risk Description", "Probability", "Impact", "Score", "Mitigation", "Owner", "Status"]
    for col, header in enumerate(headers, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    risks = [
        ("R1", "Historical documents in unknown formats", "Medium", "High", "High", "Early discovery phase with sample analysis", "Tech Lead", "Open"),
        ("R2", "TCO template variations more complex than expected", "Medium", "Medium", "Medium", "Flexible mapping engine with fallback options", "Developer", "Open"),
        ("R3", "Volume exceeds capacity", "Low", "Medium", "Low", "Scalable architecture with parallel processing", "Tech Lead", "Open"),
        ("R4", "Key resource unavailable", "Medium", "Medium", "Medium", "Cross-training; documentation", "PM", "Open"),
        ("R5", "Client SME availability", "Medium", "High", "High", "Early scheduling; async communication", "PM", "Open"),
        ("R6", "Integration dependencies", "Low", "Medium", "Low", "Modular design with clear APIs", "Tech Lead", "Open"),
        ("R7", "Performance bottlenecks", "Medium", "Medium", "Medium", "Early performance testing; optimization sprints", "Developer", "Open"),
        ("R8", "Scope creep", "Medium", "High", "High", "Change control process; clear SOW", "PM", "Open"),
    ]

    for row, risk in enumerate(risks, 2):
        for col, value in enumerate(risk, 1):
            cell = ws3.cell(row=row, column=col, value=value)
            cell.border = thin_border
            cell.alignment = wrap_alignment

    ws3.column_dimensions['A'].width = 6
    ws3.column_dimensions['B'].width = 45
    ws3.column_dimensions['C'].width = 12
    ws3.column_dimensions['D'].width = 10
    ws3.column_dimensions['E'].width = 10
    ws3.column_dimensions['F'].width = 45
    ws3.column_dimensions['G'].width = 12
    ws3.column_dimensions['H'].width = 10

    # Sheet 4: Acceptance Criteria
    ws4 = wb.create_sheet("Acceptance Criteria")

    headers = ["ID", "Type", "Criteria", "Target", "Validation Method", "Status"]
    for col, header in enumerate(headers, 1):
        cell = ws4.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    criteria = [
        ("AC1", "Functional", "System processes 20+ proposals per month", "-", "Volume test with real proposals", "Not Tested"),
        ("AC2", "Functional", "Historical proposals imported with >90% success rate", ">90%", "Batch processing report", "Not Tested"),
        ("AC3", "Functional", "All TCO template variants supported via configuration", "100%", "Template coverage test", "Not Tested"),
        ("AC4", "Functional", "Exception queue enables manual review with resolution tracking", "-", "Workflow walkthrough", "Not Tested"),
        ("AC5", "Functional", "Monitoring dashboard shows real-time status and alerts", "-", "UI demonstration", "Not Tested"),
        ("PA1", "Performance", "Single proposal processing time", "< 2 minutes", "Timing test", "Not Tested"),
        ("PA2", "Performance", "Batch of 20 proposals", "< 30 minutes", "Batch timing test", "Not Tested"),
        ("PA3", "Performance", "Dashboard response time", "< 3 seconds", "UI response test", "Not Tested"),
        ("PA4", "Performance", "System availability", "99%", "Uptime monitoring", "Not Tested"),
        ("QA1", "Quality", "Extraction accuracy (auto-accepted)", "> 95%", "Accuracy report", "Not Tested"),
        ("QA2", "Quality", "Overall processing success rate", "> 98%", "Processing report", "Not Tested"),
        ("QA3", "Quality", "Zero data loss or corruption", "100%", "Data integrity check", "Not Tested"),
        ("QA4", "Quality", "Complete audit trail for all outputs", "100%", "Audit log review", "Not Tested"),
    ]

    for row, crit in enumerate(criteria, 2):
        for col, value in enumerate(crit, 1):
            cell = ws4.cell(row=row, column=col, value=value)
            cell.border = thin_border
            cell.alignment = wrap_alignment

    ws4.column_dimensions['A'].width = 8
    ws4.column_dimensions['B'].width = 12
    ws4.column_dimensions['C'].width = 55
    ws4.column_dimensions['D'].width = 15
    ws4.column_dimensions['E'].width = 30
    ws4.column_dimensions['F'].width = 12

    wb.save('Project_Plan_Phase2.xlsx')
    print("Created: Project_Plan_Phase2.xlsx")

if __name__ == "__main__":
    print("Generating SOW and Project Plan documents...")
    print("-" * 50)

    create_sow_word_document()
    create_discovery_questions_excel()
    create_project_plan_excel()

    print("-" * 50)
    print("All documents created successfully!")
