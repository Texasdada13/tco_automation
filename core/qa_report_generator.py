"""
QA Report Generator Module

Generates comprehensive, well-formatted Quality Assurance reports
in Word document format for every extraction.

Reports include:
- Executive Summary
- Extraction Statistics
- Confidence Analysis
- Data Quality Issues
- Items Requiring Review
- Recommendations
"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class QAMetrics:
    """Container for QA metrics."""
    total_items: int = 0
    high_confidence_count: int = 0  # >= 0.90
    medium_confidence_count: int = 0  # 0.70 - 0.90
    low_confidence_count: int = 0  # < 0.70
    average_confidence: float = 0.0

    # Validation results
    items_auto_approved: int = 0  # Bucket 1
    items_need_review: int = 0  # Bucket 2

    # Issues
    missing_fields: List[Dict] = field(default_factory=list)
    low_confidence_items: List[Dict] = field(default_factory=list)
    validation_errors: List[Dict] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    # Categories
    category_distribution: Dict[str, int] = field(default_factory=dict)
    fee_type_distribution: Dict[str, int] = field(default_factory=dict)

    # Financial summary
    total_monthly: float = 0.0
    total_one_time: float = 0.0
    total_annual: float = 0.0


class QAReportGenerator:
    """
    Generates comprehensive QA reports in Word format.

    Usage:
        generator = QAReportGenerator()
        metrics = generator.analyze_extraction(extraction_data)
        generator.generate_word_report(metrics, output_path)
    """

    # Confidence thresholds
    HIGH_CONFIDENCE = 0.90
    MEDIUM_CONFIDENCE = 0.70

    def __init__(self):
        """Initialize the QA report generator."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, Word reports will not be generated")

    def analyze_extraction(
        self,
        extraction_data: Dict[str, Any],
        vendor: str = "Unknown"
    ) -> QAMetrics:
        """
        Analyze extraction data and compute QA metrics.

        Args:
            extraction_data: The AI-enhanced extraction JSON
            vendor: Vendor name

        Returns:
            QAMetrics object with computed metrics
        """
        metrics = QAMetrics()

        line_items = extraction_data.get('line_items', [])
        metrics.total_items = len(line_items)

        if not line_items:
            metrics.warnings.append("No line items extracted")
            return metrics

        confidences = []

        for item in line_items:
            confidence = item.get('overall_confidence', 0.85)
            confidences.append(confidence)

            # Confidence buckets
            if confidence >= self.HIGH_CONFIDENCE:
                metrics.high_confidence_count += 1
                metrics.items_auto_approved += 1
            elif confidence >= self.MEDIUM_CONFIDENCE:
                metrics.medium_confidence_count += 1
                metrics.items_need_review += 1
                metrics.low_confidence_items.append({
                    'solution_name': item.get('solution_name', 'Unknown'),
                    'confidence': confidence,
                    'reason': 'Medium confidence - needs verification'
                })
            else:
                metrics.low_confidence_count += 1
                metrics.items_need_review += 1
                metrics.low_confidence_items.append({
                    'solution_name': item.get('solution_name', 'Unknown'),
                    'confidence': confidence,
                    'reason': 'Low confidence - manual review required'
                })

            # Check for missing fields
            required_fields = ['solution_name', 'fee_type', 'category']
            for field in required_fields:
                if not item.get(field):
                    metrics.missing_fields.append({
                        'item': item.get('solution_name', 'Unknown'),
                        'field': field
                    })

            # Category distribution
            category = item.get('category', 'Other')
            metrics.category_distribution[category] = metrics.category_distribution.get(category, 0) + 1

            # Fee type distribution
            fee_type = item.get('fee_type', 'Unknown')
            metrics.fee_type_distribution[fee_type] = metrics.fee_type_distribution.get(fee_type, 0) + 1

            # Financial totals
            monthly = item.get('monthly_fee', 0) or item.get('per_unit_rate', 0)
            one_time = item.get('one_time_fee', 0)

            if fee_type in ['Monthly F', 'Monthly V']:
                metrics.total_monthly += monthly
            elif fee_type == 'Annual':
                metrics.total_annual += monthly
            elif fee_type == 'One-Time':
                metrics.total_one_time += one_time or monthly

        # Calculate average confidence
        metrics.average_confidence = sum(confidences) / len(confidences) if confidences else 0

        # Validation checks
        self._run_validation_checks(metrics, line_items, extraction_data)

        return metrics

    def _run_validation_checks(
        self,
        metrics: QAMetrics,
        line_items: List[Dict],
        extraction_data: Dict
    ) -> None:
        """Run validation checks and add errors/warnings."""

        # Check for duplicate items
        seen_names = {}
        for item in line_items:
            name = item.get('solution_name', '')
            fee_type = item.get('fee_type', '')
            key = f"{name}_{fee_type}"

            if key in seen_names:
                metrics.validation_errors.append({
                    'type': 'duplicate',
                    'item': name,
                    'message': f"Duplicate entry found: {name} ({fee_type})"
                })
            seen_names[key] = True

        # Check for zero amounts on required items
        for item in line_items:
            if not item.get('optional', False):
                amount = (item.get('monthly_fee', 0) or
                         item.get('one_time_fee', 0) or
                         item.get('per_unit_rate', 0))
                if amount == 0:
                    metrics.warnings.append(
                        f"Zero amount for required item: {item.get('solution_name', 'Unknown')}"
                    )

        # Check summary totals if available
        summary = extraction_data.get('summary', {})
        stated_monthly = summary.get('total_monthly_required', 0)
        if stated_monthly > 0:
            calculated_monthly = metrics.total_monthly
            diff_pct = abs(stated_monthly - calculated_monthly) / stated_monthly * 100
            if diff_pct > 5:  # More than 5% difference
                metrics.validation_errors.append({
                    'type': 'sum_mismatch',
                    'item': 'Monthly Total',
                    'message': f"Sum mismatch: stated ${stated_monthly:,.2f} vs calculated ${calculated_monthly:,.2f} ({diff_pct:.1f}% diff)"
                })

    def generate_word_report(
        self,
        metrics: QAMetrics,
        output_path: str,
        extraction_data: Dict[str, Any] = None,
        vendor: str = "Unknown",
        client: str = "Unknown",
        source_file: str = ""
    ) -> str:
        """
        Generate a comprehensive Word document QA report.

        Args:
            metrics: QAMetrics object
            output_path: Path to save the Word document
            extraction_data: Original extraction data (optional)
            vendor: Vendor name
            client: Client name
            source_file: Source file name

        Returns:
            Path to generated report
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available, cannot generate Word report")
            return ""

        doc = Document()

        # Set up document styles
        self._setup_styles(doc)

        # === TITLE ===
        title = doc.add_heading('TCO Extraction Quality Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
        meta_para.add_run(f"Vendor: {vendor} | Client: {client}").italic = True

        if source_file:
            doc.add_paragraph(f"Source: {source_file}").alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # === EXECUTIVE SUMMARY ===
        doc.add_heading('Executive Summary', level=1)

        # Overall status
        status_para = doc.add_paragraph()

        if metrics.average_confidence >= 0.90 and not metrics.validation_errors:
            status_run = status_para.add_run('PASSED')
            status_run.bold = True
            status_run.font.color.rgb = RGBColor(0, 128, 0)
            status_para.add_run(' - Extraction completed with high confidence. Ready for processing.')
        elif metrics.average_confidence >= 0.70:
            status_run = status_para.add_run('REVIEW RECOMMENDED')
            status_run.bold = True
            status_run.font.color.rgb = RGBColor(255, 165, 0)
            status_para.add_run(f' - {metrics.items_need_review} items require manual verification.')
        else:
            status_run = status_para.add_run('MANUAL REVIEW REQUIRED')
            status_run.bold = True
            status_run.font.color.rgb = RGBColor(255, 0, 0)
            status_para.add_run(' - Low confidence extraction. Please verify all items.')

        # Key metrics table
        doc.add_paragraph()
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'

        summary_data = [
            ('Total Line Items Extracted', str(metrics.total_items)),
            ('Average Confidence Score', f"{metrics.average_confidence:.1%}"),
            ('Auto-Approved (Bucket 1)', f"{metrics.items_auto_approved} items ({metrics.items_auto_approved/max(metrics.total_items,1)*100:.0f}%)"),
            ('Needs Review (Bucket 2)', f"{metrics.items_need_review} items ({metrics.items_need_review/max(metrics.total_items,1)*100:.0f}%)"),
            ('Validation Errors', str(len(metrics.validation_errors))),
            ('Warnings', str(len(metrics.warnings)))
        ]

        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True

        # === CONFIDENCE ANALYSIS ===
        doc.add_heading('Confidence Analysis', level=1)

        conf_para = doc.add_paragraph()
        conf_para.add_run('Confidence Distribution:\n').bold = True
        conf_para.add_run(f"  High (≥90%): {metrics.high_confidence_count} items\n")
        conf_para.add_run(f"  Medium (70-90%): {metrics.medium_confidence_count} items\n")
        conf_para.add_run(f"  Low (<70%): {metrics.low_confidence_count} items")

        # Confidence visual bar
        self._add_confidence_bar(doc, metrics)

        # === FINANCIAL SUMMARY ===
        doc.add_heading('Financial Summary', level=1)

        fin_table = doc.add_table(rows=4, cols=2)
        fin_table.style = 'Table Grid'

        fin_data = [
            ('Total Monthly Fees', f"${metrics.total_monthly:,.2f}"),
            ('Total Annual Fees', f"${metrics.total_annual:,.2f}"),
            ('Total One-Time Fees', f"${metrics.total_one_time:,.2f}"),
            ('Estimated Year 1 Total', f"${(metrics.total_monthly * 12 + metrics.total_annual + metrics.total_one_time):,.2f}")
        ]

        for i, (label, value) in enumerate(fin_data):
            row = fin_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True

        # === CATEGORY DISTRIBUTION ===
        doc.add_heading('Category Distribution', level=1)

        if metrics.category_distribution:
            cat_table = doc.add_table(rows=len(metrics.category_distribution) + 1, cols=2)
            cat_table.style = 'Table Grid'

            # Header
            cat_table.rows[0].cells[0].text = 'Category'
            cat_table.rows[0].cells[1].text = 'Count'
            for cell in cat_table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True

            for i, (cat, count) in enumerate(sorted(metrics.category_distribution.items()), 1):
                cat_table.rows[i].cells[0].text = cat
                cat_table.rows[i].cells[1].text = str(count)

        # === ITEMS REQUIRING REVIEW ===
        if metrics.low_confidence_items:
            doc.add_heading('Items Requiring Review', level=1)

            doc.add_paragraph(
                f"The following {len(metrics.low_confidence_items)} items have confidence scores "
                "below the auto-approval threshold and require manual verification:"
            )

            review_table = doc.add_table(rows=len(metrics.low_confidence_items) + 1, cols=3)
            review_table.style = 'Table Grid'

            # Header
            headers = ['Solution Name', 'Confidence', 'Reason']
            for j, header in enumerate(headers):
                review_table.rows[0].cells[j].text = header
                review_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

            for i, item in enumerate(metrics.low_confidence_items, 1):
                review_table.rows[i].cells[0].text = item['solution_name']
                review_table.rows[i].cells[1].text = f"{item['confidence']:.1%}"
                review_table.rows[i].cells[2].text = item['reason']

        # === VALIDATION ERRORS ===
        if metrics.validation_errors:
            doc.add_heading('Validation Errors', level=1)

            for error in metrics.validation_errors:
                error_para = doc.add_paragraph(style='List Bullet')
                error_run = error_para.add_run(f"[{error['type'].upper()}] ")
                error_run.bold = True
                error_run.font.color.rgb = RGBColor(255, 0, 0)
                error_para.add_run(error['message'])

        # === WARNINGS ===
        if metrics.warnings:
            doc.add_heading('Warnings', level=1)

            for warning in metrics.warnings:
                warn_para = doc.add_paragraph(style='List Bullet')
                warn_run = warn_para.add_run('Warning: ')
                warn_run.bold = True
                warn_run.font.color.rgb = RGBColor(255, 165, 0)
                warn_para.add_run(warning)

        # === RECOMMENDATIONS ===
        doc.add_heading('Recommendations', level=1)

        recommendations = self._generate_recommendations(metrics)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

        # === AUDIT INFORMATION ===
        doc.add_heading('Audit Information', level=1)

        audit_para = doc.add_paragraph()
        audit_para.add_run('Report Generated: ').bold = True
        audit_para.add_run(f"{datetime.now().isoformat()}\n")
        audit_para.add_run('Extraction Method: ').bold = True
        audit_para.add_run('AI-Enhanced (Claude)\n')
        audit_para.add_run('QA Validation: ').bold = True
        audit_para.add_run('4-Layer Validation System')

        # Save document
        doc.save(output_path)
        logger.info(f"QA Report saved to: {output_path}")

        return output_path

    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles."""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_confidence_bar(self, doc: Document, metrics: QAMetrics) -> None:
        """Add a visual confidence bar."""
        total = max(metrics.total_items, 1)
        high_pct = metrics.high_confidence_count / total * 100
        med_pct = metrics.medium_confidence_count / total * 100
        low_pct = metrics.low_confidence_count / total * 100

        bar_para = doc.add_paragraph()
        bar_para.add_run(f"[{'█' * int(high_pct/5)}").font.color.rgb = RGBColor(0, 128, 0)
        bar_para.add_run(f"{'█' * int(med_pct/5)}").font.color.rgb = RGBColor(255, 165, 0)
        bar_para.add_run(f"{'█' * int(low_pct/5)}]").font.color.rgb = RGBColor(255, 0, 0)
        bar_para.add_run(f" High: {high_pct:.0f}% | Medium: {med_pct:.0f}% | Low: {low_pct:.0f}%")

    def _generate_recommendations(self, metrics: QAMetrics) -> List[str]:
        """Generate recommendations based on metrics."""
        recommendations = []

        if metrics.average_confidence >= 0.95:
            recommendations.append("Extraction quality is excellent. Proceed with confidence.")
        elif metrics.average_confidence >= 0.85:
            recommendations.append("Extraction quality is good. Spot-check flagged items before finalizing.")
        elif metrics.average_confidence >= 0.70:
            recommendations.append("Review all medium and low confidence items before processing.")
        else:
            recommendations.append("Consider re-extracting with clearer document or manual entry.")

        if metrics.items_need_review > 0:
            recommendations.append(
                f"Review the {metrics.items_need_review} flagged items in the 'Items Requiring Review' section."
            )

        if metrics.validation_errors:
            recommendations.append(
                f"Address the {len(metrics.validation_errors)} validation error(s) before finalizing."
            )

        if metrics.missing_fields:
            recommendations.append(
                f"Complete the {len(metrics.missing_fields)} missing required field(s)."
            )

        if not recommendations:
            recommendations.append("No specific recommendations. Data quality appears acceptable.")

        return recommendations

    def generate_json_report(
        self,
        metrics: QAMetrics,
        output_path: str,
        extraction_data: Dict[str, Any] = None,
        vendor: str = "Unknown",
        client: str = "Unknown"
    ) -> str:
        """
        Generate a JSON format QA report.

        Args:
            metrics: QAMetrics object
            output_path: Path to save the JSON file
            extraction_data: Original extraction data
            vendor: Vendor name
            client: Client name

        Returns:
            Path to generated report
        """
        report = {
            'report_metadata': {
                'generated_at': datetime.now().isoformat(),
                'vendor': vendor,
                'client': client,
                'report_type': 'qa_validation'
            },
            'summary': {
                'total_items': metrics.total_items,
                'average_confidence': metrics.average_confidence,
                'items_auto_approved': metrics.items_auto_approved,
                'items_need_review': metrics.items_need_review,
                'validation_passed': len(metrics.validation_errors) == 0 and metrics.average_confidence >= 0.70
            },
            'confidence_distribution': {
                'high': metrics.high_confidence_count,
                'medium': metrics.medium_confidence_count,
                'low': metrics.low_confidence_count
            },
            'financial_summary': {
                'total_monthly': metrics.total_monthly,
                'total_annual': metrics.total_annual,
                'total_one_time': metrics.total_one_time,
                'estimated_year_1': metrics.total_monthly * 12 + metrics.total_annual + metrics.total_one_time
            },
            'category_distribution': metrics.category_distribution,
            'fee_type_distribution': metrics.fee_type_distribution,
            'issues': {
                'validation_errors': metrics.validation_errors,
                'warnings': metrics.warnings,
                'missing_fields': metrics.missing_fields,
                'low_confidence_items': metrics.low_confidence_items
            },
            'recommendations': self._generate_recommendations(metrics)
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, default=str)

        logger.info(f"QA JSON Report saved to: {output_path}")
        return output_path


def generate_qa_reports(
    extraction_data: Dict[str, Any],
    json_output_path: str,
    word_output_path: str,
    vendor: str = "Unknown",
    client: str = "Unknown",
    source_file: str = ""
) -> Dict[str, str]:
    """
    Convenience function to generate both JSON and Word QA reports.

    Args:
        extraction_data: AI-enhanced extraction data
        json_output_path: Path for JSON report
        word_output_path: Path for Word report
        vendor: Vendor name
        client: Client name
        source_file: Source file name

    Returns:
        Dict with paths to generated reports
    """
    generator = QAReportGenerator()
    metrics = generator.analyze_extraction(extraction_data, vendor)

    paths = {}

    # Generate JSON report
    paths['json'] = generator.generate_json_report(
        metrics, json_output_path, extraction_data, vendor, client
    )

    # Generate Word report if possible
    if DOCX_AVAILABLE:
        paths['word'] = generator.generate_word_report(
            metrics, word_output_path, extraction_data, vendor, client, source_file
        )

    return paths
