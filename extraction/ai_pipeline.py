"""
AI Extraction Pipeline

Main orchestrator that integrates all extraction components:
- IntelligentExtractor: 4-stage AI extraction
- QualityAssurance: Validation and confidence scoring
- BucketRouter: Auto-populate vs manual review routing
- ReviewReporter: Word document generation for reviews
- VendorCache: Context caching for optimization

Usage:
    from tco_automation.extraction import AIPipeline

    pipeline = AIPipeline(api_key="your-anthropic-key")
    result = pipeline.process_document(
        document_path="proposal.docx",
        output_excel="tco_output.xlsx"
    )
"""

import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field

from .intelligent_extractor import IntelligentExtractor, ExtractionResult
from .quality_assurance import QualityAssurance, QAResult, QABucket
from .bucket_router import BucketRouter, BucketResult
from .review_reporter import ReviewReporter
from .vendor_cache import VendorCache

from ..config import (
    CONTRACT_DEFAULTS,
    FALLBACK_CONFIG,
    OUTPUT_CONFIG,
    LOGGING_CONFIG
)

logger = logging.getLogger(__name__)


@dataclass
class PipelineResult:
    """Complete result from AI extraction pipeline."""
    success: bool
    extraction_result: ExtractionResult = None
    qa_result: QAResult = None
    bucket_result: BucketResult = None
    excel_path: Optional[str] = None
    review_report_path: Optional[str] = None
    audit_json_path: Optional[str] = None
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    processing_time_seconds: float = 0.0
    summary: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "excel_path": self.excel_path,
            "review_report_path": self.review_report_path,
            "audit_json_path": self.audit_json_path,
            "errors": self.errors,
            "warnings": self.warnings,
            "processing_time_seconds": self.processing_time_seconds,
            "summary": self.summary
        }


class AIPipeline:
    """
    AI-powered TCO extraction pipeline.

    Orchestrates the complete extraction workflow:
    1. Load and parse document
    2. Extract pricing data using Claude AI
    3. Run quality assurance checks
    4. Route to appropriate output buckets
    5. Generate Excel output and review reports
    """

    def __init__(
        self,
        api_key: Optional[str] = None,
        contract_term: Optional[int] = None,
        annual_increase: Optional[float] = None,
        enable_cache: bool = True,
        enable_fallback: bool = True
    ):
        """
        Initialize the AI pipeline.

        Args:
            api_key: Anthropic API key (or set ANTHROPIC_API_KEY env var)
            contract_term: Default contract term in years
            annual_increase: Default annual price increase rate
            enable_cache: Enable vendor context caching
            enable_fallback: Enable rule-based fallback if AI fails
        """
        # Initialize components
        self.extractor = IntelligentExtractor(api_key=api_key)
        self.qa = QualityAssurance()
        self.router = BucketRouter()
        self.reporter = ReviewReporter()
        self.cache = VendorCache() if enable_cache else None

        # Configuration
        self.contract_term = contract_term or CONTRACT_DEFAULTS['term_years']
        self.annual_increase = annual_increase or CONTRACT_DEFAULTS['annual_increase_rate']
        self.growth_rate = CONTRACT_DEFAULTS['volume_growth_rate']
        self.enable_fallback = enable_fallback and FALLBACK_CONFIG.get('enable_rule_based_fallback', True)

        logger.info("AIPipeline initialized")

    def process_document(
        self,
        document_path: str,
        output_excel: Optional[str] = None,
        output_review: Optional[str] = None,
        contract_term: Optional[int] = None,
        annual_increase: Optional[float] = None,
        save_audit: bool = True
    ) -> PipelineResult:
        """
        Process a vendor proposal document through the complete pipeline.

        Args:
            document_path: Path to the proposal document
            output_excel: Output Excel path (optional)
            output_review: Output review report path (optional)
            contract_term: Contract term override
            annual_increase: Annual increase override
            save_audit: Save audit JSON file

        Returns:
            PipelineResult with all outputs
        """
        start_time = datetime.now()
        result = PipelineResult(success=False)

        try:
            # Step 1: Load document
            logger.info(f"Processing document: {document_path}")
            document_text, tables = self._load_document(document_path)

            if not document_text:
                result.errors.append(f"Could not load document: {document_path}")
                return result

            # Step 2: Check cache for vendor context
            vendor_context = None
            if self.cache:
                # Try to identify vendor from document
                vendor_hint = self._detect_vendor_hint(document_text)
                if vendor_hint:
                    vendor_context = self.cache.get_vendor_context(vendor_hint)
                    if vendor_context:
                        logger.info(f"Using cached context for vendor: {vendor_hint}")

            # Step 3: Run AI extraction
            extraction_result = self.extractor.extract(
                document_text=document_text,
                tables=tables,
                contract_term=contract_term or self.contract_term,
                annual_increase=annual_increase or self.annual_increase,
                growth_rate=self.growth_rate,
                vendor_context=vendor_context
            )
            result.extraction_result = extraction_result

            if extraction_result.errors:
                result.errors.extend(extraction_result.errors)

            if not extraction_result.line_items:
                # Try fallback if enabled
                if self.enable_fallback:
                    logger.warning("AI extraction returned no items, attempting fallback")
                    extraction_result = self._try_fallback_extraction(
                        document_path, document_text
                    )
                    result.extraction_result = extraction_result

                if not extraction_result.line_items:
                    result.errors.append("No line items extracted from document")
                    return result

            # Step 4: Run QA validation
            stated_totals = extraction_result.context.get('stated_totals', {})
            qa_result = self.qa.validate(
                extraction_result.line_items,
                stated_totals,
                extraction_result.calculations
            )
            result.qa_result = qa_result

            # Step 5: Route to buckets
            bucket_result = self.router.route(
                extraction_result.line_items,
                qa_result
            )
            result.bucket_result = bucket_result

            # Step 6: Cache vendor context if successful
            if self.cache and extraction_result.overall_confidence >= 0.70:
                vendor_name = extraction_result.context.get('vendor_name', 'Unknown')
                self.cache.cache_vendor_context(
                    vendor_name,
                    extraction_result.context,
                    extraction_result.overall_confidence
                )

            # Step 7: Generate outputs
            if OUTPUT_CONFIG.get('excel_output_enabled', True) and bucket_result.bucket1_items:
                excel_path = self._populate_excel(
                    bucket_result.bucket1_items,
                    extraction_result.calculations,
                    output_excel
                )
                result.excel_path = excel_path

            if OUTPUT_CONFIG.get('review_report_enabled', True) and bucket_result.bucket2_items:
                review_path = self.reporter.generate_report(
                    bucket_result,
                    qa_result,
                    extraction_result.context,
                    output_review
                )
                result.review_report_path = review_path

            if save_audit and OUTPUT_CONFIG.get('save_extraction_json', True):
                audit_path = self._save_audit_trail(
                    document_path,
                    extraction_result,
                    qa_result,
                    bucket_result
                )
                result.audit_json_path = audit_path

            # Success
            result.success = True
            result.warnings = extraction_result.warnings

        except Exception as e:
            logger.error(f"Pipeline error: {str(e)}")
            result.errors.append(f"Pipeline error: {str(e)}")

        # Calculate processing time
        result.processing_time_seconds = (datetime.now() - start_time).total_seconds()

        # Generate summary
        result.summary = self._generate_summary(result)

        logger.info(f"Pipeline complete in {result.processing_time_seconds:.1f}s")
        return result

    def process_text(
        self,
        document_text: str,
        tables: Optional[List[Dict]] = None,
        vendor_hint: Optional[str] = None,
        contract_term: Optional[int] = None,
        annual_increase: Optional[float] = None
    ) -> PipelineResult:
        """
        Process raw text through the pipeline (for testing or API use).

        Args:
            document_text: Raw document text
            tables: Optional extracted tables
            vendor_hint: Optional vendor name hint
            contract_term: Contract term override
            annual_increase: Annual increase override

        Returns:
            PipelineResult
        """
        start_time = datetime.now()
        result = PipelineResult(success=False)

        try:
            # Check cache
            vendor_context = None
            if self.cache and vendor_hint:
                vendor_context = self.cache.get_vendor_context(vendor_hint)

            # Run extraction
            extraction_result = self.extractor.extract(
                document_text=document_text,
                tables=tables,
                contract_term=contract_term or self.contract_term,
                annual_increase=annual_increase or self.annual_increase,
                vendor_context=vendor_context
            )
            result.extraction_result = extraction_result

            if not extraction_result.line_items:
                result.errors.append("No line items extracted")
                return result

            # Run QA
            qa_result = self.qa.validate(extraction_result.line_items)
            result.qa_result = qa_result

            # Route
            bucket_result = self.router.route(extraction_result.line_items, qa_result)
            result.bucket_result = bucket_result

            result.success = True

        except Exception as e:
            result.errors.append(str(e))

        result.processing_time_seconds = (datetime.now() - start_time).total_seconds()
        result.summary = self._generate_summary(result)

        return result

    def _load_document(
        self,
        document_path: str
    ) -> Tuple[str, List[Dict]]:
        """Load and parse a document."""
        path = Path(document_path)

        if not path.exists():
            logger.error(f"Document not found: {document_path}")
            return "", []

        # Try to use existing document loaders
        try:
            from ..loaders.document_loader import DocumentLoader
            loader = DocumentLoader()
            doc_data = loader.load(document_path)
            return doc_data.get('text', ''), doc_data.get('tables', [])
        except ImportError:
            pass

        # Basic text extraction fallback
        try:
            if path.suffix.lower() == '.txt':
                with open(path, 'r', encoding='utf-8') as f:
                    return f.read(), []

            elif path.suffix.lower() == '.docx':
                from docx import Document
                doc = Document(path)
                text_parts = []
                tables = []

                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                # Extract tables with full data
                for i, table in enumerate(doc.tables):
                    table_data = {
                        'index': i,
                        'headers': [],
                        'rows': []
                    }
                    for row_idx, row in enumerate(table.rows):
                        row_data = [cell.text.strip() for cell in row.cells]
                        if row_idx == 0:
                            table_data['headers'] = row_data
                        else:
                            table_data['rows'].append(row_data)

                        # Also add to text for AI context
                        text_parts.append(' | '.join(row_data))

                    tables.append(table_data)

                return '\n'.join(text_parts), tables

            elif path.suffix.lower() in ['.xlsx', '.xls']:
                import pandas as pd
                dfs = pd.read_excel(path, sheet_name=None)
                tables = []
                text_parts = []
                for sheet_name, df in dfs.items():
                    text_parts.append(f"Sheet: {sheet_name}")
                    text_parts.append(df.to_string())
                    tables.append({
                        'sheet': sheet_name,
                        'headers': df.columns.tolist(),
                        'rows': df.values.tolist()
                    })
                return '\n\n'.join(text_parts), tables

        except Exception as e:
            logger.error(f"Error loading document: {e}")

        return "", []

    def _detect_vendor_hint(self, text: str) -> Optional[str]:
        """Try to detect vendor from document text."""
        text_lower = text.lower()

        vendor_keywords = {
            'FIS': ['fis ', 'fis,', 'horizon', 'digital one', 'payments one'],
            'Jack Henry': ['jack henry', 'jha', 'silverlake', 'xperience', 'banno'],
            'Fiserv': ['fiserv', 'dna', 'precision', 'premier'],
            'Finastra': ['finastra', 'fusion'],
        }

        for vendor, keywords in vendor_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return vendor

        return None

    def _try_fallback_extraction(
        self,
        document_path: str,
        document_text: str
    ) -> ExtractionResult:
        """Try rule-based fallback extraction."""
        result = ExtractionResult()

        try:
            # Detect vendor for appropriate extractor
            vendor = self._detect_vendor_hint(document_text)

            if vendor == 'FIS':
                from ..extractors.fis_extractor import FISExtractor
                extractor = FISExtractor()
                items = extractor.extract(document_path)
                result.line_items = items
                result.warnings.append("Used FIS rule-based fallback")

            elif vendor == 'Jack Henry':
                from ..extractors.jh_extractor import JHExtractor
                extractor = JHExtractor()
                items = extractor.extract(document_path)
                result.line_items = items
                result.warnings.append("Used Jack Henry rule-based fallback")

            # Add confidence to fallback items
            for item in result.line_items:
                if 'overall_confidence' not in item:
                    item['overall_confidence'] = 0.85  # Default confidence for rule-based

        except Exception as e:
            logger.warning(f"Fallback extraction failed: {e}")
            result.errors.append(f"Fallback failed: {e}")

        return result

    def _populate_excel(
        self,
        line_items: List[Dict],
        calculations: Dict,
        output_path: Optional[str]
    ) -> Optional[str]:
        """Populate Excel template with extracted data."""
        try:
            from ..exporters.excel_populator import ExcelPopulator

            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = f"tco_output_{timestamp}.xlsx"

            populator = ExcelPopulator()
            populator.populate(line_items, calculations, output_path)

            return output_path

        except ImportError:
            logger.warning("ExcelPopulator not available")
            return None
        except Exception as e:
            logger.error(f"Excel population failed: {e}")
            return None

    def _save_audit_trail(
        self,
        document_path: str,
        extraction_result: ExtractionResult,
        qa_result: QAResult,
        bucket_result: BucketResult
    ) -> str:
        """Save complete audit trail as JSON."""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        audit_path = f"audit_{timestamp}.json"

        audit_data = {
            "timestamp": datetime.now().isoformat(),
            "document": document_path,
            "extraction": {
                "line_items_count": len(extraction_result.line_items),
                "overall_confidence": extraction_result.overall_confidence,
                "context": extraction_result.context,
                "line_items": extraction_result.line_items,
                "calculations": extraction_result.calculations
            },
            "qa": qa_result.to_dict(),
            "routing": bucket_result.to_dict()
        }

        with open(audit_path, 'w', encoding='utf-8') as f:
            json.dump(audit_data, f, indent=2, default=str)

        return audit_path

    def _generate_summary(self, result: PipelineResult) -> Dict[str, Any]:
        """Generate pipeline result summary."""
        summary = {
            "success": result.success,
            "processing_time": f"{result.processing_time_seconds:.1f}s",
            "errors_count": len(result.errors),
            "warnings_count": len(result.warnings)
        }

        if result.extraction_result:
            summary["items_extracted"] = len(result.extraction_result.line_items)
            summary["extraction_confidence"] = f"{result.extraction_result.overall_confidence:.1%}"

        if result.bucket_result:
            summary["auto_populated"] = len(result.bucket_result.bucket1_items)
            summary["manual_review"] = len(result.bucket_result.bucket2_items)
            summary["auto_populate_rate"] = result.bucket_result.routing_summary.get('auto_populate_rate', '0%')

        if result.qa_result:
            summary["qa_passed"] = result.qa_result.passed
            summary["qa_bucket"] = result.qa_result.bucket.value

        summary["outputs"] = {
            "excel": result.excel_path or "None",
            "review_report": result.review_report_path or "None",
            "audit_json": result.audit_json_path or "None"
        }

        return summary

    def get_cache_stats(self) -> Dict[str, Any]:
        """Get vendor cache statistics."""
        if self.cache:
            return self.cache.get_cache_stats()
        return {"enabled": False}
