"""
Review Reporter Module

Generates Word document reports for manual review items (Bucket 2).
Includes source context, confidence scores, and suggested actions.
"""

import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .quality_assurance import QAResult, ValidationStatus
from .bucket_router import BucketResult
from ..config import OUTPUT_CONFIG

logger = logging.getLogger(__name__)


class ReviewReporter:
    """
    Generates review reports for Bucket 2 items.

    Creates a Word document with:
    - Executive summary
    - Items requiring review (grouped by action type)
    - Source context for each item
    - Suggested actions
    - QA check results
    """

    def __init__(self, output_config: Optional[Dict] = None):
        """
        Initialize the review reporter.

        Args:
            output_config: Output configuration override
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx package not installed. "
                "Install with: pip install python-docx"
            )

        self.config = output_config or OUTPUT_CONFIG
        self.include_source = self.config.get('include_source_context', True)
        self.include_suggestions = self.config.get('include_suggested_actions', True)

        logger.info("ReviewReporter initialized")

    def generate_report(
        self,
        bucket_result: BucketResult,
        qa_result: QAResult,
        extraction_context: Optional[Dict] = None,
        output_path: Optional[str] = None
    ) -> str:
        """
        Generate a Word document review report.

        Args:
            bucket_result: Bucket routing results
            qa_result: QA validation results
            extraction_context: Optional context from extraction (vendor, doc type)
            output_path: Output file path (default: review_report_TIMESTAMP.docx)

        Returns:
            Path to generated report
        """
        doc = Document()
        context = extraction_context or {}

        # Set up document
        self._setup_document_styles(doc)

        # Title
        title = doc.add_heading('TCO Extraction Review Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if context.get('vendor_name'):
            doc.add_paragraph(f"Vendor: {context.get('vendor_name')}")
        if context.get('document_type'):
            doc.add_paragraph(f"Document Type: {context.get('document_type')}")

        doc.add_paragraph()

        # Executive Summary
        self._add_executive_summary(doc, bucket_result, qa_result)

        # Routing Summary
        self._add_routing_summary(doc, bucket_result)

        # Items Requiring Review
        if bucket_result.bucket2_items:
            self._add_review_items_section(doc, bucket_result)

        # QA Check Results
        self._add_qa_results_section(doc, qa_result)

        # Auto-Populated Items Summary
        if bucket_result.bucket1_items:
            self._add_auto_populated_summary(doc, bucket_result)

        # Suggested Actions
        if self.include_suggestions:
            self._add_suggested_actions(doc, bucket_result, qa_result)

        # Save document
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"review_report_{timestamp}.docx"

        doc.save(output_path)
        logger.info(f"Review report generated: {output_path}")

        return output_path

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up custom document styles."""
        # Modify default styles for better readability
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_executive_summary(
        self,
        doc: Document,
        bucket_result: BucketResult,
        qa_result: QAResult
    ) -> None:
        """Add executive summary section."""
        doc.add_heading('Executive Summary', level=1)

        summary = bucket_result.routing_summary
        total = summary.get('total_items', 0)
        auto = summary.get('bucket1_auto_populate', 0)
        review = summary.get('bucket2_manual_review', 0)
        confidence = qa_result.overall_confidence

        # Summary paragraph
        para = doc.add_paragraph()
        para.add_run(f"Extraction Analysis: ").bold = True
        para.add_run(
            f"Of {total} total line items extracted, "
            f"{auto} ({auto/total*100:.0f}%) can be auto-populated "
            f"and {review} ({review/total*100:.0f}%) require manual review."
        )

        # Confidence indicator
        para2 = doc.add_paragraph()
        para2.add_run(f"Overall Confidence: ").bold = True
        confidence_text = f"{confidence:.1%}"
        run = para2.add_run(confidence_text)
        if confidence >= 0.90:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif confidence >= 0.70:
            run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        # Time saved estimate
        time_saved = summary.get('estimated_time_saved', 'N/A')
        para3 = doc.add_paragraph()
        para3.add_run(f"Estimated Time Saved: ").bold = True
        para3.add_run(time_saved)

        doc.add_paragraph()

    def _add_routing_summary(
        self,
        doc: Document,
        bucket_result: BucketResult
    ) -> None:
        """Add routing summary table."""
        doc.add_heading('Routing Summary', level=1)

        summary = bucket_result.routing_summary
        action_breakdown = summary.get('action_breakdown', {})

        # Create summary table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        rows_data = [
            ('Total Items Extracted', str(summary.get('total_items', 0))),
            ('Bucket 1 (Auto-Populate)', f"{summary.get('bucket1_auto_populate', 0)} items"),
            ('Bucket 2 (Manual Review)', f"{summary.get('bucket2_manual_review', 0)} items"),
            ('Auto-Populate Rate', summary.get('auto_populate_rate', '0%')),
            ('QA Status', summary.get('qa_bucket', 'unknown'))
        ]

        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

        doc.add_paragraph()

        # Action breakdown
        if action_breakdown:
            doc.add_heading('Action Breakdown', level=2)
            for action, count in action_breakdown.items():
                doc.add_paragraph(f"  {action.replace('_', ' ').title()}: {count} items")

        doc.add_paragraph()

    def _add_review_items_section(
        self,
        doc: Document,
        bucket_result: BucketResult
    ) -> None:
        """Add section for items requiring manual review."""
        doc.add_heading('Items Requiring Review', level=1)

        # Group by action type
        quick_review = [
            i for i in bucket_result.bucket2_items
            if i.get('_routing', {}).get('action') == 'quick_review'
        ]
        manual_entry = [
            i for i in bucket_result.bucket2_items
            if i.get('_routing', {}).get('action') == 'manual_entry'
        ]

        # Quick Review Items
        if quick_review:
            doc.add_heading('Quick Review Required', level=2)
            doc.add_paragraph(
                'These items have moderate confidence and passed basic checks. '
                'Please verify the values are correct.'
            )
            self._add_items_table(doc, quick_review, include_source=self.include_source)

        # Manual Entry Items
        if manual_entry:
            doc.add_heading('Manual Entry Required', level=2)
            doc.add_paragraph(
                'These items have low confidence or failed QA checks. '
                'Values should be manually verified against the source document.'
            )
            self._add_items_table(doc, manual_entry, include_source=True)

    def _add_items_table(
        self,
        doc: Document,
        items: List[Dict],
        include_source: bool = True
    ) -> None:
        """Add a table of items."""
        if not items:
            return

        # Determine columns
        headers = ['Solution Name', 'Fee Type', 'Category', 'Amount', 'Confidence', 'Issue']
        if include_source:
            headers.append('Source')

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Data rows
        for item in items:
            row = table.add_row()
            cells = row.cells

            cells[0].text = item.get('solution_name', 'Unknown')
            cells[1].text = item.get('fee_type', '')
            cells[2].text = item.get('category', '')

            # Determine primary amount
            amount = (
                item.get('monthly_fee') or
                item.get('annual_fee') or
                item.get('one_time_fee') or
                item.get('per_unit_rate') or
                0
            )
            cells[3].text = f"${amount:,.2f}"

            # Confidence with color
            confidence = item.get('overall_confidence', 0)
            cells[4].text = f"{confidence:.0%}"

            # Issue/reason
            routing = item.get('_routing', {})
            cells[5].text = routing.get('reason', '')

            if include_source:
                source = item.get('source_location', '') or item.get('source_text', '')[:50]
                cells[6].text = source if source else 'N/A'

        doc.add_paragraph()

    def _add_qa_results_section(
        self,
        doc: Document,
        qa_result: QAResult
    ) -> None:
        """Add QA results section."""
        doc.add_heading('Quality Assurance Results', level=1)

        summary = qa_result.summary

        # Overall status
        para = doc.add_paragraph()
        para.add_run('Overall Status: ').bold = True
        status_text = 'PASSED' if qa_result.passed else 'REVIEW REQUIRED'
        run = para.add_run(status_text)
        run.font.color.rgb = RGBColor(0, 128, 0) if qa_result.passed else RGBColor(255, 165, 0)

        # Checks summary
        doc.add_paragraph(
            f"Checks Performed: {summary.get('total_checks', 0)} | "
            f"Passed: {summary.get('checks_passed', 0)} | "
            f"Warnings: {summary.get('checks_warning', 0)} | "
            f"Failed: {summary.get('checks_failed', 0)}"
        )

        # Failed/Warning checks detail
        problem_checks = [
            c for c in qa_result.checks
            if c.status in [ValidationStatus.FAILED, ValidationStatus.WARNING]
        ]

        if problem_checks:
            doc.add_heading('Issues Found', level=2)

            for check in problem_checks[:20]:  # Limit to first 20
                status_icon = '' if check.status == ValidationStatus.FAILED else ''
                para = doc.add_paragraph()
                para.add_run(f"{status_icon} {check.check_name}: ").bold = True
                para.add_run(check.message)
                if check.item_name:
                    para.add_run(f" (Item: {check.item_name})")

        doc.add_paragraph()

    def _add_auto_populated_summary(
        self,
        doc: Document,
        bucket_result: BucketResult
    ) -> None:
        """Add summary of auto-populated items."""
        doc.add_heading('Auto-Populated Items (Bucket 1)', level=1)

        doc.add_paragraph(
            f"{len(bucket_result.bucket1_items)} items have been automatically "
            "populated to the Excel template with high confidence."
        )

        # Just show count by category
        category_counts = {}
        for item in bucket_result.bucket1_items:
            cat = item.get('category', 'Unknown')
            category_counts[cat] = category_counts.get(cat, 0) + 1

        if category_counts:
            doc.add_heading('By Category', level=2)
            for cat, count in sorted(category_counts.items()):
                doc.add_paragraph(f"  {cat}: {count} items")

        doc.add_paragraph()

    def _add_suggested_actions(
        self,
        doc: Document,
        bucket_result: BucketResult,
        qa_result: QAResult
    ) -> None:
        """Add suggested actions section."""
        doc.add_heading('Suggested Actions', level=1)

        actions = []

        # Based on bucket2 items
        quick_review_count = sum(
            1 for i in bucket_result.bucket2_items
            if i.get('_routing', {}).get('action') == 'quick_review'
        )
        manual_count = sum(
            1 for i in bucket_result.bucket2_items
            if i.get('_routing', {}).get('action') == 'manual_entry'
        )

        if quick_review_count > 0:
            actions.append(
                f"1. Review {quick_review_count} flagged items in the 'Quick Review' section. "
                "These likely have correct values but should be verified."
            )

        if manual_count > 0:
            actions.append(
                f"2. Manually enter or verify {manual_count} items in the 'Manual Entry' section. "
                "Reference the source document for accurate values."
            )

        # Based on QA failures
        failed_checks = [c for c in qa_result.checks if c.status == ValidationStatus.FAILED]
        if failed_checks:
            unique_issues = set(c.check_name for c in failed_checks)
            actions.append(
                f"3. Address {len(unique_issues)} QA issue type(s): {', '.join(list(unique_issues)[:3])}"
            )

        # General recommendations
        if qa_result.overall_confidence < 0.80:
            actions.append(
                "4. Consider re-running extraction with additional document context "
                "or checking document quality."
            )

        if not actions:
            actions.append("No additional actions required. All items processed successfully.")

        for action in actions:
            doc.add_paragraph(action)

        doc.add_paragraph()

    def generate_summary_only(
        self,
        bucket_result: BucketResult,
        qa_result: QAResult
    ) -> Dict[str, Any]:
        """
        Generate just the summary data without creating a document.

        Useful for logging or API responses.
        """
        return {
            "routing_summary": bucket_result.routing_summary,
            "qa_summary": qa_result.summary,
            "bucket1_count": len(bucket_result.bucket1_items),
            "bucket2_count": len(bucket_result.bucket2_items),
            "overall_confidence": qa_result.overall_confidence,
            "passed": qa_result.passed
        }
